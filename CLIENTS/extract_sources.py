#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
extract_sources.py
------------------
Extraction rapide de texte brut depuis un dossier client (PDF/DOCX/TXT).
- Ignore les images (on récupère uniquement le texte existant).
- Sort un JSON avec texte + métadonnées.
- Optionnel: conversion de formats Office legacy (DOC/RTF/ODT...) via LibreOffice (soffice).

Install:
  pip install pymupdf python-docx

Usage:
  python extract_sources.py --input "/path/to/client_folder" --out "out/extracted.json"
  python extract_sources.py --input "/path/to/client_folder" --out "out/extracted.json" --out-text "out/corpus.txt"
  python extract_sources.py --input "/path/to/client_folder" --out "out/extracted.json" --enable-soffice
"""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict

import fitz  # PyMuPDF
from docx import Document  # python-docx


# -------------------------
# Config / Logging
# -------------------------

LOG = logging.getLogger("extract_sources")

SUPPORTED_DIRECT = {".pdf", ".docx", ".txt"}
SUPPORTED_SOFFICE = {".doc", ".rtf", ".odt", ".docm", ".dot", ".dotx", ".dotm"}  # si --enable-soffice


# -------------------------
# Utils
# -------------------------

def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()

def normalize_text(text: str) -> str:
    # Nettoyage léger (sans "intelligence") pour rendre le corpus propre
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # réduire les gros paquets de lignes vides
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def file_mtime_iso(path: Path) -> str:
    try:
        ts = path.stat().st_mtime
        return datetime.fromtimestamp(ts).isoformat(timespec="seconds")
    except Exception:
        return ""


# -------------------------
# Extractors
# -------------------------

def extract_pdf_pymupdf(path: Path) -> Dict:
    """Extract texte brut PDF (ignore images). Retourne {text, pages}."""
    pages_text = []
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, start=1):
            t = page.get_text("text") or ""
            t = normalize_text(t)
            pages_text.append({"page": i, "text": t})

    full_text = "\n\n".join([p["text"] for p in pages_text if p["text"]]).strip()
    return {"text": full_text, "pages": pages_text}

def extract_docx_python_docx(path: Path) -> Dict:
    """Extract texte brut DOCX : paragraphes + tableaux (ignore images)."""
    doc = Document(path)
    parts: List[str] = []

    # Paragraphes
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Tableaux (important dans beaucoup de docs)
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    row_cells.append(ct)
            if row_cells:
                parts.append(" | ".join(row_cells))

    full_text = normalize_text("\n".join(parts))
    return {"text": full_text, "pages": None}

def extract_txt(path: Path) -> Dict:
    """Lecture TXT brute (avec fallback encodage)."""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
    return {"text": normalize_text(text), "pages": None}


def _soffice_available() -> Optional[str]:
    return shutil.which("soffice")

def extract_via_soffice_to_txt(path: Path, soffice_bin: str) -> Dict:
    """
    Convertit un fichier (DOC/RTF/ODT/...) en TXT via LibreOffice, puis lit le TXT.
    C'est un fallback (moins "pur Python" mais très pratique pour .doc legacy).
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        # LibreOffice écrit dans --outdir
        cmd = [
            soffice_bin,
            "--headless",
            "--nologo",
            "--nolockcheck",
            "--nodefault",
            "--nofirststartwizard",
            "--convert-to", "txt:Text (encoded):UTF8",
            "--outdir", str(tmpdir_path),
            str(path),
        ]
        LOG.debug("SOFFICE CMD: %s", " ".join(cmd))
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if proc.returncode != 0:
            raise RuntimeError(f"soffice failed ({proc.returncode}): {proc.stderr.strip()[:500]}")

        # Trouver le fichier txt généré
        # Exemple: "document.doc" -> "document.txt"
        out_txt = tmpdir_path / (path.stem + ".txt")
        if not out_txt.exists():
            # fallback: prend le 1er .txt trouvé
            candidates = list(tmpdir_path.glob("*.txt"))
            if not candidates:
                raise RuntimeError("soffice conversion produced no .txt output")
            out_txt = candidates[0]

        text = out_txt.read_text(encoding="utf-8", errors="ignore")
        return {"text": normalize_text(text), "pages": None}


# -------------------------
# Data model
# -------------------------

@dataclass
class ExtractedDoc:
    path: str
    ext: str
    size_bytes: int
    mtime_iso: str
    extractor: str
    text: str
    text_sha256: str
    pages: Optional[List[Dict]] = None
    error: Optional[str] = None


# -------------------------
# Main pipeline
# -------------------------

def extract_one(path: Path, enable_soffice: bool) -> ExtractedDoc:
    ext = path.suffix.lower()
    size = path.stat().st_size if path.exists() else 0

    try:
        if ext == ".pdf":
            res = extract_pdf_pymupdf(path)
            text = res["text"]
            pages = res["pages"]
            extractor = "pymupdf"
        elif ext == ".docx":
            res = extract_docx_python_docx(path)
            text = res["text"]
            pages = None
            extractor = "python-docx"
        elif ext == ".txt":
            res = extract_txt(path)
            text = res["text"]
            pages = None
            extractor = "txt"
        elif enable_soffice and ext in SUPPORTED_SOFFICE:
            soffice_bin = _soffice_available()
            if not soffice_bin:
                raise RuntimeError("soffice not found (install LibreOffice or disable --enable-soffice)")
            res = extract_via_soffice_to_txt(path, soffice_bin)
            text = res["text"]
            pages = None
            extractor = "soffice->txt"
        else:
            raise ValueError(f"Unsupported extension: {ext}")

        return ExtractedDoc(
            path=str(path),
            ext=ext,
            size_bytes=size,
            mtime_iso=file_mtime_iso(path),
            extractor=extractor,
            text=text,
            text_sha256=sha256_text(text),
            pages=pages,
        )

    except Exception as e:
        return ExtractedDoc(
            path=str(path),
            ext=ext,
            size_bytes=size,
            mtime_iso=file_mtime_iso(path),
            extractor="error",
            text="",
            text_sha256=sha256_text(""),
            pages=None,
            error=str(e),
        )


def walk_files(root: Path) -> List[Path]:
    files: List[Path] = []
    for p in root.rglob("*"):
        if p.is_file():
            files.append(p)
    return sorted(files)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Dossier à analyser (client)")
    parser.add_argument("--out", required=True, help="Chemin JSON de sortie")
    parser.add_argument("--out-text", default=None, help="Optionnel: concat du texte brut (corpus .txt)")
    parser.add_argument("--enable-soffice", action="store_true", help="Activer conversion DOC/RTF/ODT via LibreOffice")
    parser.add_argument("--verbose", action="store_true", help="Logs détaillés")
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)s - %(message)s",
    )

    root = Path(args.input).expanduser().resolve()
    out_path = Path(args.out).expanduser().resolve()
    out_text = Path(args.out_text).expanduser().resolve() if args.out_text else None

    if not root.exists() or not root.is_dir():
        LOG.error("Input folder not found: %s", root)
        return 2

    files = walk_files(root)
    LOG.info("Found %d files under %s", len(files), root)

    extracted: List[ExtractedDoc] = []
    ok = 0
    skipped = 0

    for p in files:
        ext = p.suffix.lower()
        # Filtrage rapide : on prend direct / ou si soffice activé
        if ext in SUPPORTED_DIRECT or (args.enable_soffice and ext in SUPPORTED_SOFFICE):
            doc = extract_one(p, enable_soffice=args.enable_soffice)
            extracted.append(doc)
            if doc.error:
                LOG.warning("FAIL %s -> %s", p.name, doc.error)
            else:
                ok += 1
                LOG.info("OK   %s (%s) [%s chars]", p.name, doc.extractor, len(doc.text))
        else:
            skipped += 1

    payload = {
        "root": str(root),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "enable_soffice": bool(args.enable_soffice),
        "counts": {"ok": ok, "errors": sum(1 for d in extracted if d.error), "skipped": skipped, "total_seen": len(files)},
        "documents": [asdict(d) for d in extracted],
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    LOG.info("Wrote JSON -> %s", out_path)

    if out_text:
        out_text.parent.mkdir(parents=True, exist_ok=True)
        # Concat simple : séparation claire par fichier
        chunks = []
        for d in extracted:
            if d.text:
                chunks.append(f"\n\n===== FILE: {d.path} =====\n{d.text}\n")
        out_text.write_text("".join(chunks).strip() + "\n", encoding="utf-8")
        LOG.info("Wrote corpus -> %s", out_text)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
