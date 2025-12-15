#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Génère les champs du rapport en lots (format CHAMP=VALEUR)."""

from __future__ import annotations

import argparse
import json
import re
import textwrap
import urllib.request
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

from docx import Document

# On réutilise tes fonctions existantes de build_context.py
# (Ton build_context.py ne s'exécute pas à l'import grâce au if __name__ == "__main__")
from build_context import BM25Index, make_chunks


# -----------------------------
DEFAULT_FIELDS = [
    {
        "key": "PROFESSION",
        "query": "Profession actuelle, expériences, postes occupés, responsabilités",
        "instructions": "Synthétise les informations clés sur la situation professionnelle.",
    },
    {
        "key": "FORMATION",
        "query": "Formations, diplômes, certifications, parcours scolaire",
        "instructions": "Présente les formations et diplômes mentionnés dans les sources.",
    },
    {
        "key": "DISCUSSION_ASSURE",
        "query": "Résultats de la discussion avec l'assuré : motivations, freins, points d'appui",
        "instructions": "Résumé structuré de la discussion avec l'assuré.",
    },
    {
        "key": "ORIENTATION",
        "query": "Orientation, pistes métiers, projet professionnel, recommandations",
        "instructions": "Propose des orientations cohérentes basées uniquement sur les sources.",
    },
    {
        "key": "STAGE",
        "query": "Stage, immersion, objectifs, résultats, suite recommandée",
        "instructions": "Décris le stage s'il existe, sinon indique l'absence d'information.",
    },
    {
        "key": "CONCLUSION",
        "query": "Conclusion globale, résumé final, prochaines étapes",
        "instructions": "Synthèse finale en quelques phrases.",
    },
]

DEFAULT_RULES = textwrap.dedent(
    """
    FR uniquement.
    Format obligatoire : une ligne par champ CHAMP=VALEUR (pas de puces, pas de code).
    Valeur limitée à 50 caractères maximum.
    Si l'information est introuvable : CHAMP= | WHY: explication courte.
    Interdit : JSON, Markdown, backticks, balises de code.
    Tu ne renvoies rien d'autre que ces lignes.
    """
).strip()

PLACEHOLDER_RE = re.compile(r"\{\{([^{}]+)\}\}")


def ollama_generate(model: str, prompt: str, host: str, temperature: float, top_p: float) -> str:
    url = host.rstrip("/") + "/api/generate"
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": temperature, "top_p": top_p},
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=300) as resp:
        out = json.loads(resp.read().decode("utf-8"))
    return out.get("response", "")


def extract_placeholders_from_docx(template_path: Path) -> List[str]:
    doc = Document(str(template_path))
    placeholders: List[str] = []

    def register(text: str) -> None:
        if not text:
            return
        for match in PLACEHOLDER_RE.findall(text):
            key = match.strip()
            if key and key not in placeholders:
                placeholders.append(key)

    for paragraph in doc.paragraphs:
        register(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    register(paragraph.text)

    return placeholders


def build_field_specs(placeholders: Sequence[str], fallback_defs: Optional[Sequence[Dict[str, Any]]] = None) -> List[Dict[str, Any]]:
    fallback_lookup = {item["key"]: item for item in (fallback_defs or []) if item.get("key")}
    default_lookup = {item["key"]: item for item in DEFAULT_FIELDS}
    specs: List[Dict[str, Any]] = []
    seen: set[str] = set()

    for raw in placeholders:
        key = raw.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        base = fallback_lookup.get(key) or default_lookup.get(key)
        query = base.get("query") if base else key.replace("_", " ")
        instructions = base.get("instructions") if base else f"Synthétise les informations pertinentes pour « {key} »"
        specs.append({"key": key, "query": query, "instructions": instructions})
    return specs


def path_allowed(p: str, include: Optional[Sequence[str]], exclude: Optional[Sequence[str]]) -> bool:
    low = p.lower()
    if exclude:
        for ex in exclude:
            if ex and ex.lower() in low:
                return False
    if include:
        return any(inc.lower() in low for inc in include if inc)
    return True


def chunked(items: Sequence[Dict[str, Any]], size: int) -> Iterable[List[Dict[str, Any]]]:
    size = max(1, size)
    for idx in range(0, len(items), size):
        yield list(items[idx : idx + size])


def collect_contexts(
    specs: Sequence[Dict[str, Any]],
    index: BM25Index,
    chunks,
    topk: int,
    debug_path: Optional[Path] = None,
) -> Dict[str, List[Dict[str, Any]]]:
    contexts: Dict[str, List[Dict[str, Any]]] = {}
    for spec in specs:
        top_hits = index.topk(spec["query"], topk)
        entries: List[Dict[str, Any]] = []
        for chunk_idx, score in top_hits:
            ch = chunks[chunk_idx]
            entries.append(
                {
                    "chunk_id": ch.chunk_id,
                    "source_path": ch.source_path,
                    "page": ch.page,
                    "text": ch.text,
                    "score": score,
                }
            )
        contexts[spec["key"]] = entries
        if debug_path:
            (debug_path / f"debug_{spec['key']}.json").write_text(
                json.dumps({"field": spec["key"], "context": entries}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
    return contexts


def truncate_value(value: str, max_len: int) -> str:
    value = (value or "").strip()
    if max_len <= 0 or len(value) <= max_len:
        return value
    return value[: max(1, max_len - 1)].rstrip() + "…"


def build_batch_prompt(
    batch_specs: Sequence[Dict[str, Any]],
    contexts: Dict[str, List[Dict[str, Any]]],
    rules_text: str,
    max_len: int,
) -> str:
    lines = [
        "Tu es un analyste francophone. Tu rédiges exclusivement en français.",
        f"Valeur maximale autorisée : {max_len} caractères par champ.",
        "Règles absolues :",
    ]
    for rule in rules_text.splitlines():
        rule = rule.strip()
        if rule:
            lines.append(f"- {rule}")
    lines.append("")
    lines.append("Utilise UNIQUEMENT les extraits fournis pour chaque champ.")

    for spec in batch_specs:
        lines.append(f"=== Champ : {spec['key']} ===")
        lines.append(f"Consigne : {spec['instructions']}")
        ctx_list = contexts.get(spec["key"], [])
        if not ctx_list:
            lines.append("(Aucun passage pertinent n'a été trouvé.)")
        else:
            lines.append("Passages autorisés :")
            for idx, ctx in enumerate(ctx_list, start=1):
                snippet = " ".join(ctx["text"].split())
                if len(snippet) > 400:
                    snippet = snippet[:400] + "…"
                source_name = Path(ctx["source_path"]).name
                reference = f"[{idx}] {source_name}"
                if ctx.get("page") is not None:
                    reference += f" (page {ctx['page']})"
                lines.append(f"{reference} : {snippet}")
        lines.append("")

    lines.append("Réponds maintenant en respectant strictement le format CHAMP=VALEUR.")
    return "\n".join(lines).strip()


def parse_llm_response(
    raw: str,
    expected_fields: Sequence[str],
    *,
    max_len: int,
) -> Dict[str, Dict[str, Optional[str]]]:
    results: Dict[str, Dict[str, Optional[str]]] = {key: {"value": "", "why": None} for key in expected_fields}
    field_set = set(expected_fields)
    for line in raw.splitlines():
        if "=" not in line:
            continue
        left, right = line.split("=", 1)
        key = left.strip()
        if key not in field_set:
            continue
        value = right.strip()
        why = None
        if "|" in value:
            value_part, _, remainder = value.partition("|")
            value = value_part.strip()
            if remainder.strip().upper().startswith("WHY:"):
                why = remainder.split("WHY:", 1)[1].strip()
        results[key]["value"] = truncate_value(value, max_len)
        if why:
            results[key]["why"] = why
    return results


def generate_fields_from_payload(
    payload: Dict[str, Any],
    *,
    template_fields: Sequence[str],
    model: str,
    host: str,
    topk: int,
    temperature: float,
    top_p: float,
    batch_size: int,
    max_value_len: int,
    rules_text: str,
    exclude: Optional[Sequence[str]] = None,
    include: Optional[Sequence[str]] = None,
    debug_dir: Optional[Path] = None,
    output_path: Optional[Path] = None,
    missing_debug_path: Optional[Path] = None,
    custom_fields: Optional[Sequence[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    if not template_fields:
        raise ValueError("Aucun champ {{...}} trouvé dans le template DOCX.")

    specs = build_field_specs(template_fields, custom_fields)

    chunks = make_chunks(payload, chunk_size=1200, overlap=200)
    filtered_chunks = [c for c in chunks if path_allowed(c.source_path, include, exclude)]
    if not filtered_chunks:
        raise RuntimeError("Aucun passage exploitable après filtrage include/exclude.")
    index = BM25Index(filtered_chunks)

    debug_path = debug_dir.expanduser().resolve() if debug_dir else None
    if debug_path:
        debug_path.mkdir(parents=True, exist_ok=True)

    contexts = collect_contexts(specs, index, filtered_chunks, topk, debug_path)

    answers: Dict[str, Any] = {}
    missing_reasons: Dict[str, Dict[str, Any]] = {}

    for batch in chunked(specs, batch_size):
        prompt = build_batch_prompt(batch, contexts, rules_text, max_value_len)
        raw = ollama_generate(model, prompt, host, temperature, top_p)
        parsed = parse_llm_response(raw, [spec["key"] for spec in batch], max_len=max_value_len)

        for spec in batch:
            data = parsed.get(spec["key"], {"value": "", "why": None})
            value = data.get("value") or ""
            why = data.get("why") or None
            if not value:
                if not why:
                    why = "Information introuvable ou absente des documents."
                missing_reasons[spec["key"]] = {"why": why, "instructions": spec["instructions"]}
            answers[spec["key"]] = {
                "field": spec["key"],
                "value": value,
                "why": why if not value else None,
                "sources_used": [ctx["chunk_id"] for ctx in contexts.get(spec["key"], [])],
            }
            print(f"OK: {spec['key']} (len={len(value)})")

    if output_path:
        out_path = output_path.expanduser().resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(json.dumps(answers, ensure_ascii=False, indent=2), encoding="utf-8")

    if missing_reasons:
        target = missing_debug_path or (output_path.expanduser().with_name("debug_missing.json") if output_path else None)
        if target:
            target_path = Path(target).expanduser().resolve()
            target_path.parent.mkdir(parents=True, exist_ok=True)
            target_path.write_text(json.dumps(missing_reasons, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"⚠️  Détails des WHY -> {target_path}")

    return answers


def main() -> int:
    ap = argparse.ArgumentParser(description="Génère les champs du rapport en lots (CHAMP=VALEUR)")
    ap.add_argument("--template", required=True, help="Template DOCX contenant les {{placeholders}}")
    ap.add_argument("--extracted", required=True, help="Payload JSON produit par extract_sources.py")
    ap.add_argument("--out", default="out/answers.json", help="Fichier answers.json de sortie")
    ap.add_argument("--missing-debug", default="", help="Fichier debug_missing.json (optionnel)")
    ap.add_argument("--model", default="mistral:latest", help="Modèle Ollama à utiliser")
    ap.add_argument("--host", default="http://localhost:11434", help="Hôte du serveur Ollama")
    ap.add_argument("--topk", type=int, default=8, help="Nombre de passages BM25 par champ")
    ap.add_argument("--batch-size", type=int, default=10, help="Nombre de champs traités par appel LLM")
    ap.add_argument("--max-len", type=int, default=50, help="Longueur maximale des valeurs retournées")
    ap.add_argument("--temperature", type=float, default=0.2)
    ap.add_argument("--top_p", type=float, default=0.9)
    ap.add_argument("--exclude", default="", help="Mots-clés chemins à exclure (séparés par virgules)")
    ap.add_argument("--include", default="", help="Mots-clés chemins à inclure (séparés par virgules)")
    ap.add_argument("--debug-dir", default="out/debug", help="Dossier de traces contextuelles")
    ap.add_argument("--fields", default="", help="JSON facultatif de configuration des champs")
    ap.add_argument("--rules-text", default=DEFAULT_RULES, help="Texte des règles imposées au LLM")
    args = ap.parse_args()

    template_path = Path(args.template).expanduser().resolve()
    placeholders = extract_placeholders_from_docx(template_path)
    if not placeholders:
        raise SystemExit("Aucun champ {{...}} détecté dans le template.")

    payload = json.loads(Path(args.extracted).expanduser().read_text(encoding="utf-8"))

    custom_fields = None
    if args.fields:
        custom_fields = json.loads(Path(args.fields).expanduser().read_text(encoding="utf-8"))

    exclude = [s.strip() for s in args.exclude.split(",") if s.strip()]
    include = [s.strip() for s in args.include.split(",") if s.strip()]

    missing_debug = Path(args.missing_debug).expanduser() if args.missing_debug else None

    generate_fields_from_payload(
        payload,
        template_fields=placeholders,
        model=args.model,
        host=args.host,
        topk=args.topk,
        temperature=args.temperature,
        top_p=args.top_p,
        batch_size=args.batch_size,
        max_value_len=args.max_len,
        rules_text=args.rules_text.strip() or DEFAULT_RULES,
        exclude=exclude,
        include=include,
        debug_dir=Path(args.debug_dir),
        output_path=Path(args.out),
        missing_debug_path=missing_debug,
        custom_fields=custom_fields,
    )

    print(f"\nOK: answers -> {Path(args.out).expanduser().resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
