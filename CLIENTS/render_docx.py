#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
render_docx.py
--------------
Injecte les champs générés (out/answers.json) dans le template DOCX
en remplaçant le "contenu variable" des sections.

Sections gérées (par défaut) :
- Profession
- Formation
- Résultats de la discussion avec l’assuré
- Orientation
- Stage
- Conclusion

Usage :
python3 render_docx.py \
  --template "../TEMPLATE_V1 2 2.docx" \
  --answers "out/answers.json" \
  --output "out/rapport_final.docx" \
  --name "Malik" --surname "Karaoui" --civility "Monsieur"

Option PDF (LibreOffice) :
soffice --headless --convert-to pdf --outdir out out/rapport_final.docx
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Optional, Dict, Any, List, Union

from docx import Document
from docx.text.paragraph import Paragraph


# ----------------------------
# Utils DOCX (delete/insert)
# ----------------------------

def _norm(s: str) -> str:
    """Normalise le texte pour matcher facilement (ignore NBSP, ponctuation simple, espaces)."""
    s = (s or "").replace("\u00a0", " ").strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = s.rstrip(":")
    return s

def _style_ok(p: Paragraph, must_start_with: Optional[List[str]] = None) -> bool:
    name = getattr(getattr(p, "style", None), "name", "") or ""
    if name.startswith("TOC"):  # évite Table of Contents
        return False
    if not must_start_with:
        return True
    return any(name.startswith(prefix) for prefix in must_start_with)

def find_paragraph(doc: Document, text: str, *, after: int = 0, style_prefixes: Optional[List[str]] = None) -> Optional[Paragraph]:
    target = _norm(text)
    for p in doc.paragraphs[after:]:
        if not _style_ok(p, style_prefixes):
            continue
        if _norm(p.text) == target:
            return p
    return None

def delete_paragraph(p: Paragraph) -> None:
    """Supprime un paragraphe (hack XML courant)."""
    el = p._element
    el.getparent().remove(el)

def insert_paragraph_after(p: Paragraph, text: str, style_name: Optional[str] = None) -> Paragraph:
    """Insère un paragraphe après p."""
    new_p = p._element.addnext(p._element.__class__())
    para = Paragraph(new_p, p._parent)
    if style_name:
        try:
            para.style = style_name
        except Exception:
            pass
    if text is not None:
        para.add_run(text)
    return para

def replace_text_everywhere(doc: Document, mapping: Dict[str, str]) -> None:
    """Remplace du texte simple partout (paragraphes + tableaux)."""

    def replace_in_paragraph(par: Paragraph):
        if not par.runs:
            text = par.text
        else:
            text = "".join(run.text for run in par.runs)

        replaced = False
        for old, new in mapping.items():
            if not old:
                continue
            if old not in text:
                continue
            text = text.replace(old, new)
            replaced = True

        if replaced:
            if par.runs:
                par.runs[0].text = text
                for r in par.runs[1:]:
                    r.text = ""
            else:
                par.text = text

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    replace_in_paragraph(par)


def build_moustache_mapping(answers: Dict[str, Any]) -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    for key, value in answers.items():
        if isinstance(value, dict):
            answer_text = (value.get("value") or value.get("answer") or "").strip()
        else:
            answer_text = str(value).strip()
        if not answer_text:
            continue
        placeholder = f"{{{{{key}}}}}"
        placeholder_lower = f"{{{{{key.lower()}}}}}"
        mapping[placeholder] = answer_text
        mapping.setdefault(placeholder_lower, answer_text)
    return mapping


# ----------------------------
# Remplacement de sections
# ----------------------------

def replace_section(
    doc: Document,
    *,
    start_text: str,
    end_text: str,
    answer_text: str,
    start_style_prefixes: List[str],
    end_style_prefixes: Optional[List[str]] = None,
    keep_end: bool = True,
) -> None:
    """
    Remplace tous les paragraphes entre start_text et end_text (exclus)
    par le contenu answer_text (split par lignes).
    """
    # (re)trouve start
    start_p = find_paragraph(doc, start_text, style_prefixes=start_style_prefixes)
    if not start_p:
        raise RuntimeError(f"Start section introuvable: '{start_text}'")

    # trouve end après start
    start_idx = doc.paragraphs.index(start_p)
    end_p = find_paragraph(doc, end_text, after=start_idx + 1, style_prefixes=end_style_prefixes)
    if not end_p:
        raise RuntimeError(f"End section introuvable après '{start_text}': '{end_text}'")

    end_idx = doc.paragraphs.index(end_p)

    # paragraphes à supprimer (entre start et end)
    between = doc.paragraphs[start_idx + 1 : end_idx]
    base_style = None
    for bp in between:
        if bp.text.strip():
            base_style = getattr(getattr(bp, "style", None), "name", None)
            break
    if not base_style:
        # si rien, on prend le style du premier "entre" si dispo
        if between:
            base_style = getattr(getattr(between[0], "style", None), "name", None)
    if not base_style:
        base_style = "Corps A A"  # fallback (chez toi c’est souvent ça)

    # supprime le contenu existant
    for bp in list(between):
        delete_paragraph(bp)

    # insère le nouveau contenu
    answer_text = (answer_text or "").strip()
    lines = [ln.rstrip() for ln in answer_text.splitlines()]
    lines = [ln for ln in lines if ln.strip()]

    cursor = start_p
    if not lines:
        # rien à mettre → on met au moins une ligne vide pour éviter section "collée"
        insert_paragraph_after(cursor, "", base_style)
        return

    for ln in lines:
        # normalise bullets
        if ln.startswith(("- ", "* ")):
            ln = "• " + ln[2:].strip()
        cursor = insert_paragraph_after(cursor, ln, base_style)


# ----------------------------
# Main
# ----------------------------

def render_report_from_answers(
    template: Union[str, Path],
    answers: Union[Dict[str, Any], str, Path],
    output: Union[str, Path],
    *,
    name: str = "",
    surname: str = "",
    civility: str = "Monsieur",
) -> Path:
    template_path = Path(template).expanduser().resolve()
    output_path = Path(output).expanduser().resolve()

    doc = Document(str(template_path))

    if isinstance(answers, (str, Path)):
        answers_dict: Dict[str, Any] = json.loads(Path(answers).expanduser().read_text(encoding="utf-8"))
    else:
        answers_dict = answers

    def get_answer(key: str) -> str:
        obj = answers_dict.get(key) or {}
        if isinstance(obj, dict):
            return (obj.get("value") or obj.get("answer") or "").strip()
        return str(obj).strip()

    mapping = {}
    if name:
        mapping["{NAME}"] = name
        mapping["{monsieur ou madame NAME}"] = f"{civility} {name}".strip()
    if surname:
        mapping["{surname}"] = surname
        mapping["XX"] = f"{civility} {surname}".strip()
    if mapping:
        replace_text_everywhere(doc, mapping)

    moustache_mapping = build_moustache_mapping(answers_dict)
    if moustache_mapping:
        replace_text_everywhere(doc, moustache_mapping)

    replace_section(
        doc,
        start_text="Profession",
        end_text="Formation",
        answer_text=get_answer("PROFESSION"),
        start_style_prefixes=["TITRE 2"],
        end_style_prefixes=["TITRE 2"],
    )

    replace_section(
        doc,
        start_text="Formation",
        end_text="Tests",
        answer_text=get_answer("FORMATION"),
        start_style_prefixes=["TITRE 2"],
        end_style_prefixes=["Heading 1"],
    )

    replace_section(
        doc,
        start_text="RÉSULTATS DE LA DISCUSSION AVEC L’ASSURÉ",
        end_text="Compétences Professionnelles & Sociales",
        answer_text=get_answer("DISCUSSION_ASSURE"),
        start_style_prefixes=["TITRE 2"],
        end_style_prefixes=["Heading 1"],
    )

    replace_section(
        doc,
        start_text="Orientation",
        end_text="Stage",
        answer_text=get_answer("ORIENTATION"),
        start_style_prefixes=["TITRE 2"],
        end_style_prefixes=["TITRE 2"],
    )

    replace_section(
        doc,
        start_text="Stage",
        end_text="Formation",
        answer_text=get_answer("STAGE"),
        start_style_prefixes=["TITRE 2"],
        end_style_prefixes=["TITRE 2"],
    )

    replace_section(
        doc,
        start_text="Conclusion",
        end_text="Lieu & Date",
        answer_text=get_answer("CONCLUSION"),
        start_style_prefixes=["Heading 1"],
        end_style_prefixes=None,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True, help="Chemin du template DOCX")
    ap.add_argument("--answers", required=True, help="Chemin out/answers.json")
    ap.add_argument("--output", required=True, help="Chemin de sortie DOCX")
    ap.add_argument("--name", default="", help="Prénom (pour placeholders)")
    ap.add_argument("--surname", default="", help="Nom (pour placeholders)")
    ap.add_argument("--civility", default="Monsieur", help="Monsieur/Madame")
    args = ap.parse_args()

    render_report_from_answers(
        template=args.template,
        answers=args.answers,
        output=args.output,
        name=args.name,
        surname=args.surname,
        civility=args.civility,
    )
    print(f"OK: rendu DOCX -> {Path(args.output).expanduser().resolve()}")

if __name__ == "__main__":
    main()
