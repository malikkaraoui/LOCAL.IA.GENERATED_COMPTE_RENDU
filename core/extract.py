"""Extraction des sources client (TXT/PDF/DOCX/MSG...)."""

from __future__ import annotations

import hashlib
import json
import logging
import re
import shutil
import subprocess
import tempfile
from collections.abc import Sequence
from datetime import datetime
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document

from .errors import Result, ExtractionError
from .logger import get_logger
from .models import SourceDoc

# ✅ Import lazy du module .msg
try:
    from .extractors.msg_extractor import MSG_SUPPORT_AVAILABLE, extract_msg_safe
except ImportError:
    MSG_SUPPORT_AVAILABLE = False
    extract_msg_safe = None  # type: ignore

LOG = get_logger("core.extract")

SUPPORTED_DIRECT = {".pdf", ".docx", ".txt"}
SUPPORTED_SOFFICE = {".doc", ".rtf", ".odt", ".docm", ".dot", ".dotx", ".dotm"}
SUPPORTED_MSG = {".msg"}  # ✅ Emails Outlook


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def file_mtime_iso(path: Path) -> str:
    try:
        return datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
    except Exception:
        return ""


def extract_pdf(path: Path) -> Result[dict]:
    """Extrait le texte d'un PDF avec PyMuPDF.
    
    Args:
        path: Chemin vers le fichier PDF
        
    Returns:
        Result[dict]: Succès avec {"text", "pages"} ou échec avec ExtractionError
    """
    try:
        LOG.info("Extraction PDF: %s", path.name)
        pages_text = []
        with fitz.open(path) as doc:
            for i, page in enumerate(doc, start=1):
                t = page.get_text("text") or ""
                t = normalize_text(t)
                pages_text.append({"page": i, "text": t})
        full_text = "\n\n".join(p["text"] for p in pages_text if p["text"]).strip()
        LOG.debug("PDF extrait: %d pages, %d caractères", len(pages_text), len(full_text))
        return Result.ok({"text": full_text, "pages": pages_text})
    except Exception as exc:
        error = ExtractionError(f"Échec extraction PDF {path.name}: {exc}")
        LOG.error("Erreur extraction PDF %s: %s", path.name, exc)
        return Result.fail(error)


def extract_docx(path: Path) -> Result[dict]:
    """Extrait le texte d'un DOCX avec python-docx.
    
    Args:
        path: Chemin vers le fichier DOCX
        
    Returns:
        Result[dict]: Succès avec {"text", "pages": None} ou échec avec ExtractionError
    """
    try:
        LOG.info("Extraction DOCX: %s", path.name)
        doc = Document(path)
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        full_text = normalize_text("\n".join(parts))
        LOG.debug("DOCX extrait: %d paragraphes, %d caractères", len(parts), len(full_text))
        return Result.ok({"text": full_text, "pages": None})
    except Exception as exc:
        error = ExtractionError(f"Échec extraction DOCX {path.name}: {exc}")
        LOG.error("Erreur extraction DOCX %s: %s", path.name, exc)
        return Result.fail(error)


def extract_txt(path: Path) -> Result[dict]:
    """Extrait le texte d'un fichier TXT avec détection d'encodage.
    
    Args:
        path: Chemin vers le fichier TXT
        
    Returns:
        Result[dict]: Succès avec {"text", "pages": None} ou échec avec ExtractionError
    """
    try:
        LOG.info("Extraction TXT: %s", path.name)
        try:
            text = path.read_text(encoding="utf-8")
            encoding = "utf-8"
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
            encoding = "latin-1"
        normalized = normalize_text(text)
        LOG.debug("TXT extrait (%s): %d caractères", encoding, len(normalized))
        return Result.ok({"text": normalized, "pages": None})
    except Exception as exc:
        error = ExtractionError(f"Échec extraction TXT {path.name}: {exc}")
        LOG.error("Erreur extraction TXT %s: %s", path.name, exc)
        return Result.fail(error)


def extract_msg(path: Path, output_dir: Optional[Path] = None) -> Result[dict]:
    """Extrait le contenu d'un email Outlook .msg.
    
    Args:
        path: Chemin vers le fichier .msg
        output_dir: Dossier pour extraire les pièces jointes (optionnel)
        
    Returns:
        Result[dict]: Succès avec {"text", "pages": None, "meta"} ou échec avec ExtractionError
        
    Notes:
        - Requiert extract-msg installé (pip install extract-msg)
        - Extrait automatiquement les pièces jointes PDF/DOCX/DOC/TXT dans output_dir si fourni
        - Format texte : [EMAIL_MSG] Subject/From/To/Date/Attachments + Body
    """
    if not MSG_SUPPORT_AVAILABLE or extract_msg_safe is None:
        error = ExtractionError("extract-msg non installé (pip install extract-msg>=0.48.0)")
        LOG.warning("Tentative extraction .msg sans extract-msg : %s", path.name)
        return Result.fail(error)
    
    try:
        LOG.info("Extraction MSG: %s", path.name)
        text, meta, error = extract_msg_safe(path, output_dir)
        
        if error:
            return Result.fail(ExtractionError(error))
        
        if text is None:
            return Result.fail(ExtractionError("Extraction .msg retourné None"))
        
        LOG.debug(
            "MSG extrait: %d caractères, %d pièces jointes",
            len(text),
            meta.get("attachments_count", 0) if meta else 0,
        )
        
        return Result.ok({
            "text": text,
            "pages": None,
            "meta": meta,
        })
        
    except Exception as exc:
        error = ExtractionError(f"Échec extraction MSG {path.name}: {exc}")
        LOG.error("Erreur extraction MSG %s: %s", path.name, exc)
        return Result.fail(error)


def soffice_available() -> Optional[str]:
    return shutil.which("soffice")


def extract_via_soffice(path: Path, soffice_bin: str) -> Result[dict]:
    """Extrait le texte via LibreOffice en convertissant en TXT.
    
    Args:
        path: Chemin vers le fichier à convertir
        soffice_bin: Chemin vers l'exécutable soffice
        
    Returns:
        Result[dict]: Succès avec {"text", "pages": None} ou échec avec ExtractionError
    """
    try:
        LOG.info("Extraction soffice: %s", path.name)
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            cmd = [
                soffice_bin,
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--nodefault",
                "--nofirststartwizard",
                "--convert-to",
                "txt:Text (encoded):UTF8",
                "--outdir",
                str(tmpdir_path),
                str(path),
            ]
            proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if proc.returncode != 0:
                error_msg = f"soffice failed ({proc.returncode}): {proc.stderr.strip()[:500]}"
                LOG.error("Échec soffice %s: %s", path.name, error_msg)
                return Result.fail(ExtractionError(error_msg))
            out_txt = tmpdir_path / (path.stem + ".txt")
            if not out_txt.exists():
                candidates = list(tmpdir_path.glob("*.txt"))
                if not candidates:
                    error_msg = "soffice conversion produced no .txt output"
                    LOG.error("Aucun output soffice pour %s", path.name)
                    return Result.fail(ExtractionError(error_msg))
                out_txt = candidates[0]
            text = out_txt.read_text(encoding="utf-8", errors="ignore")
            normalized = normalize_text(text)
            LOG.debug("Soffice extrait: %d caractères", len(normalized))
            return Result.ok({"text": normalized, "pages": None})
    except Exception as exc:
        error = ExtractionError(f"Échec extraction soffice {path.name}: {exc}")
        LOG.error("Erreur soffice %s: %s", path.name, exc)
        return Result.fail(error)


def walk_files(root: Path) -> list[Path]:
    """Liste récursivement tous les fichiers, en excluant les fichiers/dossiers cachés."""
    def _is_hidden(path: Path) -> bool:
        """Retourne True si le fichier ou un de ses parents est caché."""
        for part in path.parts:
            if part.startswith('.'):
                return True
        return False

    return sorted([p for p in root.rglob("*") if p.is_file() and not _is_hidden(p.relative_to(root))])


def extract_sources(
    root: Path,
    *,
    enable_soffice: bool = False,
    enable_msg: bool = True,  # ✅ Support .msg activé par défaut
    include_extensions: Optional[Sequence[str]] = None,
    msg_attachments_dir: Optional[Path] = None,  # ✅ Dossier pour extraire pièces jointes .msg
) -> dict:
    """Extrait le texte de tous les documents d'un dossier.
    
    Args:
        root: Dossier racine contenant les documents
        enable_soffice: Activer conversion via LibreOffice
        enable_msg: Activer extraction emails .msg (défaut: True)
        include_extensions: Liste d'extensions à inclure (None = toutes supportées)
        msg_attachments_dir: Dossier pour extraire pièces jointes des .msg
        
    Returns:
        dict: Payload avec counts et liste des documents extraits
    """
    root = Path(root).expanduser().resolve()
    if not root.exists() or not root.is_dir():
        raise FileNotFoundError(f"Dossier introuvable: {root}")

    allow_exts = set(include_extensions or [])
    files = walk_files(root)
    documents: list[SourceDoc] = []
    ok = errors = skipped = 0
    
    # ✅ Liste des fichiers additionnels extraits depuis pièces jointes .msg
    additional_files: list[Path] = []

    for path in files:
        ext = path.suffix.lower()
        allowed = not allow_exts or ext in allow_exts
        
        # ✅ Vérifier support (inclut .msg si enable_msg=True)
        supported = (
            ext in SUPPORTED_DIRECT
            or (enable_soffice and ext in SUPPORTED_SOFFICE)
            or (enable_msg and ext in SUPPORTED_MSG)
        )
        
        if not supported or not allowed:
            skipped += 1
            continue

        result: Optional[Result[dict]] = None
        extractor = "unknown"
        
        try:
            if ext == ".pdf":
                result = extract_pdf(path)
                extractor = "pymupdf"
            elif ext == ".docx":
                result = extract_docx(path)
                extractor = "python-docx"
            elif ext == ".txt":
                result = extract_txt(path)
                extractor = "txt"
            elif enable_msg and ext == ".msg":
                # ✅ Extraction email .msg avec pièces jointes
                result = extract_msg(path, output_dir=msg_attachments_dir)
                extractor = "extract-msg"
                
                # ✅ Si pièces jointes extraites, les ajouter à la liste pour traitement
                if result and result.success and result.value.get("meta"):
                    att_paths = result.value["meta"].get("extracted_attachments_paths", [])
                    for att_path_str in att_paths:
                        att_path = Path(att_path_str)
                        if att_path.exists() and att_path not in additional_files:
                            additional_files.append(att_path)
                            LOG.info("Pièce jointe .msg à indexer : %s", att_path.name)
            elif enable_soffice and ext in SUPPORTED_SOFFICE:
                soffice_bin = soffice_available()
                if not soffice_bin:
                    result = Result.fail(ExtractionError("LibreOffice (soffice) non trouvé"))
                else:
                    result = extract_via_soffice(path, soffice_bin)
                extractor = "soffice->txt"
            else:
                skipped += 1
                continue
                
            if result and result.success:
                data = result.value
                doc = SourceDoc(
                    path=str(path),
                    ext=ext,
                    size_bytes=path.stat().st_size,
                    mtime_iso=file_mtime_iso(path),
                    extractor=extractor,
                    text=data["text"],
                    text_sha256=sha256_text(data["text"]),
                    pages=data.get("pages"),
                )
                documents.append(doc)
                ok += 1
            elif result:
                # Échec explicite avec Result.fail
                documents.append(
                    SourceDoc(
                        path=str(path),
                        ext=ext,
                        size_bytes=path.stat().st_size if path.exists() else 0,
                        mtime_iso=file_mtime_iso(path),
                        extractor="error",
                        text="",
                        text_sha256=sha256_text(""),
                        pages=None,
                        error=str(result.error),
                    )
                )
                errors += 1
        except Exception as exc:
            # Gestion des exceptions imprévues
            documents.append(
                SourceDoc(
                    path=str(path),
                    ext=ext,
                    size_bytes=path.stat().st_size if path.exists() else 0,
                    mtime_iso=file_mtime_iso(path),
                    extractor="error",
                    text="",
                    text_sha256=sha256_text(""),
                    pages=None,
                    error=str(exc),
                )
            )
            errors += 1
            LOG.warning("FAIL %s -> %s", path.name, exc)
    
    # ✅ Traiter les pièces jointes extraites des .msg
    if additional_files:
        LOG.info("Traitement de %d pièces jointes extraites depuis .msg", len(additional_files))
        for att_path in additional_files:
            ext = att_path.suffix.lower()
            result = None
            extractor = "msg-attachment"
            
            try:
                if ext == ".pdf":
                    result = extract_pdf(att_path)
                    extractor = "msg-attachment:pdf"
                elif ext == ".docx":
                    result = extract_docx(att_path)
                    extractor = "msg-attachment:docx"
                elif ext == ".txt":
                    result = extract_txt(att_path)
                    extractor = "msg-attachment:txt"
                elif ext == ".doc" and enable_soffice:
                    soffice_bin = soffice_available()
                    if soffice_bin:
                        result = extract_via_soffice(att_path, soffice_bin)
                        extractor = "msg-attachment:doc"
                
                if result and result.success:
                    data = result.value
                    doc = SourceDoc(
                        path=str(att_path),
                        ext=ext,
                        size_bytes=att_path.stat().st_size,
                        mtime_iso=file_mtime_iso(att_path),
                        extractor=extractor,
                        text=data["text"],
                        text_sha256=sha256_text(data["text"]),
                        pages=data.get("pages"),
                    )
                    documents.append(doc)
                    ok += 1
                elif result:
                    errors += 1
            except Exception as exc:
                LOG.warning("FAIL pièce jointe %s -> %s", att_path.name, exc)
                errors += 1

    payload = {
        "root": str(root),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "enable_soffice": enable_soffice,
        "enable_msg": enable_msg,  # ✅ Ajouter flag .msg
        "msg_attachments_extracted": len(additional_files),  # ✅ Stats pièces jointes
        "counts": {
            "ok": ok,
            "errors": errors,
            "skipped": skipped,
            "total_seen": len(files),
        },
        "documents": [doc.__dict__ for doc in documents],
    }
    return payload


def write_payload(payload: dict, out_path: Path) -> Path:
    out_path = Path(out_path).expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return out_path
