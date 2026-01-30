"""Rendu du rapport DOCX à partir d'un template."""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any, Optional, Union

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

LOGGER = logging.getLogger(__name__)


def _norm(text: str) -> str:
    text = (text or "").replace("\u00a0", " ").strip().lower()
    text = text.replace(":", "")
    text = " ".join(text.split())
    return text


def _style_ok(paragraph: Paragraph, prefixes: Optional[list[str]]) -> bool:
    name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if name.startswith("TOC"):
        return False
    if not prefixes:
        return True
    return any(name.startswith(prefix) for prefix in prefixes)


def find_paragraph(
    doc: Document,
    text: str,
    *,
    after: int = 0,
    style_prefixes: Optional[list[str]] = None,
) -> tuple[Optional[int], Optional[Paragraph]]:
    target = _norm(text)
    for idx in range(after, len(doc.paragraphs)):
        p = doc.paragraphs[idx]
        if not _style_ok(p, style_prefixes):
            continue
        if _norm(p.text) == target:
            return idx, p
    return None, None


def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: Optional[str]) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._element.addnext(new_element)
    para = Paragraph(new_element, paragraph._parent)
    if style_name:
        try:
            para.style = style_name
        except Exception:
            pass
    if text is not None:
        para.add_run(text)
    return para


def replace_text_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    # Replace longer tokens first to avoid partial overlaps (e.g. {NAME} vs {{NAME}})
    ordered_items = sorted(
        ((old, new) for old, new in mapping.items() if old),
        key=lambda item: len(item[0]),
        reverse=True,
    )

    def replace_in_par(par: Paragraph):
        if not ordered_items:
            return
        text = "".join(run.text for run in par.runs) if par.runs else par.text
        replaced = False
        for old, new in ordered_items:
            if old not in text:
                continue
            text = text.replace(old, new)
            replaced = True
        if replaced:
            if par.runs:
                par.runs[0].text = text
                for r in par.runs[1:]:
                    r.text = ""
            else:
                par.text = text

    for paragraph in doc.paragraphs:
        replace_in_par(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_par(paragraph)

    # Headers et footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            if header_footer and header_footer.is_linked_to_previous:
                continue
            try:
                for paragraph in header_footer.paragraphs:
                    replace_in_par(paragraph)
                for table in header_footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_in_par(paragraph)
            except Exception:
                pass


def _stringify_answer(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        answer = value.get("value")
        if isinstance(answer, str):
            return answer.strip()
        answer = value.get("answer")
        if isinstance(answer, str):
            return answer.strip()
        if answer is None:
            return ""
        return json.dumps(answer, ensure_ascii=False)
    if value is None:
        return ""
    return str(value).strip()


def build_moustache_mapping(answers: dict[str, Any]) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for key, value in answers.items():
        text = _stringify_answer(value)
        if not text:
            continue
        placeholder = f"{{{{{key}}}}}"
        mapping[placeholder] = text
        mapping.setdefault(f"{{{{{key.lower()}}}}}", text)
    return mapping


def replace_section(
    doc: Document,
    *,
    start_text: str,
    end_text: str,
    answer_text: str,
    start_style_prefixes: Optional[list[str]] = None,
    end_style_prefixes: Optional[list[str]] = None,
    remove_if_empty: bool = True,
) -> None:
    start_idx, start_par = find_paragraph(doc, start_text, style_prefixes=start_style_prefixes)
    if start_par is None or start_idx is None:
        raise RuntimeError(f"Section '{start_text}' introuvable")
    end_idx, end_par = find_paragraph(doc, end_text, after=start_idx + 1, style_prefixes=end_style_prefixes)
    if end_par is None or end_idx is None:
        LOGGER.warning(
            "Section fin '%s' introuvable après '%s' – insertion en fin de document.", end_text, start_text
        )
        end_idx = len(doc.paragraphs)
    between = doc.paragraphs[start_idx + 1 : end_idx]
    base_style = None
    for paragraph in between:
        if paragraph.text.strip():
            base_style = getattr(getattr(paragraph, "style", None), "name", None)
            break
    if not base_style and between:
        base_style = getattr(getattr(between[0], "style", None), "name", None)
    if not base_style:
        base_style = "Corps"
    for paragraph in list(between):
        delete_paragraph(paragraph)
    text = (answer_text or "").strip()
    # Considérer "Non renseigné" et "[]" comme vide
    if text.lower() in ("non renseigné", "non renseigne", "[]", "non évalué", "non evalue", "vide"):
        text = ""
    # Détection "CHAMP : VIDE" ou "CHAMP : Vide"
    if re.match(r'^[\w\s]+:\s*vide\s*$', text, re.IGNORECASE):
        text = ""
    if not text:
        if remove_if_empty:
            # Supprimer la section entière (titre inclus)
            delete_paragraph(start_par)
        else:
            insert_paragraph_after(start_par, "", base_style)
        return
    cursor = start_par
    for line in [ln.strip() for ln in text.splitlines() if ln.strip()]:
        if line.startswith(("- ", "* ")):
            line = "• " + line[2:].strip()
        cursor = insert_paragraph_after(cursor, line, base_style)


def render_report(
    template: Union[str, Path],
    answers: Union[dict[str, Any], str, Path],
    output: Union[str, Path],
    *,
    name: str = "",
    surname: str = "",
    civility: str = "Monsieur",
    location_date: str = "",
) -> Path:
    template_path = Path(template).expanduser().resolve()
    output_path = Path(output).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(template_path))

    if isinstance(answers, (str, Path)):
        answers_dict: dict[str, Any] = json.loads(Path(answers).expanduser().read_text(encoding="utf-8"))
    else:
        answers_dict = answers

    simple_mapping = {}
    if name:
        simple_mapping["{NAME}"] = name
        simple_mapping["{{NAME}}"] = name
        simple_mapping["{monsieur ou madame NAME}"] = f"{civility} {name}".strip()
    if surname:
        simple_mapping["{surname}"] = surname
        simple_mapping["{{SURNAME}}"] = surname
        simple_mapping["XX"] = f"{civility} {surname}".strip()
    simple_mapping["{{MONSIEUR_OU_MADAME}}"] = civility
    if location_date:
        simple_mapping["{LIEU_ET_DATE}"] = location_date
        simple_mapping["{{LIEU_ET_DATE}}"] = location_date
    if simple_mapping:
        replace_text_everywhere(doc, simple_mapping)

    moustache_mapping = build_moustache_mapping(answers_dict)
    if moustache_mapping:
        replace_text_everywhere(doc, moustache_mapping)

    # Nettoyage : supprimer les placeholders {{...}} orphelins et leurs titres parents
    _clean_orphan_placeholders(doc)
    _remove_empty_sections(doc)

    doc.save(str(output_path))
    return output_path


RE_MOUSTACHE = re.compile(r"\{\{[^}]+\}\}")


def _clean_orphan_placeholders(doc: Document) -> None:
    """Supprime tout texte {{...}} restant dans le document (placeholders non remplacés)."""

    def _clean_par(par: Paragraph) -> None:
        text = "".join(run.text for run in par.runs) if par.runs else (par.text or "")
        if not RE_MOUSTACHE.search(text):
            return
        cleaned = RE_MOUSTACHE.sub("", text).strip()
        if par.runs:
            par.runs[0].text = cleaned
            for r in par.runs[1:]:
                r.text = ""
        else:
            par.text = cleaned
        # Si le paragraphe est maintenant vide, le supprimer
        if not cleaned:
            try:
                delete_paragraph(par)
            except Exception:
                pass

    for paragraph in list(doc.paragraphs):
        _clean_par(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    _clean_par(paragraph)

    # Nettoyer aussi les headers et footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            if header_footer and header_footer.is_linked_to_previous:
                continue
            try:
                for paragraph in list(header_footer.paragraphs):
                    _clean_par(paragraph)
                for table in header_footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in list(cell.paragraphs):
                                _clean_par(paragraph)
            except Exception:
                pass


# Valeurs considérées comme vides
_EMPTY_VALUES = {"non renseigné", "non renseigne", "non évalué", "non evalue", "[]", "vide", ""}
_RE_CHAMP_VIDE = re.compile(r'^[\w\s]+:\s*vide\s*$', re.IGNORECASE)

# Styles de titre (paragraphe qui précède un contenu)
_TITLE_STYLE_PREFIXES = ("Heading", "TITRE")


def _is_title_style(paragraph: Paragraph) -> bool:
    name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    return any(name.startswith(prefix) for prefix in _TITLE_STYLE_PREFIXES)


def _is_empty_content(text: str) -> bool:
    t = (text or "").strip()
    if t.lower() in _EMPTY_VALUES:
        return True
    if _RE_CHAMP_VIDE.match(t):
        return True
    return not t


def _remove_empty_sections(doc: Document) -> None:
    """Supprime les titres dont le contenu suivant est vide.

    Parcourt les paragraphes : si un titre est suivi uniquement de paragraphes
    vides (avant le prochain titre ou la fin), supprime le titre et ces paragraphes.
    Remonte aussi : si un titre parent (Heading 1, TITRE 2) n'a plus aucun
    contenu sous lui, il est également supprimé.
    """
    # Deux passes : d'abord supprimer les sous-sections vides, puis les sections parentes vides
    for _ in range(3):  # max 3 passes pour gérer la cascade
        paragraphs = list(doc.paragraphs)
        if not paragraphs:
            break
        to_delete = []
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            if _is_title_style(p):
                # Collecter les paragraphes de contenu après ce titre
                j = i + 1
                content_paragraphs = []
                while j < len(paragraphs) and not _is_title_style(paragraphs[j]):
                    content_paragraphs.append(paragraphs[j])
                    j += 1
                # Vérifier si tout le contenu est vide
                all_empty = all(_is_empty_content(cp.text) for cp in content_paragraphs)
                if all_empty and content_paragraphs:
                    to_delete.append(p)
                    to_delete.extend(content_paragraphs)
                elif not content_paragraphs:
                    # Titre sans contenu du tout (titre suivi directement d'un autre titre)
                    # Ne supprimer que si ce n'est pas un Heading 1 (section parente)
                    style_name = getattr(getattr(p, "style", None), "name", "") or ""
                    if not style_name.startswith("Heading 1"):
                        to_delete.append(p)
            i += 1

        if not to_delete:
            break
        for p in to_delete:
            try:
                delete_paragraph(p)
            except Exception:
                pass

    # Dernière passe : supprimer les Heading 1 qui n'ont plus de contenu
    paragraphs = list(doc.paragraphs)
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        style_name = getattr(getattr(p, "style", None), "name", "") or ""
        if style_name.startswith("Heading 1"):
            j = i + 1
            has_content = False
            while j < len(paragraphs):
                sn = getattr(getattr(paragraphs[j], "style", None), "name", "") or ""
                if sn.startswith("Heading 1"):
                    break
                if paragraphs[j].text.strip():
                    has_content = True
                    break
                j += 1
            if not has_content:
                try:
                    delete_paragraph(p)
                except Exception:
                    pass
        i += 1
