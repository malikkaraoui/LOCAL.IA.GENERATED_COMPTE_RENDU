"""Utilitaires pour détecter les champs {{...}} d'un template DOCX."""

from __future__ import annotations

import re
from collections.abc import Sequence
from pathlib import Path
from typing import Any, Optional

from docx import Document

from .field_specs import get_field_spec

PLACEHOLDER_RE = re.compile(r"\{\{([^{}]+)\}\}")


def extract_placeholders_from_docx(template_path: Path) -> list[str]:
    """Retourne les placeholders uniques trouvés dans un template DOCX."""

    doc = Document(str(Path(template_path).expanduser().resolve()))
    placeholders: list[str] = []

    def register(text: str) -> None:
        if not text:
            return
        for match in PLACEHOLDER_RE.findall(text):
            key = match.strip()
            if key and key not in placeholders:
                placeholders.append(key)

    for paragraph in doc.paragraphs:
        register(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    register(paragraph.text)

    return placeholders


def build_field_specs(
    placeholders: Sequence[str],
    fallback_defs: Optional[Sequence[dict[str, Any]]] = None,
) -> list[dict[str, Any]]:
    """Construit les spécifications de champs à partir des placeholders."""

    fallback_lookup = {item["key"]: item for item in (fallback_defs or []) if item.get("key")}
    specs: list[dict[str, Any]] = []
    seen: set[str] = set()

    for raw in placeholders:
        key = raw.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        base = fallback_lookup.get(key) if fallback_lookup else None
        spec = get_field_spec(key)
        query = base.get("query") if base and base.get("query") else spec.query
        instructions = base.get("instructions") if base and base.get("instructions") else spec.instructions
        specs.append({
            "key": key,
            "query": query,
            "instructions": instructions,
        })

    return specs
