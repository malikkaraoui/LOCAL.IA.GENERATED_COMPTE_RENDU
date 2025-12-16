import os
import re
import unicodedata

import requests
from docx import Document
from pypdf import PdfReader

# ============================================================
# PARAMÉTRAGE GÉNÉRAL
# ============================================================

# URL de l'API Ollama locale
OLLAMA_URL = "http://localhost:11434/api/chat"

# Nom du modèle à utiliser (doit exister dans `ollama list`)
MODEL_NAME = "qwen3-vl:2b"

# Dossier racine où se trouve CE script (ex: SCRIPT.IA)
SCRIPT_ROOT = os.path.dirname(os.path.abspath(__file__))

# Dossier qui contient tous les clients :
#  SCRIPT_ROOT/
#    CLIENTS/
#      Karaoui/
#      Louisa_Cabard/
#      ...
DOSSIERS_CLIENTS_ROOT = os.path.join(SCRIPT_ROOT, "CLIENTS")

# Nom du client utilisé si on tape simplement "/rapport"
CLIENT_DEFAUT = "Malik Karaoui"

# Template Word commun à tous les clients
TEMPLATE_RAPPORT = os.path.join(SCRIPT_ROOT, "TemplateRapportStage.docx")

# Nom du placeholder dans le template où sera injecté le texte du rapport
PLACEHOLDER_CONTENU = "<<CONTENU_RAPPORT>>"

# Placeholder optionnel pour le nom du stagiaire (si tu veux l'utiliser)
PLACEHOLDER_NOM = "<<NOM_STAGIAIRE>>"


# ============================================================
# OUTILS DIVERS
# ============================================================

def slugify(nom: str) -> str:
    """
    Transforme 'Louisa Cabard' -> 'louisa_cabard'
    Sert à créer des noms de fichiers/dossiers propres.
    """
    # On enlève les accents
    nfkd = unicodedata.normalize("NFKD", nom)
    s = "".join(c for c in nfkd if not unicodedata.combining(c))
    # On remplace tout ce qui n'est pas [a-z0-9] par '_'
    s = re.sub(r"[^a-z0-9]+", "_", s.lower())
    # On supprime les '_' au début/fin
    s = s.strip("_")
    return s or "client"


def trouver_dossier_client(nom_client: str):
    """
    À partir d'un nom saisi (ex: 'Louisa Cabard'),
    essaie de retrouver le bon dossier dans DOSSIERS_CLIENTS_ROOT.

    Exemple :
      - dossier réel : CLIENTS/Louisa_Cabard
      - commande : "fais moi le rapport de stage de louisa cabard"
      -> on compare via slugify et on matche.
    """
    if not os.path.isdir(DOSSIERS_CLIENTS_ROOT):
        raise FileNotFoundError(
            f"Le dossier clients n'existe pas : {DOSSIERS_CLIENTS_ROOT}"
        )

    slug_recherche = slugify(nom_client)
    meilleur_score = -1
    meilleur_path = None
    meilleur_nom_dossier = None

    for entry in os.listdir(DOSSIERS_CLIENTS_ROOT):
        full_path = os.path.join(DOSSIERS_CLIENTS_ROOT, entry)
        if not os.path.isdir(full_path):
            continue

        slug_dossier = slugify(entry)

        # Petit scoring pour trouver le dossier le plus proche
        if slug_dossier == slug_recherche:
            score = 3
        elif slug_recherche in slug_dossier or slug_dossier in slug_recherche:
            score = 2
        else:
            morceaux = slug_recherche.split("_")
            if all(m in slug_dossier for m in morceaux):
                score = 1
            else:
                score = 0

        if score > meilleur_score:
            meilleur_score = score
            meilleur_path = full_path
            meilleur_nom_dossier = entry

    if meilleur_path is None or meilleur_score == 0:
        raise FileNotFoundError(
            f"Aucun dossier client trouvé pour : '{nom_client}'. "
            f"Vérifie le contenu de {DOSSIERS_CLIENTS_ROOT}"
        )

    return meilleur_path, meilleur_nom_dossier


# ============================================================
# FONCTIONS LLM / OLLAMA
# ============================================================

def call_llm(messages):
    """
    Appelle le modèle local via l'API chat d'Ollama.

    Paramètre:
      messages : liste de dictionnaires
        ex: [{"role": "user", "content": "Bonjour"}]

    Retour:
      texte de la réponse du modèle (string)
    """
    payload = {
        "model": MODEL_NAME,
        "messages": messages,
        "stream": False,
    }
    resp = requests.post(OLLAMA_URL, json=payload)
    resp.raise_for_status()
    data = resp.json()
    return data["message"]["content"]


# ============================================================
# LECTURE DES DOCUMENTS CLIENT
# ============================================================

def extraire_texte_fichier(path):
    """
    Lit un fichier texte / PDF / DOCX et renvoie son contenu sous forme de string.
    Les autres formats sont ignorés.
    """
    path_lower = path.lower()

    if path_lower.endswith(".txt"):
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()

    if path_lower.endswith(".pdf"):
        reader = PdfReader(path)
        texte = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            texte += page_text + "\n"
        return texte

    if path_lower.endswith(".docx"):
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    # autres formats ignorés
    return ""


def lire_dossier_client(dossier_sources):
    """
    Parcourt le dossier 'sources' du client et lit tous les fichiers
    TXT / PDF / DOCX pour constituer une liste de (nom_fichier, texte).
    """
    documents = []
    for root, _, files in os.walk(dossier_sources):
        for fname in files:
            # on ignore les scripts Python au cas où
            if fname.endswith(".py"):
                continue

            full_path = os.path.join(root, fname)
            texte = extraire_texte_fichier(full_path)
            if texte.strip():
                documents.append((fname, texte))
    return documents


# ============================================================
# SYNTHÈSE & RAPPORT
# ============================================================

def resumer_documents(documents, nom_stagiaire: str):
    """
    documents : liste de (nom_fichier, texte)
    nom_stagiaire : 'Louisa Cabard', 'Malik Karaoui', etc.

    Retourne un gros texte structuré = contenu du rapport.
    """

    system_msg = {
        "role": "system",
        "content": (
            "Tu es un assistant qui prépare des rapports de stage en français. "
            "Tu écris dans un style professionnel, clair et synthétique."
        ),
    }

    # 1) Résumé par document
    resumes = []
    for name, content in documents:
        user_prompt = (
            f"Voici le contenu du document « {name} ».\n"
            f"Résume les informations importantes pour un rapport de stage "
            f"(missions, contexte, compétences, résultats, etc.).\n\n"
            f"CONTENU DU DOCUMENT :\n{content}"
        )

        summary = call_llm([system_msg, {"role": "user", "content": user_prompt}])
        resumes.append((name, summary))

    # 2) Fusion des résumés en un seul rapport structuré
    resumes_text = ""
    for name, summary in resumes:
        resumes_text += f"--- Résumé de {name} ---\n{summary}\n\n"

    final_prompt = (
        f"À partir des résumés ci-dessous, rédige un rapport de stage complet pour "
        f"{nom_stagiaire}. Structure le rapport avec des titres clairs :\n"
        "1. Introduction\n"
        "2. Présentation de l'entreprise et du contexte\n"
        "3. Missions réalisées\n"
        "4. Compétences développées\n"
        "5. Bilan et perspectives\n\n"
        "Le ton doit être professionnel, mais accessible. "
        "N'écris pas de bullet points, uniquement des paragraphes.\n\n"
        "RÉSUMÉS DE DOCUMENTS :\n"
        f"{resumes_text}"
    )

    rapport = call_llm([system_msg, {"role": "user", "content": final_prompt}])
    return rapport


def generer_word(texte_rapport: str, template_path: str, output_path: str, nom_stagiaire: str):
    """
    Crée le fichier Word final.

    - Si un template est défini et existe :
        - remplace <<NOM_STAGIAIRE>> si présent
        - remplace <<CONTENU_RAPPORT>> par le texte généré
        - si <<CONTENU_RAPPORT>> est absent, ajoute le rapport en fin de doc.
    - Sinon :
        - crée un docx simple avec un titre + le contenu.
    """

    if template_path and os.path.exists(template_path):
        doc = Document(template_path)

        # Remplacement du nom du stagiaire (optionnel)
        if PLACEHOLDER_NOM:
            for para in doc.paragraphs:
                if PLACEHOLDER_NOM in para.text:
                    para.text = para.text.replace(PLACEHOLDER_NOM, nom_stagiaire)

        # Remplacement du contenu du rapport
        replaced = False
        if PLACEHOLDER_CONTENU:
            for para in doc.paragraphs:
                if PLACEHOLDER_CONTENU in para.text:
                    para.text = para.text.replace(PLACEHOLDER_CONTENU, texte_rapport)
                    replaced = True
                    break

        # Si pas de placeholder de contenu, on ajoute tout à la fin
        if not replaced:
            doc.add_page_break()
            doc.add_heading(f"Rapport de stage de {nom_stagiaire}", level=1)
            for line in texte_rapport.split("\n"):
                doc.add_paragraph(line)
    else:
        # Pas de template => document minimaliste
        doc = Document()
        doc.add_heading(f"Rapport de stage de {nom_stagiaire}", level=1)
        for line in texte_rapport.split("\n"):
            doc.add_paragraph(line)

    doc.save(output_path)


def generer_rapport_client(nom_client: str):
    """
    Pipeline complet pour un client :
      - trouver le bon dossier
      - localiser 'sources' + 'Rapport'
      - lire les documents
      - appeler l'IA
      - générer le .docx final
    """
    print(f"➡️ Recherche du dossier pour : {nom_client}...")

    # 1) Localiser le dossier du client
    dossier_client, nom_dossier = trouver_dossier_client(nom_client)
    nom_affichage = nom_client  # ce qu'on affiche dans le rapport

    # 2) Dossier 'sources' (où sont les docs à analyser)
    dossier_sources = os.path.join(dossier_client, "sources")
    if not os.path.isdir(dossier_sources):
        # fallback : on prend directement le dossier client si pas de 'sources'
        dossier_sources = dossier_client

    # 3) Dossier 'Rapport' (sortie) => créé si besoin
    dossier_rapport = os.path.join(dossier_client, "Rapport")
    os.makedirs(dossier_rapport, exist_ok=True)

    # 4) Chemin de sortie du .docx
    slug_client = slugify(nom_affichage)
    output_path = os.path.join(
        dossier_rapport,
        f"rapport_stage_{slug_client}.docx",
    )

    print(f"➡️ Lecture du dossier client : {dossier_sources}")
    documents = lire_dossier_client(dossier_sources)
    if not documents:
        print("⚠️ Aucun document lisible trouvé pour ce client.")
        return

    print(f"✅ {len(documents)} document(s) trouvé(s). Analyse par l'IA...")
    texte_rapport = resumer_documents(documents, nom_affichage)

    print("📝 Génération du fichier Word...")
    generer_word(texte_rapport, TEMPLATE_RAPPORT, output_path, nom_affichage)

    print(f"✅ Rapport généré : {output_path}")


# ============================================================
# COMPRÉHENSION D'UNE PHRASE NATURELLE
# ============================================================

def extraire_nom_client_depuis_phrase(user_input: str):
    """
    Essaie de détecter une demande de rapport de stage formulée en langage naturel
    et d'en extraire le nom du client.

    Exemples détectés :
      - "fais moi le rapport de stage de Louisa Cabard"
      - "peux-tu préparer le rapport de stage de Monsieur Karaoui ?"

    Retourne :
      - le nom du client (string) si on a détecté une commande de rapport
      - None sinon (=> c'est un message normal pour le chat)
    """

    lower = user_input.lower()

    # On considère que c'est une demande de rapport s'il y a "rapport" ET "stage"
    if "rapport" not in lower or "stage" not in lower:
        return None

    # On cherche le DERNIER "de " dans la phrase, en supposant que le nom est après
    idx = lower.rfind("de ")
    if idx != -1:
        # On récupère la partie après "de "
        name = user_input[idx + 3 :]
        # On enlève espaces et ponctuation en trop
        name = name.strip(" .!?\"'")
        if name:
            return name

    # Si on n'a pas réussi à extraire un nom, on revient au client par défaut
    return CLIENT_DEFAUT


# ============================================================
# BOUCLE DE CHAT INTERACTIVE
# ============================================================

def boucle_chat():
    """
    Petite boucle de chat en console pour :
      - discuter avec le modèle
      - déclencher un rapport de façon naturelle
        ("fais moi le rapport de stage de Louisa Cabard")
      - OU avec une commande explicite : /rapport [Nom Prenom]
    """
    history = [
        {
            "role": "system",
            "content": (
                "Tu es un assistant utile qui discute en français "
                "et répond de manière concise."
            ),
        }
    ]

    print(f"Chat local avec le modèle {MODEL_NAME}")
    print("Tape 'exit' pour quitter.")
    print("Tu peux dire par exemple :")
    print("   - fais moi le rapport de stage de Malik Karaoui")
    print("   - fais moi le rapport de stage de Louisa Cabard")
    print("Ou utiliser :")
    print("   - /rapport                    (client par défaut)")
    print("   - /rapport Louisa Cabard\n")

    while True:
        try:
            user_input = input("Toi : ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nAu revoir !")
            break

        if not user_input:
            continue

        if user_input.lower() in ("exit", "quit"):
            print("Au revoir !")
            break

        # 1) Commande explicite : /rapport [...]
        if user_input.lower().startswith("/rapport"):
            morceaux = user_input.split(maxsplit=1)
            if len(morceaux) == 1:
                nom_client = CLIENT_DEFAUT
            else:
                nom_client = morceaux[1].strip().strip('"').strip("'")

            print(f"🤖 OK, je prépare le rapport de stage de {nom_client}...")
            try:
                generer_rapport_client(nom_client)
            except FileNotFoundError as e:
                print(f"❌ Erreur : {e}")
            except Exception as e:
                print(f"❌ Erreur inattendue pendant la génération : {e}")
            continue

        # 2) Phrase naturelle du type "fais moi le rapport de stage de ..."
        nom_naturel = extraire_nom_client_depuis_phrase(user_input)
        if nom_naturel is not None:
            print(f"🤖 J'ai compris que tu veux le rapport de stage de {nom_naturel}...")
            try:
                generer_rapport_client(nom_naturel)
            except FileNotFoundError as e:
                print(f"❌ Erreur : {e}")
            except Exception as e:
                print(f"❌ Erreur inattendue pendant la génération : {e}")
            continue

        # 3) Sinon -> message normal pour le chat
        history.append({"role": "user", "content": user_input})
        try:
            response = call_llm(history)
        except Exception as e:
            print("Erreur en appelant le modèle :", e)
            history.pop()
            continue

        print("Assistant :", response)
        history.append({"role": "assistant", "content": response})


# ============================================================
# POINT D'ENTRÉE
# ============================================================

if __name__ == "__main__":
    boucle_chat()
