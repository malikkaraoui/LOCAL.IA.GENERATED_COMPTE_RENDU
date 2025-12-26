#!/usr/bin/env python3
"""
Crée un document DOCX de test pour le parser RH-Pro
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_sample_bilan():
    """Crée un bilan RH-Pro sample"""
    doc = Document()
    
    # En-tête
    header = doc.add_heading('BILAN D\'ORIENTATION PROFESSIONNELLE', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Identité
    doc.add_heading('Identité', level=1)
    p = doc.add_paragraph('M. Jean DUPONT')
    p.add_run('\nAVS: 756.1234.5678.90')
    
    # Participation au programme
    doc.add_heading('Participation au programme', level=1)
    doc.add_paragraph(
        'Le bénéficiaire a participé activement au programme d\'orientation '
        'professionnelle pendant 6 mois. Il a montré une grande motivation '
        'et un engagement constant dans les différentes étapes du processus.'
    )
    
    # Profession & Formation
    doc.add_heading('Profession & Formation', level=1)
    
    doc.add_heading('Profession', level=2)
    doc.add_paragraph(
        'Le bénéficiaire a travaillé pendant 15 ans en tant que technicien '
        'en informatique dans le secteur bancaire. Il possède une solide '
        'expérience en support technique et en maintenance de systèmes.'
    )
    
    doc.add_heading('Formation', level=2)
    doc.add_paragraph(
        'CFC d\'informaticien obtenu en 2005. Formation continue en '
        'cybersécurité suivie en 2018. Certificat CISCO obtenu en 2019.'
    )
    
    # Tests
    doc.add_heading('Tests', level=1)
    
    doc.add_heading('Evolution', level=2)
    doc.add_paragraph(
        'Le bénéficiaire montre une évolution positive dans sa capacité '
        'd\'adaptation et sa confiance en soi. Il a progressé dans la '
        'définition de son projet professionnel.'
    )
    
    doc.add_heading('Ressources professionnelles', level=2)
    
    doc.add_heading('Ressources comportementales', level=3)
    doc.add_paragraph('Points d\'appui:')
    doc.add_paragraph('- Grande rigueur et précision dans le travail', style='List Bullet')
    doc.add_paragraph('- Capacité d\'analyse et de résolution de problèmes', style='List Bullet')
    doc.add_paragraph('- Bon esprit d\'équipe', style='List Bullet')
    
    doc.add_paragraph('\nPoints de vigilance:')
    doc.add_paragraph('- Tendance au perfectionnisme pouvant ralentir l\'exécution', style='List Bullet')
    doc.add_paragraph('- Difficulté à déléguer certaines tâches', style='List Bullet')
    
    doc.add_heading('Ressources motivationnelles', level=3)
    doc.add_paragraph(
        'Fort intérêt pour les technologies émergentes et la sécurité '
        'informatique. Souhaite évoluer vers des responsabilités de '
        'chef de projet ou d\'expert technique.'
    )
    
    # Compétences
    doc.add_heading('Compétences Professionnelles & Sociales', level=1)
    
    doc.add_heading('Sociales', level=2)
    doc.add_paragraph(
        'Bonnes capacités de communication, à l\'aise en équipe. '
        'Capacité d\'écoute et de médiation dans les situations conflictuelles.'
    )
    
    doc.add_heading('Professionnelles', level=2)
    doc.add_paragraph(
        'Expertise technique en systèmes Windows et Linux. Maîtrise des '
        'protocoles réseaux et de la sécurité informatique. Connaissance '
        'des bases de données SQL.'
    )
    
    # Orientation & Formation
    doc.add_heading('Orientation & Formation', level=1)
    
    doc.add_heading('Orientation', level=2)
    doc.add_paragraph(
        'Orientation vers le domaine de la cybersécurité avec spécialisation '
        'en analyse de risques. Formation complémentaire recommandée en '
        'gestion de projet agile.'
    )
    
    doc.add_heading('Stage', level=2)
    doc.add_paragraph(
        'Stage de 3 mois recommandé dans une entreprise spécialisée en '
        'sécurité informatique pour consolider les compétences pratiques.'
    )
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'Le bilan est positif. Le bénéficiaire dispose des compétences et '
        'de la motivation nécessaires pour réussir sa réorientation '
        'professionnelle vers la cybersécurité. Un accompagnement sur 6 mois '
        'est recommandé pour faciliter la transition.'
    )
    
    # Sauvegarder
    output_dir = Path(__file__).parent.parent / 'data' / 'samples'
    output_dir.mkdir(parents=True, exist_ok=True)
    
    output_path = output_dir / 'bilan_rhpro_sample.docx'
    doc.save(str(output_path))
    
    print(f"✓ Document sample créé: {output_path}")
    return output_path


if __name__ == '__main__':
    create_sample_bilan()
