#!/usr/bin/env python3
"""
Crée un DOCX synthétique pour les tests avec toutes les sections Step 6
"""
from docx import Document
from pathlib import Path

def create_synthetic_sample():
    """Crée un sample DOCX avec les sections nécessaires pour Step 6"""
    doc = Document()
    
    # Titre du document (devrait être ignoré)
    doc.add_heading("BILAN D'ORIENTATION PROFESSIONNELLE", level=0)
    
    # 1. Identité
    doc.add_heading("Identité", level=1)
    doc.add_paragraph("Nom: Dupont Jean")
    doc.add_paragraph("Date de naissance: 15/03/1985")
    doc.add_paragraph("Adresse: Rue de la Paix 12, 1000 Lausanne")
    
    # 2. Participation au programme
    doc.add_heading("Participation au programme", level=1)
    doc.add_paragraph("Le bénéficiaire a participé activement au programme durant 6 mois.")
    
    # 3. Profession & Formation (avec sous-sections inline)
    doc.add_heading("Profession & Formation", level=1)
    doc.add_paragraph("Profession")
    doc.add_paragraph("Le bénéficiaire a travaillé pendant 15 ans en tant que technicien en informatique dans le secteur bancaire. Il possède une solide expérience en support technique et en maintenance de systèmes.")
    doc.add_paragraph("Formation")
    doc.add_paragraph("CFC d'informaticien obtenu en 2005. Formation continue en cybersécurité suivie en 2018. Certificat CISCO obtenu en 2019.")
    
    # 4. Tests
    doc.add_heading("Tests", level=1)
    doc.add_paragraph("Evolution")
    doc.add_paragraph("Le bénéficiaire montre une évolution positive dans sa capacité d'adaptation et sa confiance en soi. Il a progressé dans la définition de son projet professionnel.")
    doc.add_paragraph("Ressources professionnelles")
    doc.add_paragraph("Ressources comportementales")
    doc.add_paragraph("Points d'appui:")
    doc.add_paragraph("- Grande rigueur et précision dans le travail")
    doc.add_paragraph("- Capacité d'analyse et de résolution de problèmes")
    doc.add_paragraph("- Bon esprit d'équipe")
    doc.add_paragraph("Points de vigilance:")
    doc.add_paragraph("- Tendance au perfectionnisme pouvant ralentir l'exécution")
    doc.add_paragraph("- Difficulté à déléguer certaines tâches")
    doc.add_paragraph("Ressources motivationnelles")
    doc.add_paragraph("Fort intérêt pour les technologies émergentes et la sécurité informatique. Souhaite évoluer vers des responsabilités de chef de projet ou d'expert technique.")
    
    # 5. Compétences (avec sous-sections inline)
    doc.add_heading("Compétences Professionnelles & Sociales", level=1)
    doc.add_paragraph("Compétences Sociales")
    doc.add_paragraph("Bonnes capacités de communication, à l'aise en équipe. Capacité d'écoute et de médiation dans les situations conflictuelles.")
    doc.add_paragraph("Compétences Professionnelles")
    doc.add_paragraph("Expertise technique en systèmes Windows et Linux. Maîtrise des protocoles réseaux et de la sécurité informatique. Connaissance des bases de données SQL.")
    
    # 6. Orientation & Formation (avec sous-sections inline)
    doc.add_heading("Orientation & Formation", level=1)
    doc.add_paragraph("Orientation")
    doc.add_paragraph("Orientation vers le domaine de la cybersécurité avec spécialisation en analyse de risques. Formation complémentaire recommandée en gestion de projet agile.")
    doc.add_paragraph("Stage")
    doc.add_paragraph("Stage de 3 mois recommandé dans une entreprise spécialisée en sécurité informatique pour consolider les compétences pratiques.")
    
    # 7. Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Le bénéficiaire dispose d'un profil solide avec une expérience significative dans le domaine informatique. Son projet professionnel est cohérent et réalisable.")
    
    return doc

if __name__ == '__main__':
    # Créer le document
    doc = create_synthetic_sample()
    
    # Sauvegarder dans client_01
    output_path = Path(__file__).parent.parent / 'data' / 'samples' / 'client_01' / 'source.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    print(f"✅ Sample synthétique créé: {output_path}")
