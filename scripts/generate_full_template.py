"""
Génère le template DOCX complet avec les 53 sections de field_specs_v2.

Usage:
    python scripts/generate_full_template.py

Produit: templates/TEMPLATE_FULL_V2.docx
"""

import sys
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def add_paragraph_with_style(doc: Document, text: str, style_name: str) -> Paragraph:
    """Add a paragraph with a specific style, falling back gracefully."""
    p = doc.add_paragraph(text)
    try:
        p.style = style_name
    except Exception:
        pass
    return p


def build_full_template():
    """Build the full V2 template with all 53 sections."""
    # Start from the existing template to inherit styles
    source = Path(__file__).resolve().parent.parent / "uploaded_templates" / "TEMPLATE_SIMPLE_BASE1.docx"
    doc = Document(str(source))

    # Clear all existing paragraphs (keep styles defined in the doc)
    for p in list(doc.paragraphs):
        el = p._element
        el.getparent().remove(el)

    # Also clear tables if any
    for t in list(doc.tables):
        el = t._element
        el.getparent().remove(el)

    # ===== HEADER =====
    add_paragraph_with_style(doc, "{{MONSIEUR_OU_MADAME}} {{NAME}}  {{SURNAME}}", "Titre B")
    add_paragraph_with_style(doc, "AVS {{NUMERO_AVS}}", "Titre B")
    add_paragraph_with_style(doc, "Participation au programme", "Titre B")
    add_paragraph_with_style(doc, "", "Corps B")

    # ===== SECTION 1: PROFESSION & FORMATION =====
    add_paragraph_with_style(doc, "Profession & Formation", "Heading 1")

    add_paragraph_with_style(doc, "Profession", "TITRE 2")
    add_paragraph_with_style(doc, "{{PROFESSION}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Formation", "TITRE 2")
    add_paragraph_with_style(doc, "{{FORMATION}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Présentation", "TITRE 2")
    add_paragraph_with_style(doc, "{{PRESENTATION}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Relation à la carrière", "TITRE 2")
    add_paragraph_with_style(doc, "{{RELATION_A_LA_CARRIERE}}", "Corps SPEC")

    # ===== SECTION 2: TESTS =====
    add_paragraph_with_style(doc, "Tests", "Heading 1")

    # -- Positionnement de niveau: Langues --
    add_paragraph_with_style(doc, "Positionnement de niveau", "TITRE 2")

    add_paragraph_with_style(doc, "Français – positionnement de niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{FRANCAIS_POSITIONNEMENT_DE_NIVEAU}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Anglais – positionnement de niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{ANGLAIS_POSITIONNEMENT_DE_NIVEAU}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Allemand – positionnement de niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{ALLEMAND_POSITIONNEMENT_DE_NIVEAU}}", "Corps SPEC")

    # -- Positionnement de niveau: Bureautique --
    add_paragraph_with_style(doc, "Word – positionnement de niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{WORD_POSITIONNEMENT_DE_NIVEAU}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Excel – positionnement de niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{EXCEL_POSITIONNEMENT_DE_NIVEAU}}", "Corps SPEC")

    add_paragraph_with_style(doc, "PowerPoint – positionnement de niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{POWERPOINT_POSITIONNEMENT_DE_NIVEAU}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Outlook – positionnement de niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{OUTLOOK_POSITIONNEMENT_DE_NIVEAU}}", "Corps SPEC")

    # -- Tests psychométriques --
    add_paragraph_with_style(doc, "Tests psychométriques", "TITRE 2")

    add_paragraph_with_style(doc, "Tri et classement", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{TRI_ET_CLASSEMENT}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Test attention administratif", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{TEST_ATTENTION_ADMINISTRATIF}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Calcul niveau", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{CALCUL_NIVEAU}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Dimensions volumes et mesures", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{DIMENSIONS_VOLUMES_ET_MESURES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Test niveau comptabilité", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{TEST_NIVEAU_COMPTABILITE}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Test compréhension consignes", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{TEST_COMPREHENSION_CONSIGNES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Test saisie commandes", "TITRE 2.2 A")
    add_paragraph_with_style(doc, "{{TEST_SAISIE_COMMANDES}}", "Corps SPEC")

    # ===== SECTION 3: ÉVOLUTION / RESSOURCES =====
    add_paragraph_with_style(doc, "Évolution", "Heading 1")

    add_paragraph_with_style(doc, "Ressources professionnelles", "TITRE 2.2 A")

    add_paragraph_with_style(doc, "Ressources motivationnelles", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{RESSOURCES_MOTIVATIONNELLES_PRINCIPAUX}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Relation au marché de l'emploi", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{RELATION_AU_MARCHE_DE_LEMPLOI}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Stratégies comportementales", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{STRATEGIES_COMPORTEMENTALES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Ressources comportementales – points d'appui", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{RESSOURCES_COMPORTEMENTALES_POINTS_APPUI}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Ressources comportementales – points de vigilance", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{RESSOURCES_COMPORTEMENTALES_POINTS_VIGILANCE}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Ressources interpersonnelles", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{RESSOURCES_INTERPERSONNELLES_PRINCIPALES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Conditions de succès", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{CONDITIONS_DE_SUCCES}}", "Corps SPEC")

    # -- Profil emploi --
    add_paragraph_with_style(doc, "Profil emploi", "TITRE 2.2 A")

    add_paragraph_with_style(doc, "Contexte organisation privilégiée", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{CONTEXTE_ORGANISATION_PRIVILEGIEE}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Contexte rôle privilégié", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{CONTEXTE_ROLE_PRIVILEGIE}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Activités", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{ACTIVITES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Activités privilégiées", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{ACTIVITES_PRIVILEGIEES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Secteurs privilégiés", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{SECTEURS_PRIVILEGIES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Fonctions privilégiées", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{FONCTIONS_PRIVILEGIEES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Métiers privilégiés envisageables", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{METIERS_PRIVILEGIES_ENVISAGEABLES}}", "Corps SPEC")

    # -- Vocatio / RIASEC --
    add_paragraph_with_style(doc, "Vocatio", "TITRE 2")

    add_paragraph_with_style(doc, "Vocatio", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{VOCATIO}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Domaines professionnels exemples", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{DOMAINES_PROFESSIONNELS_EXEMPLES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "RIASEC correspondance score", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{RIASEC_CORRESPONDANCE_SCORE}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Rôles professionnels", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{ROLES_PROFESSIONNELS}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Professions", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{PROFESSIONS}}", "Corps SPEC")

    # -- Formations supérieures --
    add_paragraph_with_style(doc, "Formations supérieures", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{FORMATIONS_SUPERIEURES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Formations hautes écoles", "TITRE 2.3 A A")
    add_paragraph_with_style(doc, "{{FORMATIONS_HAUTES_ECOLES}}", "Corps SPEC")

    # ===== SECTION 4: DISCUSSION ASSURÉ =====
    add_paragraph_with_style(doc, "RÉSULTATS DE LA DISCUSSION AVEC L'ASSURÉ", "TITRE 2")
    add_paragraph_with_style(doc, "{{DISCUSSION_ASSURE}}", "Corps SPEC")

    # ===== SECTION 5: COMPÉTENCES =====
    add_paragraph_with_style(doc, "Compétences Professionnelles & Sociales", "Heading 1")

    add_paragraph_with_style(doc, "Sociales", "TITRE 2")
    add_paragraph_with_style(doc, "{{COMPETENCES_SOCIALES}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Professionnelles", "TITRE 2")
    add_paragraph_with_style(doc, "{{COMPETENCES_PRO}}", "Corps SPEC")

    # ===== SECTION 6: OBSTACLES =====
    add_paragraph_with_style(doc, "Incertitudes & Obstacles", "Heading 1")
    add_paragraph_with_style(doc, "{{OBSTACLES}}", "Corps SPEC")

    # ===== SECTION 7: ORIENTATION =====
    add_paragraph_with_style(doc, "Orientation", "Heading 1")

    add_paragraph_with_style(doc, "Orientation", "TITRE 2")
    add_paragraph_with_style(doc, "{{ORIENTATION}}", "Corps SPEC")

    add_paragraph_with_style(doc, "Stage", "TITRE 2")
    add_paragraph_with_style(doc, "{{STAGE}}", "Corps SPEC")

    # ===== SECTION 8: DOSSIER =====
    add_paragraph_with_style(doc, "Dossier & Présentation", "Heading 1")

    add_paragraph_with_style(doc, "Lettre de motivation", "TITRE 2")
    add_paragraph_with_style(doc, "{{LETTRE_DE_MOTIVATION}}", "Corps SPEC")

    add_paragraph_with_style(doc, "CV", "TITRE 2")
    add_paragraph_with_style(doc, "{{CV}}", "Corps SPEC")

    # ===== SECTION 9: CONCLUSION =====
    add_paragraph_with_style(doc, "Conclusion", "Heading 1")
    add_paragraph_with_style(doc, "{{CONCLUSION}}", "Corps SPEC")

    # ===== FOOTER =====
    add_paragraph_with_style(doc, "", "Corps SPEC")
    add_paragraph_with_style(doc, "Lieu & Date", "TITRE 2")
    add_paragraph_with_style(doc, "{{LIEU_ET_DATE}}", "Corps SPEC")

    # Save
    output = Path(__file__).resolve().parent.parent / "templates" / "TEMPLATE_FULL_V2.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Template saved to {output}")

    # Verify: count placeholders
    doc2 = Document(str(output))
    import re
    placeholders = set()
    for p in doc2.paragraphs:
        for m in re.finditer(r'\{\{(\w+)\}\}', p.text):
            placeholders.add(m.group(1))
    print(f"Placeholders found: {len(placeholders)}")
    for ph in sorted(placeholders):
        print(f"  {{{{{ph}}}}}")


if __name__ == "__main__":
    build_full_template()
