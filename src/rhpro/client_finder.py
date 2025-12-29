"""
Client Finder — Recherche tolérante de dossiers clients par nom
"""
import unicodedata
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from difflib import SequenceMatcher
from src.utils.file_filters import is_ignored_filename
import re
import re


def contains_keyword(text: str, keywords: List[str]) -> bool:
    """
    Vérifie si un texte contient un des mots-clés (case-insensitive).
    
    Args:
        text: Texte à vérifier
        keywords: Liste de mots-clés à rechercher
        
    Returns:
        True si au moins un mot-clé est trouvé
        
    Example:
        >>> contains_keyword("02 Devis", ["devis"])
        True
        >>> contains_keyword("Rapport final", ["devis"])
        False
    """
    text_lower = text.lower()
    return any(keyword.lower() in text_lower for keyword in keywords)


def normalize_text(text: str) -> str:
    """
    Normalise un texte : minuscules + suppression accents
    
    Args:
        text: Texte à normaliser
        
    Returns:
        Texte normalisé (minuscules, sans accents)
        
    Example:
        >>> normalize_text("ARIFI Élodie")
        'arifi elodie'
    """
    # Normalisation NFD (décomposition des accents)
    nfd = unicodedata.normalize('NFD', text)
    # Supprimer les accents (catégorie Mn = Mark, Nonspacing)
    without_accents = ''.join(char for char in nfd if unicodedata.category(char) != 'Mn')
    # Minuscules
    return without_accents.lower()


def fuzzy_score(query: str, target: str) -> float:
    """
    Calcule un score de similarité entre query et target
    
    Args:
        query: Texte recherché
        target: Texte cible
        
    Returns:
        Score entre 0.0 et 1.0 (1.0 = match parfait)
        
    Example:
        >>> fuzzy_score("arifi", "ARIFI Elodie")
        0.85
    """
    # Normaliser les deux textes
    q_norm = normalize_text(query)
    t_norm = normalize_text(target)
    
    # Score basé sur SequenceMatcher
    base_score = SequenceMatcher(None, q_norm, t_norm).ratio()
    
    # Bonus si query est contenu dans target
    if q_norm in t_norm:
        base_score += 0.3
    
    # Bonus si query est au début de target (nom/prénom)
    if t_norm.startswith(q_norm):
        base_score += 0.2
    
    # Bonus si tous les mots de query sont dans target
    query_words = set(q_norm.split())
    target_words = set(t_norm.split())
    if query_words and query_words.issubset(target_words):
        base_score += 0.3
    
    # Limiter à 1.0
    return min(base_score, 1.0)


def find_client_folders(root_dir: str, query: str = None, min_score: float = 0.3) -> List[Dict[str, Any]]:
    """
    Trouve tous les dossiers clients dans root_dir
    Si query fourni, filtre et trie par pertinence (fuzzy matching)
    
    Args:
        root_dir: Dossier racine contenant les dossiers clients
        query: Terme de recherche optionnel (ex: "ARIFI")
        min_score: Score minimum pour inclure un résultat (défaut: 0.3)
        
    Returns:
        Liste de dicts avec:
        - path: Path absolu du dossier
        - name: Nom du dossier
        - score: Score de pertinence (si query fourni)
        - has_docx: True si contient au moins un .docx
        
    Example:
        >>> results = find_client_folders("/path/to/dataset", "ARIFI")
        >>> print(results[0]['name'])
        'ARIFI Elodie'
    """
    root = Path(root_dir)
    
    if not root.exists():
        raise FileNotFoundError(f"Root directory not found: {root_dir}")
    
    if not root.is_dir():
        raise NotADirectoryError(f"Not a directory: {root_dir}")
    
    # Lister tous les sous-dossiers directs
    all_folders = [
        d for d in root.iterdir()
        if d.is_dir() and not d.name.startswith('.')
    ]
    
    results = []
    
    for folder in all_folders:
        # Vérifier si contient au moins un fichier (docx, pdf, txt, etc.)
        has_docx = any(folder.glob("*.docx"))
        has_pdf = any(folder.glob("*.pdf"))
        has_txt = any(folder.glob("*.txt"))
        has_audio = any(folder.glob("*.mp3")) or any(folder.glob("*.wav")) or any(folder.glob("*.m4a"))
        has_files = has_docx or has_pdf or has_txt or has_audio
        
        entry = {
            'path': folder,
            'name': folder.name,
            'has_docx': has_docx,
            'has_pdf': has_pdf,
            'has_txt': has_txt,
            'has_audio': has_audio,
            'has_files': has_files
        }
        
        # Si query fourni, calculer le score
        if query:
            score = fuzzy_score(query, folder.name)
            if score >= min_score:
                entry['score'] = score
                results.append(entry)
        else:
            entry['score'] = 1.0  # Tous les dossiers ont score max si pas de query
            results.append(entry)
    
    # Trier par score décroissant
    results.sort(key=lambda x: x['score'], reverse=True)
    
    return results


def find_client_folder(root_dir: str, query: str, exact: bool = False) -> Tuple[Path, List[Dict[str, Any]]]:
    """
    Trouve UN dossier client par nom (recherche tolérante)
    
    Args:
        root_dir: Dossier racine contenant les dossiers clients
        query: Terme de recherche (ex: "ARIFI", "arifi elodie")
        exact: Si True, nécessite un match exact (après normalisation)
        
    Returns:
        Tuple (best_match_path, all_matches)
        - best_match_path: Path du meilleur résultat (ou None si aucun)
        - all_matches: Liste complète des résultats triés par score
        
    Raises:
        FileNotFoundError: Si root_dir n'existe pas
        
    Example:
        >>> path, matches = find_client_folder("/dataset", "ARIFI")
        >>> if len(matches) > 1:
        ...     print(f"Ambigu: {len(matches)} résultats")
        >>> print(path.name)
        'ARIFI Elodie'
    """
    # Recherche avec score minimum
    min_score = 0.9 if exact else 0.3
    matches = find_client_folders(root_dir, query, min_score)
    
    if not matches:
        return None, []
    
    # Si match exact demandé, vérifier score parfait
    if exact:
        best = matches[0]
        if best['score'] >= 0.95:
            return best['path'], matches
        else:
            return None, matches
    
    # Retourner le meilleur résultat
    return matches[0]['path'], matches


def format_search_results(matches: List[Dict[str, Any]], max_results: int = 10) -> str:
    """
    Formate les résultats de recherche pour affichage console
    
    Args:
        matches: Liste des résultats de find_client_folders()
        max_results: Nombre max de résultats à afficher
        
    Returns:
        str formaté pour console
    """
    if not matches:
        return "⚠️  Aucun résultat trouvé"
    
    lines = []
    lines.append(f"🔍 {len(matches)} résultat(s) trouvé(s):")
    lines.append("")
    
    for i, match in enumerate(matches[:max_results], 1):
        score = match.get('score', 0.0)
        name = match['name']
        
        # Indicateurs de fichiers
        indicators = []
        if match.get('has_docx'):
            indicators.append('📄')
        if match.get('has_pdf'):
            indicators.append('📕')
        if match.get('has_audio'):
            indicators.append('🎤')
        
        indicators_str = ''.join(indicators) if indicators else '📁'
        
        # Affichage
        lines.append(f"{i:2d}. [{score:.2f}] {indicators_str} {name}")
    
    if len(matches) > max_results:
        lines.append(f"    ... et {len(matches) - max_results} autre(s)")
    
    return "\n".join(lines)


def discover_client_documents(client_folder: Path) -> Dict[str, List[Path]]:
    """
    Découvre tous les documents dans un dossier client (racine uniquement, legacy).
    
    ⚠️ DEPRECATED : Utiliser discover_client_documents_recursive() avec max_depth=0 pour équivalent.
    
    Args:
        client_folder: Path du dossier client
        
    Returns:
        Dict avec clés 'docx', 'pdf', 'txt', 'audio'
        Chaque valeur est une liste de Path
        
    Example:
        >>> docs = discover_client_documents(Path("/dataset/ARIFI Elodie"))
        >>> print(f"DOCX: {len(docs['docx'])}")
        DOCX: 3
    """
    if not client_folder.exists():
        raise FileNotFoundError(f"Client folder not found: {client_folder}")
    
    documents = {
        'docx': [f for f in client_folder.glob("*.docx") if not is_ignored_filename(f)],
        'pdf': [f for f in client_folder.glob("*.pdf") if not is_ignored_filename(f)],
        'txt': [f for f in client_folder.glob("*.txt") if not is_ignored_filename(f)],
        'audio': []
    }
    
    # Audio formats
    for ext in ['mp3', 'wav', 'm4a', 'ogg', 'flac']:
        for audio_file in client_folder.glob(f"*.{ext}"):
            if not is_ignored_filename(audio_file):
                documents['audio'].append(audio_file)
    
    return documents


def discover_client_documents_recursive(
    client_folder: Path,
    max_depth: int = 2,
    include_subfolders: bool = True,
    max_files: int = 5000,
    allowed_exts: Optional[set] = None,
    ignore_dirs: Optional[set] = None,
    exclude_dir_keywords: Optional[List[str]] = None,
    exclude_file_keywords: Optional[List[str]] = None,
    follow_symlinks: bool = False
) -> Dict[str, Any]:
    """
    Découvre tous les documents dans un dossier client avec scan récursif contrôlé.
    
    Args:
        client_folder: Path du dossier client
        max_depth: Profondeur maximale de scan (0 = racine uniquement, 1 = sous-dossiers directs, etc.)
        include_subfolders: Si False, force max_depth=0
        max_files: Limite maximale de fichiers scannés (évite freeze UI)
        allowed_exts: Extensions autorisées (défaut: {'.pdf','.docx','.doc','.txt','.msg','.m4a','.mp3','.wav'})
        ignore_dirs: Dossiers à ignorer (défaut: {'.git','__MACOSX','node_modules','.venv','venv','artifacts','output'})
        exclude_dir_keywords: Mots-clés pour exclure des dossiers (défaut: ['devis'])
        exclude_file_keywords: Mots-clés pour exclure des fichiers (défaut: ['devis'])
        follow_symlinks: Suivre les liens symboliques
        
    Returns:
        Dict avec:
        - 'files': Dict[str, List[Path]] avec clés 'docx', 'pdf', 'txt', 'audio', 'msg'
        - 'stats_by_type': Dict[str, int] avec nombre de fichiers par type
        - 'stats_by_subfolder': Dict[str, Dict[str, int]] avec stats par sous-dossier (top 10)
        - 'total_files': int
        - 'truncated': bool (si max_files atteint)
        - 'excluded_dirs': List[str] dossiers exclus
        - 'excluded_files': int nombre de fichiers exclus par keyword
        
    Example:
        >>> result = discover_client_documents_recursive(
        ...     Path("/dataset/ARIFI Elodie"),
        ...     max_depth=2,
        ...     include_subfolders=True,
        ...     exclude_dir_keywords=['devis']
        ... )
        >>> print(f"Total: {result['total_files']} fichiers")
        >>> print(f"DOCX: {len(result['files']['docx'])}")
        >>> print(f"Exclus: {result['excluded_dirs']}")
    """
    import os
    from collections import defaultdict
    
    if not client_folder.exists():
        raise FileNotFoundError(f"Client folder not found: {client_folder}")
    
    # Defaults
    if allowed_exts is None:
        allowed_exts = {'.pdf', '.docx', '.doc', '.txt', '.msg', '.m4a', '.mp3', '.wav', '.ogg', '.flac'}
    
    if ignore_dirs is None:
        ignore_dirs = {'.git', '__MACOSX', 'node_modules', '.venv', 'venv', 'artifacts', 'output', '__pycache__'}
    
    if exclude_dir_keywords is None:
        exclude_dir_keywords = ['devis']
    
    if exclude_file_keywords is None:
        exclude_file_keywords = ['devis']
    
    # Force depth=0 si include_subfolders=False
    if not include_subfolders:
        max_depth = 0
    
    # Structures de résultats
    files = {
        'docx': [],
        'pdf': [],
        'txt': [],
        'audio': [],
        'msg': []
    }
    
    stats_by_type = defaultdict(int)
    stats_by_subfolder = defaultdict(lambda: defaultdict(int))
    total_files = 0
    truncated = False
    excluded_dirs = []
    excluded_files_count = 0
    
    # Scan récursif avec os.walk
    for dirpath, dirnames, filenames in os.walk(client_folder, followlinks=follow_symlinks):
        # Calculer profondeur actuelle
        rel_path = os.path.relpath(dirpath, client_folder)
        if rel_path == '.':
            depth = 0
            subfolder_name = "Racine"
        else:
            depth = rel_path.count(os.sep) + 1
            # Nom du premier sous-dossier (ex: "01 Dossier personnel")
            subfolder_name = rel_path.split(os.sep)[0]
        
        # Arrêter la descente si profondeur max atteinte
        if depth > max_depth:
            dirnames[:] = []  # Empêche de descendre plus profond
            continue
        
        # Filtrer les sous-dossiers à ignorer (ignore_dirs standard)
        dirnames[:] = [d for d in dirnames if d not in ignore_dirs]
        
        # Filtrer les sous-dossiers par keywords (ex: devis)
        original_dirnames = dirnames[:]
        dirnames[:] = [d for d in dirnames if not contains_keyword(d, exclude_dir_keywords)]
        
        # Tracker les dossiers exclus
        for excluded_dir in set(original_dirnames) - set(dirnames):
            excluded_dirs.append(os.path.join(rel_path if rel_path != '.' else '', excluded_dir))
        
        # Scanner les fichiers
        for filename in filenames:
            # Limite max_files
            if total_files >= max_files:
                truncated = True
                break
            
            # Filtrer fichiers temporaires Office et fichiers système
            if is_ignored_filename(filename):
                continue
            
            # Filtrer fichiers par keywords (ex: devis)
            if contains_keyword(filename, exclude_file_keywords):
                excluded_files_count += 1
                continue
            
            # Vérifier extension
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in allowed_exts:
                continue
            
            # Ajouter le fichier
            file_path = Path(dirpath) / filename
            
            # Classer par type
            if file_ext in {'.docx', '.doc'}:
                files['docx'].append(file_path)
                file_type = 'docx'
            elif file_ext == '.pdf':
                files['pdf'].append(file_path)
                file_type = 'pdf'
            elif file_ext == '.txt':
                files['txt'].append(file_path)
                file_type = 'txt'
            elif file_ext == '.msg':
                files['msg'].append(file_path)
                file_type = 'msg'
            elif file_ext in {'.mp3', '.wav', '.m4a', '.ogg', '.flac'}:
                files['audio'].append(file_path)
                file_type = 'audio'
            else:
                continue
            
            # Stats
            stats_by_type[file_type] += 1
            stats_by_subfolder[subfolder_name][file_type] += 1
            total_files += 1
        
        if truncated:
            break
    
    # Limiter stats_by_subfolder aux top 10 pour UI
    top_subfolders = dict(sorted(
        stats_by_subfolder.items(),
        key=lambda x: sum(x[1].values()),
        reverse=True
    )[:10])
    
    return {
        'files': files,
        'stats_by_type': dict(stats_by_type),
        'stats_by_subfolder': top_subfolders,
        'total_files': total_files,
        'truncated': truncated,
        'excluded_dirs': excluded_dirs,
        'excluded_files_count': excluded_files_count
    }


def select_best_source_docx(
    docx_paths: List[Path],
    profile: str = "bilan_complet"
) -> Tuple[Optional[Path], str]:
    """
    Sélectionne automatiquement le meilleur DOCX source pour l'extraction RH-Pro.
    
    Stratégie (Option B - Patch 4):
    - Rejette les docs administratifs/évaluation stage (contrat, devis, facture, attestation, evaluation, stage, certificat)
    - Privilégie les docs RH-Pro structurants (bilan, rapport, orientation, synthese, final, lai, bilan final, bilan d'orientation)
    - Analyse rapide des headings pour détecter la structure RH-Pro
    - Fallback sur le doc le plus long (hors blacklist)
    - Si profile=bilan_complet: interdire docs évaluation/stage/contrat comme source structurante
    
    Args:
        docx_paths: Liste de chemins vers des fichiers DOCX
        profile: Profil gate ('bilan_complet', 'placement_suivi', etc.)
        
    Returns:
        Tuple (best_docx_path, selection_mode):
        - best_docx_path: Path du meilleur DOCX ou None si aucun candidat
        - selection_mode: "AUTO_PRIORITY" | "AUTO_FALLBACK" | "NONE"
        
    Example:
        >>> docx_files = [Path("Contrat.docx"), Path("RH-Pro Bilan final.docx")]
        >>> best, mode = select_best_source_docx(docx_files, profile="bilan_complet")
        >>> print(best.name, mode)
        RH-Pro Bilan final.docx AUTO_PRIORITY
    """
    if not docx_paths:
        return None, "NONE"
    
    # Patch 4 + PATCH 9: Mots-clés à rejeter comme source structurante
    # (mais gardés pour RAG si nécessaire)
    REJECT_KEYWORDS = [
        'contrat', 'convention', 'devis', 'facture', 'attestation',
        'convocation', 'invitation', 'cv', 'curriculum', 'certificat',
        # Patch 4: rejeter aussi evaluation/stage pour bilan_complet
        # PATCH 9: ajouter journal
        'evaluation', 'évaluation', 'stage', 'journal'
    ]
    
    # Patch 2 + PATCH 9: Keywords composés par ordre de priorité
    # Priorité MAX : vrais bilans/rapports finaux
    COMPOSITE_KEYWORDS_HIGH = [
        'bilan final', 'rapport final', 'bilan général', 'bilan general',
        'bilan d\'orientation', 'bilan orientation', 'synthèse finale',
        'synthese finale'
    ]
    
    # Priorité MOYENNE : rapports RH-Pro génériques
    COMPOSITE_KEYWORDS_MEDIUM = [
        'rapport rh-pro', 'rapport rhpro'
    ]
    
    # Patch 2: Mots-clés à privilégier (boost bas)
    BOOST_KEYWORDS = [
        'bilan', 'rapport', 'orientation', 'synthese', 'synthèse',
        'final', 'lai', 'rh-pro', 'rhpro'
    ]
    
    # Anchors RH-Pro connus (pour détection de structure)
    RHPRO_ANCHORS = [
        'identity', 'profession_formation', 'orientation_formation',
        'competences', 'projet', 'preconisations', 'conclusion'
    ]
    
    candidates = []
    
    for docx_path in docx_paths:
        filename = docx_path.name.lower()
        score = 0.0
        
        # Patch 4 + PATCH 9: Rejet strict pour bilan_complet
        # evaluation/stage/contrat/devis/journal ne doivent PAS être source structurante
        if any(keyword in filename for keyword in REJECT_KEYWORDS):
            continue  # Skip complètement (mais restera en RAG)
        
        # PATCH 9: Priorité HAUTE pour vrais rapports finaux
        high_composite_match = False
        for composite in COMPOSITE_KEYWORDS_HIGH:
            if composite in filename:
                score += 50.0  # Boost MAX pour bilan final, rapport final, etc.
                high_composite_match = True
                break
        
        # PATCH 9: Priorité MOYENNE pour rapports RH-Pro génériques
        if not high_composite_match:
            for composite in COMPOSITE_KEYWORDS_MEDIUM:
                if composite in filename:
                    score += 30.0  # Boost moyen
                    break
        
        # Bonus pour mots-clés simples (seulement si pas de composite match HIGH)
        if not high_composite_match:
            for keyword in BOOST_KEYWORDS:
                if keyword in filename:
                    score += 8.0  # Réduit de 10 à 8
        
        # Analyse rapide de la structure (headings) - poids réduit
        try:
            from docx import Document
            doc = Document(str(docx_path))
            
            # Compter les headings (poids réduit si composite HIGH match)
            heading_count = sum(1 for para in doc.paragraphs 
                              if para.style.name.startswith('Heading'))
            structure_bonus = min(heading_count / 10.0, 3.0)  # Max 3 points (réduit de 5)
            if high_composite_match:
                structure_bonus *= 0.5  # Diviser par 2 si déjà HIGH priority
            score += structure_bonus
            
            # Détecter anchors RH-Pro dans le texte (poids réduit)
            full_text = '\n'.join(para.text.lower() for para in doc.paragraphs[:100])  # Premiers 100 paras
            anchor_matches = sum(1 for anchor in RHPRO_ANCHORS if anchor in full_text)
            anchor_bonus = anchor_matches * 2.0  # Réduit de 3.0 à 2.0
            if high_composite_match:
                anchor_bonus *= 0.5  # Diviser par 2 si déjà HIGH priority
            score += anchor_bonus
            
            # Bonus taille (nb paragraphes) - poids réduit
            para_count = len(doc.paragraphs)
            size_bonus = 0.0
            if para_count > 80:
                size_bonus = 3.0  # Réduit de 5
            elif para_count > 50:
                size_bonus = 2.0  # Réduit de 3
            if high_composite_match:
                size_bonus *= 0.5
            score += size_bonus
            
        except Exception:
            # Si erreur de lecture, pas bloquant
            pass
        
        candidates.append((docx_path, score))
    
    if not candidates:
        # Fallback: prendre le plus long DOCX (hors blacklist)
        fallback_candidates = []
        for docx_path in docx_paths:
            filename = docx_path.name.lower()
            if any(keyword in filename for keyword in REJECT_KEYWORDS):
                continue
            try:
                size = docx_path.stat().st_size
                fallback_candidates.append((docx_path, size))
            except Exception:
                continue
        
        if fallback_candidates:
            fallback_candidates.sort(key=lambda x: x[1], reverse=True)
            return fallback_candidates[0][0], "AUTO_FALLBACK"
        else:
            return None, "NONE"
    
    # Trier par score décroissant
    candidates.sort(key=lambda x: x[1], reverse=True)
    best_docx, best_score = candidates[0]
    
    mode = "AUTO_PRIORITY" if best_score > 5.0 else "AUTO_FALLBACK"
    return best_docx, mode
