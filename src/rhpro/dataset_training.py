"""
Module de training dataset pour RH-Pro.

Analyse un dataset de dossiers clients, extrait des patterns de structure/rédaction,
et produit un état persistant (training_state) réutilisable pour la génération RAG+DOCX.
"""
import json
import hashlib
import unicodedata
import re
import shutil
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from collections import Counter, defaultdict
import statistics
from difflib import SequenceMatcher

from .client_scanner import scan_client_folder
from .validation_profiles import validate_report, ValidationProfile
from .gold_diagnostics import diagnose_gold_missing, write_diagnostics_jsonl, write_diagnostics_summary


# ============================================================================
# Sections canoniques RH-Pro V1
# ============================================================================

# META HEADERS — titres administratifs à ignorer (ne comptent pas comme unknown_titles)
# ✅ PRIORITÉ 3: Normalisation identique à celle utilisée pour section_title_map
META_HEADERS_RAW = {
    "PARTICIPATION AU PROGRAMME",
}

# Fonction de normalisation pour les titres (identique à celle du mapper/normalizer)
def _normalize_title_for_meta(title: str) -> str:
    """Normalise un titre de la même manière que section_title_map et unknown_titles"""
    if not title:
        return ""
    # Supprimer ponctuation finale (: … etc), strip, uppercase
    normalized = title.strip().rstrip(':…').strip().upper()
    # Remplacer accents (NFD decomposition)
    normalized = unicodedata.normalize('NFD', normalized)
    normalized = ''.join(c for c in normalized if unicodedata.category(c) != 'Mn')
    return normalized

# FIX ESSAI 100: Ajouter titres méta récurrents à ignorer (AC2)
META_HEADERS_RAW_ADDITIONS = [
    "PARTICIPATION AU PROGRAMME",
    "A L'ATTENTION DE",
    "LIEU ET DATE",
]

# PATCH v1.1 (AC5): Titres administratifs à ignorer complètement
# Ces titres ne créent PAS de section ET ne comptent PAS dans unknown_titles
IGNORED_TITLES_ADMIN = [
    "PARTICIPATION AU PROGRAMME",
    "A L'ATTENTION DE",
    "A L ATTENTION DE",
    "LIEU ET DATE",
    "OFFICE CANTONAL DES ASSURANCES SOCIALES OCAS",
    "OFFICE CANTONAL DES ASSURANCES SOCIALES",
    "OCAS",
    "ASSURANCE INVALIDITE",
    "SERVICE DE L ASSURANCE INVALIDITE",
    "REPUBLIQUE ET CANTON",
    "DEPARTEMENT DE LA SECURITE",
    "EN TETE ADMINISTRATIF",  # Cas générique
]

# Fusionner avec les méta headers existants
ALL_META_HEADERS = list(META_HEADERS_RAW) + META_HEADERS_RAW_ADDITIONS + IGNORED_TITLES_ADMIN

# Pré-calculer les meta headers normalisés
META_HEADERS_NORM = {_normalize_title_for_meta(h) for h in ALL_META_HEADERS}

CANONICAL_SECTIONS = {
    "identity": "Identité",
    "situation_professionnelle": "Situation professionnelle",
    "formation": "Formation",
    "competences": "Compétences",
    "ressources_points_appui": "Ressources comportementales – Points d'appui",
    "ressources_points_vigilance": "Ressources comportementales – Points de vigilance",
    "motivations_valeurs": "Motivations / valeurs",
    "contraintes_freins": "Contraintes / freins",
    "objectifs": "Objectifs",
    "pistes_metiers": "Pistes métiers / pistes d'orientation",
    "plan_action": "Plan d'action",
    "synthese_conclusion": "Synthèse / conclusion"
}

# ============================================================================
# Sections internes (non-canoniques) - Micro-fix v3
# ============================================================================
# Ces sections sont extraites mais ne comptent pas dans les métriques canoniques

INTERNAL_SECTIONS = {
    "tests": "Tests et évaluations"
}

# Toutes les sections reconnues (canoniques + internes)
ALL_SECTIONS = {**CANONICAL_SECTIONS, **INTERNAL_SECTIONS}

# ============================================================================
# Conteneurs / sous-titres (Micro-fix v3)
# ============================================================================
# Ces titres ne doivent PAS ouvrir de nouvelle section ni être comptés en unknown

CONTAINER_HEADINGS = {
    "RESSOURCES COMPORTEMENTALES",
    "SOCIALES",
    "PROFESSIONNELLES",
    "RESSOURCES",
}


# Seed mapping : titre normalisé -> section canonique
SEED_SECTION_TITLE_MAP = {
    # Identity
    "IDENTITE": "identity",
    "ETAT CIVIL": "identity",
    "INFORMATIONS PERSONNELLES": "identity",
    "DONNEES PERSONNELLES": "identity",
    
    # Situation professionnelle
    "SITUATION PROFESSIONNELLE": "situation_professionnelle",
    "PARCOURS PROFESSIONNEL": "situation_professionnelle",
    "EXPERIENCE PROFESSIONNELLE": "situation_professionnelle",
    "EXPERIENCES": "situation_professionnelle",
    "PARCOURS": "situation_professionnelle",
    "STAGE EN LAI 15": "situation_professionnelle",
    "LAI 15": "situation_professionnelle",
    "CONTEXTE ET DEROULEMENT DU STAGE": "situation_professionnelle",
    "PROFESSION": "situation_professionnelle",
    
    # Formation
    "FORMATION": "formation",
    "PARCOURS FORMATION": "formation",
    "DIPLOMES": "formation",
    "CERTIFICATIONS": "formation",
    "CURSUS": "formation",
    
    # Compétences
    "COMPETENCES": "competences",
    "COMPETENCES TECHNIQUES": "competences",
    "SAVOIR FAIRE": "competences",
    "APTITUDES": "competences",
    "DISCIPLINE AU TRAVAIL": "competences",
    "INTEGRATION AUPRES DES COLLABORATEURS": "competences",
    "RYTHME ET QUANTITE DE TRAVAIL": "competences",
    "QUALITE DU TRAVAIL FOURNI": "competences",
    "QUALITE DU TRAVAIL": "competences",
    "ORGANISATION PRISES D INITIATIVES AUTONOMIE": "competences",
    "ORGANISATION": "competences",
    "PRISES D INITIATIVES": "competences",
    "AUTONOMIE": "competences",
    "PRESENTATION": "competences",
    "SELON L EVALUATION DE STAGE FINALE LES TACHES REALISEES ONT ETE LES SUIVANTES": "competences",
    "DANS SON STAGE SES TACHES SONT LES SUIVANTES": "competences",
    
    # Ressources points d'appui
    "RESSOURCES COMPORTEMENTALES POINTS D APPUI": "ressources_points_appui",
    "POINTS D APPUI": "ressources_points_appui",
    "FORCES": "ressources_points_appui",
    "ATOUTS": "ressources_points_appui",
    
    # Ressources points de vigilance
    "RESSOURCES COMPORTEMENTALES POINTS DE VIGILANCE": "ressources_points_vigilance",
    "POINTS DE VIGILANCE": "ressources_points_vigilance",
    "AXES DE VIGILANCE": "ressources_points_vigilance",
    "VIGILANCES": "ressources_points_vigilance",
    
    # Motivations/valeurs
    "MOTIVATIONS": "motivations_valeurs",
    "VALEURS": "motivations_valeurs",
    "INTERETS": "motivations_valeurs",
    "CENTRES D INTERET": "motivations_valeurs",
    "ENGAGEMENT ET PERSEVERANCE": "motivations_valeurs",
    "TEST EVOLUTION": "motivations_valeurs",
    
    # Contraintes/freins
    "CONTRAINTES": "contraintes_freins",
    "FREINS": "contraintes_freins",
    "LIMITES": "contraintes_freins",
    "LIMITATIONS": "contraintes_freins",
    "SANTE": "contraintes_freins",
    "INCERTITUDES & OBSTACLES": "contraintes_freins",
    "INCERTITUDES & OBSTACLES (LIMITATIONS)": "contraintes_freins",
    "ACTUELLEMENT LES LIMITATIONS FONCTIONNELLES RETENUES SONT LES SUIVANTES": "contraintes_freins",
    "LES LIMITATIONS MEDICALES DE L ASSURE SONT LES SUIVANTES": "contraintes_freins",
    "DIFFICULTEES RENCONTREES": "contraintes_freins",
    
    # Objectifs
    "OBJECTIFS": "objectifs",
    "OBJECTIF": "objectifs",
    "PROJET": "objectifs",
    "PROJET PROFESSIONNEL": "objectifs",
    "RELATION AU MARCHE DE L EMPLOI": "objectifs",
    
    # Pistes métiers
    "PISTES": "pistes_metiers",
    "PISTES METIERS": "pistes_metiers",
    "ORIENTATION": "pistes_metiers",
    "PISTES D ORIENTATION": "pistes_metiers",
    "VOCATIO": "pistes_metiers",
    # Micro-fix v3: mapper RESULTATS DE LA DISCUSSION vers pistes_metiers
    "RESULTATS DE LA DISCUSSION AVEC L'ASSURE": "pistes_metiers",
    "RESULTATS DE LA DISCUSSION AVEC L ASSURE": "pistes_metiers",  # variante sans apostrophe
    "RESULTATS DE LA DISCUSSION": "pistes_metiers",
    
    # Plan d'action
    "PLAN D ACTION": "plan_action",
    "ACTIONS": "plan_action",
    "PLANIFICATION": "plan_action",
    "ETAPES": "plan_action",
    
    # Synthèse/conclusion
    "SYNTHESE": "synthese_conclusion",
    "CONCLUSION": "synthese_conclusion",
    "BILAN": "synthese_conclusion",
    "RECAPITULATIF": "synthese_conclusion",
    "QUALITE GENERALE": "synthese_conclusion",
    
    # Tests et évaluations (section interne non-canonique) - Micro-fix v3
    "EVALUATIONS": "tests",
    "EVALUATION": "tests",
    "TESTS METIERS": "tests",
    "TESTS": "tests",  # Note: filtre NOISE gère "TESTS" seul en minuscule
    "FRANCAIS NIVEAU 2": "tests",
    "FRANCAIS NIVEAU 3": "tests",
    "FRANCAIS - NIVEAU 2": "tests",
    "FRANCAIS - NIVEAU 3": "tests",
    "FRANCAIS - NIVEAU 2/3": "tests",
    "POSITIONNEMENT DE NIVEAU DE FRANCAIS": "tests",
    "VITESSE DE FRAPPE EN FRANCAIS": "tests",
    "WORD POSITIONNEMENT DE NIVEAU": "tests",
    "WORD - POSITIONNEMENT DE NIVEAU": "tests",
    "EXCEL POSITIONNEMENT DE NIVEAU": "tests",
    "EXCEL - POSITIONNEMENT DE NIVEAU": "tests",
    "POWERPOINT POSITIONNEMENT DE NIVEAU": "tests",
    "POWERPOINT - POSITIONNEMENT DE NIVEAU": "tests",
    "OUTLOOK 2010": "tests",
    "OUTLOOK": "tests",
    "POSITIONNEMENT": "tests",
    # FIX ESSAI 100: Ajouter top titres inconnus (AC2)
    "FRANCAIS - POSITIONNEMENT DE NIVEAU": "tests",
    "ANGLAIS - POSITIONNEMENT DE NIVEAU": "tests",
    "ALLEMAND - POSITIONNEMENT DE NIVEAU": "tests",
    "CALCUL NIVEAU 1": "tests",
    "CALCUL NIVEAU 2": "tests",
    "CALCUL NIVEAU 3": "tests",
    "CALCUL NIVEAU 2/3": "tests",
    "TRI ET CLASSEMENT": "tests",
    "TEST ADMINISTRATIF BUREAUTIQUE": "tests",
    "DIMENSIONS, VOLUMES ET MESURES": "tests",
    "DIMENSIONS VOLUMES ET MESURES": "tests",
    "SAISIE DE COMMANDES": "tests",
}


# ============================================================================
# Normalisation des titres
# ============================================================================

def is_noise_cell_text(text: str) -> bool:
    """
    Détecte si le contenu d'une cellule/ligne est du bruit (labels formulaire, PII).
    Plus strict que is_noise_title - utilisé pour filtrer contenu, pas seulement titres.
    
    Returns:
        True si le texte est du bruit (à ignorer dans contenu)
    """
    if not text or len(text.strip()) <= 2:
        return True
    
    text_upper = text.strip().upper()
    
    # Normaliser le texte pour comparaison (enlever caractères spéciaux)
    text_normalized = text_upper.replace('°', ' ').replace('°', ' ').strip()
    text_normalized = re.sub(r'\s+', ' ', text_normalized)
    
    # Libellés de formulaire
    form_labels = {
        'NOM', 'PRENOM', 'PRENOM NOM', 'NOM PRENOM',
        'AVS', 'N AVS', 'NUMERO AVS', 'NO AVS', 'N AVS',
        'DATE', 'DATES', 'DATE DE NAISSANCE', 'DATE NAISSANCE',
        'DATES DE LA MESURE', 'PERIODE',
        'CONSEILLER', 'CONSEILLERE', 'RESPONSABLE',
        'TELEPHONE', 'TEL', 'MAIL', 'EMAIL', 'ADRESSE',
        'STAGE', 'STAGE EN QUALITE DE', 'LIEU DE STAGE',
        'EVALUATION', 'EVALUATION DE STAGE',
        'ENTREPRISE', 'L ENTREPRISE', 'EMPLOYEUR',
        'STAGIAIRE', 'LE STAGIAIRE', 'LA STAGIAIRE',
        'SIGNATURE', 'SIGNATURES',
        'PARTIES CONCERNEES', 'CONTACT', 'COORDONNEES'
    }
    if text_normalized in form_labels:
        return True
    
    # AVS suisse pattern
    if re.search(r'\b756[\s\.]?\d{4}[\s\.]?\d{4}[\s\.]?\d{2}\b', text):
        return True
    
    # Dates
    if re.search(r'\b\d{1,2}[\/\.\s]\d{1,2}[\/\.\s]\d{2,4}\b', text):
        return True
    
    # Trop de chiffres (>= 6 digits)
    digit_count = sum(c.isdigit() for c in text)
    if digit_count >= 6:
        return True
    
    # Uniquement chiffres
    if text_upper.replace(' ', '').isdigit():
        return True
    
    # Uniquement ponctuation
    if all(c in '.,;:!?-_/*+= \t\n' for c in text):
        return True
    
    return False


def is_useful_line(text: str) -> bool:
    """
    Détermine si une ligne contient du contenu utile pour le training.
    
    Returns:
        True si ligne utile (pas de bruit, assez substantielle)
    """
    if not text:
        return False
    
    stripped = text.strip()
    
    # Trop court
    if len(stripped) <= 3:
        return False
    
    # Bruit
    if is_noise_cell_text(stripped):
        return False
    
    # OK si au moins quelques mots réels
    words = [w for w in stripped.split() if len(w) >= 2]
    return len(words) >= 2


def is_noise_heading(text: str) -> bool:
    """
    Détecte si un texte détecté comme heading contient des PII ou libellés de formulaire.
    Plus strict que is_noise_title - empêche création de section ET ajout dans unknown_titles.
    
    CORRECTIF B: Filtre "NOM AYNE PRENOM MICKAEL", AVS, dates des unknown_titles.
    ENHANCEMENT: Filtre MONSIEUR/MADAME, phrases intro répétitives, noise patterns.
    
    Returns:
        True si le texte contient des données nominatives ou patterns formulaire
    """
    if not text or len(text.strip()) < 2:
        return True
    
    text_upper = text.strip().upper()
    # Normaliser apostrophes (' → ')
    text_upper = text_upper.replace("'", "'").replace("`", "'")
    text_normalized = re.sub(r'\s+', ' ', text_upper)
    
    # 1. Patterns nominatifs directs : "NOM xxx PRENOM yyy"
    if re.search(r'\bNOM\s+\w+\s+PRENOM\s+\w+', text_normalized):
        return True
    if re.search(r'\bPRENOM\s+\w+\s+NOM\s+\w+', text_normalized):
        return True
    
    # 2. Contient NOM et PRENOM dans le même heading (peu importe l'ordre)
    has_nom = 'NOM' in text_normalized.split()
    has_prenom = 'PRENOM' in text_normalized.split()
    if has_nom and has_prenom:
        return True
    
    # 3. Contient MONSIEUR ou MADAME (PII) - regex flexible pour M. et MME
    # Pattern qui capture M. avec ou sans espace, Monsieur, Madame, etc.
    if re.search(r'\b(MONSIEUR|MADAME|M\s*\.|M\.|MME|MR)\b', text_normalized):
        return True
    # Pattern spécifique pour "M. NOM" avec point collé
    if re.search(r'\bM\.\s*[A-Z]', text_normalized):
        return True
    
    # 4. AVS suisse : 756.xxxx.xxxx.xx
    if re.search(r'\b756[\s\.]?\d{4}[\s\.]?\d{4}[\s\.]?\d{2}\b', text):
        return True
    
    # 5. Dates : dd/mm/yyyy, dd.mm.yyyy
    if re.search(r'\b\d{1,2}[\/\.\s]\d{1,2}[\/\.\s]\d{2,4}\b', text):
        return True
    
    # 6. Trop de chiffres (>= 8 digits) = probablement données perso
    digit_count = sum(c.isdigit() for c in text)
    if digit_count >= 8:
        return True
    
    # 7. Libellés de formulaire seuls
    form_labels = {
        'NOM', 'PRENOM', 'PRENOM NOM', 'NOM PRENOM',
        'AVS', 'N AVS', 'NUMERO AVS', 'NO AVS', 'N AVS',
        'DATE', 'DATES', 'DATE DE NAISSANCE', 'DATE NAISSANCE',
        'NUMERO', 'NO', 'REF', 'REFERENCE'
    }
    if text_normalized in form_labels:
        return True
    
    # 8. Patterns de noise récurrents (sous-intros de tableaux)
    noise_patterns = [
        'LES RESULTATS DETAILLES SONT LES SUIVANTS',
        'CI DESSOUS LES RESULTATS DETAILLES',
        'RESULTATS DE LA DISCUSSION AVEC L\'ASSURE',  # Avec apostrophe normalisée
        'RESULTATS DE LA DISCUSSION AVEC L ASSURE',   # Sans apostrophe
        'TESTS',  # Seul, c'est souvent un conteneur vide
    ]
    if text_normalized in noise_patterns:
        return True
    
    # 9. Phrases très longues avec intro générique (souvent PII ou noise)
    if len(text_normalized) > 60 and any(intro in text_normalized for intro in [
        'LES MOTIVATEURS PRINCIPAUX DE',
        'VOICI',
        'APRES DISCUSSION',
        'SUITE A',
    ]):
        return True
    
    return False
    
    # 9. Phrases très longues avec intro générique (souvent PII ou noise)
    if len(text_normalized) > 60 and any(intro in text_normalized for intro in [
        'LES MOTIVATEURS PRINCIPAUX DE',
        'VOICI',
        'APRES DISCUSSION',
        'SUITE A',
    ]):
        return True
    
    return False


def is_noise_title(text: str) -> bool:
    """
    Détecte les titres NOISE à ignorer (copilot.md section 0 et 5).
    
    Utilise un set lookup pour matching exact des patterns NOISE.
    NE filtre PAS le PII (délégué à is_pii_title).
    
    Returns:
        True si le titre est du NOISE (à ignorer)
    """
    if not text or len(text) < 2:
        return True
    
    # Normaliser pour matching cohérent (micro-fix v2: sans accents)
    text_norm = normalize_heading_for_titles(text)
    
    # ✅ Patterns NOISE exactes (copilot.md v2 - sans accents)
    NOISE_TITLES = {
        "LES RESULTATS DETAILLES SONT LES SUIVANTS",
        "CI DESSOUS LES RESULTATS DETAILLES",
        "RESULTATS DE LA DISCUSSION AVEC L'ASSURE",  # apostrophe normalisée, sans accents
        "TESTS",
    }
    
    if text_norm in NOISE_TITLES:
        return True
    
    # Liste noire explicite (chiffres romains, lettres seules)
    noise_tokens = {'X', 'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII',
                    'TS', 'PS', 'S', 'N', 'P', 'R', 'T', 'A', 'B', 'C', 'D', 'E', 'F', 'G'}
    if text_norm in noise_tokens:
        return True
    
    # ✅ Libellés de champs formulaires (V4)
    form_labels = {
        'NOM', 'PRENOM', 'PRENOM NOM', 'NOM PRENOM',
        'AVS', 'N AVS', 'NUMERO AVS', 'NO AVS',
        'DATE', 'DATES', 'DATE DE NAISSANCE', 'DATE NAISSANCE',
        'DATES DE LA MESURE', 'PERIODE',
        'CONSEILLER', 'CONSEILLERE', 'RESPONSABLE',
        'TELEPHONE', 'TEL', 'MAIL', 'EMAIL', 'ADRESSE',
        'STAGE', 'STAGE EN QUALITE DE', 'LIEU DE STAGE',
        'EVALUATION', 'EVALUATION DE STAGE',
        'ENTREPRISE', "L'ENTREPRISE", "L ENTREPRISE", 'EMPLOYEUR',
        'STAGIAIRE', 'LE STAGIAIRE', 'LA STAGIAIRE',
        'SIGNATURE', 'SIGNATURES',
        'PARTIES CONCERNEES', 'CONTACT', 'COORDONNEES'
    }
    if text_norm in form_labels:
        return True
    
    # Trop court (< 4 caractères)
    if len(text_norm) < 4:
        return True
    
    # Uniquement chiffres
    if text_norm.replace(' ', '').isdigit():
        return True
    
    # Uniquement ponctuation
    if all(c in '.,;:!?-_/*+= \t' for c in text_norm):
        return True
    
    # Un seul token ET trop court (< 3 caractères)
    tokens = text_norm.split()
    if len(tokens) == 1 and len(tokens[0]) < 3:
        return True
    
    return False


def is_pii_title(text: str) -> bool:
    """
    Détecte les titres contenant du PII (copilot.md section 0 et 5).
    
    Filtre :
    - NOM ... PRENOM ... (dans n'importe quel ordre)
    - MONSIEUR/MADAME en début
    - AVS suisse (756.xxxx.xxxx.xx)
    - Dates (dd/mm/yyyy)
    - Trop de chiffres (>= 6 digits)
    
    Returns:
        True si le titre contient du PII (à IGNORER et NE JAMAIS STOCKER)
    """
    if not text or len(text) < 2:
        return False
    
    # Normaliser pour matching cohérent (micro-fix v2: sans accents)
    text_norm = normalize_heading_for_titles(text)
    
    # 1. Patterns NOM + PRENOM (copilot.md v2)
    # Détecte "NOM ... PRENOM ..." ou "PRENOM ... NOM ..."
    # Supporte séparateurs : ":" espaces, "-", "/", etc.
    # Ex: "NOM : X PRENOM : Y", "NOM X PRENOM Y", "NOM- X / PRENOM- Y"
    if re.search(r'\bNOM\b.*\bPRENOM\b|\bPRENOM\b.*\bNOM\b', text_norm):
        return True
    
    # 2. MONSIEUR ou MADAME en début (copilot.md section 0)
    # M. peut avoir un espace après le point : "M. DUBOIS" ou "M.DUBOIS"
    if re.match(r'^\s*(MONSIEUR|MADAME|M\.\s*|MME|MR)\b', text_norm):
        return True
    
    # 3. AVS suisse : 756.xxxx.xxxx.xx
    if re.search(r'\b756[\s\.]?\d{4}[\s\.]?\d{4}[\s\.]?\d{2}\b', text):
        return True
    
    # 4. Dates : dd/mm/yyyy, dd.mm.yyyy, dd mm yyyy
    if re.search(r'\b\d{1,2}[\/\.\s]\d{1,2}[\/\.\s]\d{2,4}\b', text):
        return True
    
    # 5. Trop de chiffres (>= 6 digits) = probablement données perso
    digit_count = sum(c.isdigit() for c in text)
    if digit_count >= 6:
        return True
    
    return False


def normalize_title(title: str) -> str:
    """
    Normalise un titre de section pour matching robuste.
    
    Règles :
    - uppercase
    - strip accents (é → E)
    - trim + collapse espaces
    - remplacer apostrophes typographiques
    - enlever ponctuation faible en fin (;.,)
    - compacter tirets/puces
    - conserver chiffres
    
    FIX ESSAI 100 (AC3): Normalisation durcie pour matcher variantes :
    - Guillemets typographiques (" ") → " 
    - Tirets longs (– —) → -
    - Virgules → espace
    
    Exemple:
        "Ressources comportementales : Points d'appui" 
        → "RESSOURCES COMPORTEMENTALES POINTS D APPUI"
    """
    if not title:
        return ""
    
    # Uppercase
    text = title.upper()
    
    # Strip accents
    text = unicodedata.normalize('NFD', text)
    text = ''.join(c for c in text if unicodedata.category(c) != 'Mn')
    
    # FIX ESSAI 100: Remplacer guillemets typographiques par guillemets droits
    text = text.replace('"', '"').replace('"', '"').replace('«', '"').replace('»', '"')
    
    # Remplacer apostrophes courbes et droites par espace
    text = text.replace(''', ' ').replace(''', ' ').replace("'", ' ')
    
    # FIX ESSAI 100: Remplacer tirets longs (– —) par tiret classique
    text = re.sub(r'[–—]', '-', text)
    
    # Remplacer tirets multiples/puces par espace
    text = re.sub(r'[-•]', ' ', text)
    
    # FIX ESSAI 100: Remplacer virgules et points-virgules par espace
    text = text.replace(',', ' ').replace(';', ' ')
    
    # Enlever ponctuation faible en fin
    text = re.sub(r'[:;.,]+$', '', text)
    text = re.sub(r'[:.,]', ' ', text)
    
    # Collapse espaces
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


def normalize_heading_for_titles(text: str) -> str:
    """
    Normalisation stricte pour filtrage NOISE/PII dans unknown_titles.
    
    Applique (copilot.md v2) :
    - strip + collapse espaces multiples en 1
    - .upper()
    - suppression des accents (É → E, ASSURÉ → ASSURE)
    - apostrophe typographique ' → '
    - retirer ponctuation terminale (., ..., etc.)
    - normaliser tirets multiples en -
    
    Exemple:
        "RÉSULTATS DE LA DISCUSSION AVEC L'ASSURÉ..." 
        → "RESULTATS DE LA DISCUSSION AVEC L'ASSURE"
    """
    if not text:
        return ""
    
    # Strip et uppercase
    text = text.strip().upper()
    
    # Suppression des accents (micro-fix v2)
    text = unicodedata.normalize('NFD', text)
    text = ''.join(c for c in text if unicodedata.category(c) != 'Mn')
    
    # Normaliser apostrophes typographiques : ' ' ` → '
    text = text.replace(''', "'").replace(''', "'").replace('`', "'")
    
    # Normaliser tirets multiples en -
    text = re.sub(r'[-–—]+', '-', text)
    
    # Retirer ponctuation terminale : ., ..., !!!, etc.
    text = re.sub(r'[.!?,;:]+$', '', text)
    
    # Collapse espaces multiples
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


def match_title_to_canonical(title: str, learned_map: Optional[Dict[str, str]] = None) -> Optional[str]:
    """
    Match un titre vers une section canonique.
    
    Stratégie :
    1. Exact match dans seed + learned_map
    2. Heuristiques (keywords)
    3. Fuzzy match (ratio >= 0.85)
    
    Args:
        title: Titre original
        learned_map: Mappings appris (optionnel)
        
    Returns:
        Section canonique ou None
    """
    normalized = normalize_title(title)
    
    if not normalized:
        return None
    
    # 1. Exact match
    full_map = {**SEED_SECTION_TITLE_MAP}
    if learned_map:
        full_map.update(learned_map)
    
    if normalized in full_map:
        return full_map[normalized]
    
    # 2. Heuristiques (keywords)
    keywords_map = {
        "formation": ["FORMATION", "DIPLOME", "CURSUS", "CERTIF"],
        "situation_professionnelle": ["PROFESSIONNEL", "EXPERIENCE", "PARCOURS"],
        "competences": ["COMPETENCE", "APTITUDE", "SAVOIR"],
        "ressources_points_appui": ["APPUI", "FORCE", "ATOUT"],
        "ressources_points_vigilance": ["VIGILANCE", "AXES DE VIGILANCE"],
        "motivations_valeurs": ["MOTIVATION", "VALEUR", "INTERET"],
        "contraintes_freins": ["CONTRAINTE", "FREIN", "LIMITE"],
        "objectifs": ["OBJECTIF", "PROJET"],
        "pistes_metiers": ["PISTE", "ORIENTATION", "METIER"],
        "plan_action": ["PLAN", "ACTION"],
        "synthese_conclusion": ["SYNTHESE", "CONCLUSION", "BILAN"],
        "identity": ["IDENTITE", "ETAT CIVIL", "PERSONNEL"]
    }
    
    for canonical, keywords in keywords_map.items():
        if any(kw in normalized for kw in keywords):
            return canonical
    
    # 3. Fuzzy match
    best_match = None
    best_ratio = 0.85
    
    for seed_title, canonical in full_map.items():
        ratio = SequenceMatcher(None, normalized, seed_title).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_match = canonical
    
    return best_match


def is_container_heading(title: str) -> bool:
    """
    Détermine si un titre est un conteneur/sous-titre (Micro-fix v3).
    
    Les conteneurs ne doivent PAS :
    - ouvrir une nouvelle section
    - être comptés en unknown_titles
    
    Args:
        title: Titre normalisé
        
    Returns:
        True si c'est un conteneur
    """
    normalized = normalize_heading_for_titles(title)
    
    # 1. Match exact dans CONTAINER_HEADINGS
    if normalized in CONTAINER_HEADINGS:
        return True
    
    # 2. Règle heuristique : 1-2 mots courts (sauf si mappé explicitement)
    tokens = normalized.split()
    if len(tokens) <= 2 and len(normalized) <= 20:
        # Vérifier que ce n'est pas un titre mappé explicitement
        if normalized not in SEED_SECTION_TITLE_MAP:
            return True
    
    return False


def is_subheading(title: str) -> bool:
    """
    Détermine si un titre est un sous-titre (Micro-fix v3.1).
    
    Les sous-titres ne doivent PAS :
    - ouvrir une nouvelle section
    - être comptés en unknown_titles
    
    Règles de détection :
    1. Questions : contient '?'
    2. Listes numérotées : commence par \\d+\\.
    3. Phrases longues : > 8 mots (heuristique)
    4. Étiquettes : "MOT : ..." ou "MOT MOT : ..." (préfixe ≤ 2 mots)
    
    Args:
        title: Titre à analyser (peut être normalisé ou non)
        
    Returns:
        True si c'est un sous-titre
    """
    # Garder titre original pour détecter '?' et ':' avant normalisation
    original = title.upper().strip()
    
    # Règle 1 : Questions (détection sur titre original)
    if '?' in original:
        return True
    
    # Normaliser pour les autres règles
    normalized = normalize_heading_for_titles(title)
    tokens = normalized.split()
    
    # Règle 2 : Listes numérotées (commence par 1., 2., etc.)
    if re.match(r'^\d+\.', normalized):
        return True
    
    # Règle 3 : Phrases longues (> 8 mots)
    if len(tokens) > 8:
        return True
    
    # Règle 4 : Étiquettes avec ':' sur titre original (avant que normalisation supprime ':')
    # Format attendu : "MOT : ...", "MOT MOT : ..." mais PAS "MOT MOT MOT : ..."
    if ':' in original:
        # Split sur ':' depuis l'original
        parts_raw = original.split(':', 1)
        if len(parts_raw) == 2:
            prefix_raw = parts_raw[0].strip()
            suffix_raw = parts_raw[1].strip()
            # Compter mots dans préfixe original
            prefix_tokens_raw = prefix_raw.split()
            # Si préfixe court (1-2 mots) ET suffixe non vide → étiquette
            if 1 <= len(prefix_tokens_raw) <= 2 and len(suffix_raw) > 0:
                return True
    
    return False


def apply_max_lines(text: str, max_lines: int) -> str:
    """
    Applique une limite de lignes sur un texte (Micro-fix v3).
    
    Stratégie heuristique (sans invention) :
    - Nettoyer lignes vides / puces répétitives
    - Garder l'ordre original
    - Si > max_lines : garder (max_lines - 1) premières + fusionner reste
    
    Args:
        text: Texte source
        max_lines: Nombre max de lignes
        
    Returns:
        Texte compressé (≤ max_lines)
    """
    if not text or max_lines <= 0:
        return text
    
    # Split en lignes et nettoyer
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Si déjà OK, retourner tel quel
    if len(lines) <= max_lines:
        return '\n'.join(lines)
    
    # Garder (max_lines - 1) premières lignes
    kept_lines = lines[:max_lines - 1]
    
    # Fusionner le reste dans la dernière ligne
    remaining = lines[max_lines - 1:]
    # Limiter la fusion pour éviter une ligne trop longue
    merged = ' ; '.join(remaining)
    if len(merged) > 2000:  # Augmenté de 200 à 2000 pour correspondre au max_chars
        merged = merged[:1997] + '...'
    
    kept_lines.append(merged)
    
    return '\n'.join(kept_lines)


def is_probable_heading(para_text: str, para_obj=None) -> bool:
    """
    Détermine si un paragraphe est probablement un titre de section.
    Avec filtres anti-bruit stricts.
    
    Args:
        para_text: Texte du paragraphe (déjà strippé)
        para_obj: Objet python-docx Paragraph (optionnel, pour style/runs)
        
    Returns:
        True si probable titre, False sinon
    """
    if not para_text or len(para_text) < 2:
        return False
    
    # ❌ FILTRES ANTI-BRUIT (priorité absolue)
    normalized = normalize_title(para_text)
    if is_noise_title(normalized):
        return False
    
    # Trop long pour un titre
    if len(para_text) > 150:
        return False
    
    # ✅ SIGNAUX DE TITRE
    signals = []
    
    # 1. Style
    if para_obj:
        style_name = para_obj.style.name.lower()
        if any(kw in style_name for kw in ['heading', 'titre', 'title']):
            signals.append('style')
    
    # 2. Majuscules (au moins 70% de lettres majuscules)
    if para_text.isupper() and len(para_text.split()) >= 2:
        signals.append('uppercase')
    
    # 3. Gras + court (mais avec au moins 2 mots)
    if para_obj and len(para_text) < 80 and len(para_text.split()) >= 2:
        if para_obj.runs and any(run.bold for run in para_obj.runs):
            signals.append('bold')
    
    # 4. Se termine par ':' et >= 2 mots
    if para_text.endswith(':') and len(para_text.split()) >= 2:
        signals.append('colon')
    
    # 5. Numérotés: "1. Titre", "2.1 Titre"
    if len(para_text) < 100 and re.match(r'^\s*\d+(\.\d+)?[)\.-]\s+\S+', para_text):
        signals.append('numbered')
    
    return len(signals) > 0


def detect_identity_presence(doc) -> bool:
    """
    Détecte si un document contient probablement une section identité,
    même sans titre explicite (ex: dans un tableau).
    
    Args:
        doc: Document python-docx
        
    Returns:
        True si indices d'identité trouvés
    """
    # Mots-clés identité
    keywords = ['NOM', 'PRENOM', 'AVS', 'DATE DE NAISSANCE', 'NAISSANCE', 
                'ADRESSE', 'NATIONALITE', 'ETAT CIVIL', 'SEXE', 'AGE']
    
    # Chercher dans tout le texte (paragraphs + tables)
    all_text = []
    
    # Paragraphes
    for para in doc.paragraphs:
        all_text.append(para.text.upper())
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text.append(para.text.upper())
    
    # Compter combien de keywords différents sont présents
    full_text = ' '.join(all_text)
    matches = sum(1 for kw in keywords if kw in full_text)
    
    # Si au moins 2 keywords différents → identité présente
    return matches >= 2


def score_docx_for_training(docx_path: Path, gold_path: Optional[Path] = None, is_gold_mode: bool = False) -> tuple[int, List[str]]:
    """
    Score un DOCX pour déterminer s'il est adapté à l'extraction de sections (rapport/bilan).
    
    Args:
        docx_path: Path du DOCX à scorer
        gold_path: Path du GOLD (si existe et est DOCX)
        is_gold_mode: Si True, on est en mode GOLD strict (bonus amplifié)
        
    Returns:
        (score, reasons) où score est un int et reasons une liste de justifications
    """
    score = 0
    reasons = []
    
    filename = docx_path.stem.lower()
    
    # ✅ BONUS FORTS
    # +100 si c'est le GOLD
    if gold_path and docx_path == gold_path:
        score += 100
        reasons.append("+100 (GOLD)")
    
    # +40 si nom contient keywords positifs
    positive_keywords = ['rapport', 'bilan', 'synthese', 'conclusion', 'orientation', 'final', 'compte_rendu']
    for kw in positive_keywords:
        if kw in filename:
            score += 40
            reasons.append(f"+40 (keyword: {kw})")
            break
    
    # ❌ MALUS FORTS (formulaires/annexes)
    negative_keywords = ['stage', 'evaluation', 'lai', 'fiche', 'formulaire', 'annexe', 
                         'convention', 'engage', 'entreprise', 'contrat', 'demande']
    for kw in negative_keywords:
        if kw in filename:
            score -= 60
            reasons.append(f"-60 (annexe/formulaire: {kw})")
            break
    
    # ❌ V4.1: MALUS ULTRA-FORTS (journaux/transcriptions/conversations)
    ultra_negative = ['journal', 'chatgpt', 'transcription', 'vous avez dit', 'a dit', 
                     'conversation', 'whatsapp', 'notes', 'entretien', 'discussion']
    for kw in ultra_negative:
        if kw in filename:
            score -= 80
            reasons.append(f"-80 (journal/transcript: {kw})")
            break
    
    # Analyser le contenu
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Compter sections canoniques uniques détectées
        sections_found = set()
        headings_from_tables = 0
        total_headings = 0
        form_headings = 0
        
        # Analyser paragraphes uniquement (V4: pas les tables)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            if is_probable_heading(text, para):
                total_headings += 1
                canonical = match_title_to_canonical(text)
                if canonical:
                    sections_found.add(canonical)
                
                # Détecter headings formulaires
                norm = normalize_title(text)
                if norm in {'NOM', 'PRENOM', 'N AVS', 'AVS', 'DATE'}:
                    form_headings += 1
        
        # Bonus sections canoniques
        nb_sections = len(sections_found)
        if nb_sections > 0:
            bonus_sections = nb_sections * 10
            score += bonus_sections
            reasons.append(f"+{bonus_sections} ({nb_sections} sections canon.)")
        
        # Malus si < 2 sections canoniques
        if nb_sections < 2:
            score -= 20
            reasons.append(f"-20 (< 2 sections canon.)")
        
        # Malus si headings formulaires
        if form_headings > 0:
            malus = min(30, form_headings * 10)
            score -= malus
            reasons.append(f"-{malus} ({form_headings} form headings)")
        
        # Bonus taille fichier (petit, pas dominant)
        filesize_mb = docx_path.stat().st_size / (1024 * 1024)
        bonus_size = min(20, int(filesize_mb * 2))
        score += bonus_size
        reasons.append(f"+{bonus_size} (size: {filesize_mb:.1f}MB)")
        
    except Exception as e:
        score -= 50
        reasons.append(f"-50 (error: {str(e)[:30]})")
    
    return score, reasons


def select_best_docx_for_sections(client_folder: Path, scan_result: Dict) -> tuple[Optional[Path], Dict]:
    """
    Sélectionne le meilleur DOCX pour extraire les sections canoniques.
    
    V4.1: Mode GOLD strict
    - Si GOLD DOCX présent → scorer UNIQUEMENT les GOLD
    - Sinon → scorer tous les DOCX avec malus forts sur journaux/transcripts
    
    Args:
        client_folder: Dossier client
        scan_result: Résultat du scan
        
    Returns:
        (best_docx_path, debug_info) où debug_info contient score/reasons/candidates
    """
    debug_info = {
        "selected_path": None,
        "selected_score": 0,
        "selected_reasons": [],
        "candidates": [],
        "gold_mode": False
    }
    
    # ✅ V4.1: Identifier TOUS les DOCX GOLD (pas seulement le premier)
    gold_docx_list = []
    gold_path = None
    
    if scan_result.get("gold"):
        gold_candidate = Path(scan_result["gold"]["path"])
        if gold_candidate.suffix.lower() == ".docx":
            gold_path = gold_candidate
            gold_docx_list.append(gold_candidate)
    
    # Collecter tous les DOCX candidats
    all_docx = []
    for source in scan_result.get("rag_sources", []):
        path = Path(source["path"])
        if path.suffix.lower() == ".docx":
            all_docx.append(path)
            # Détecter GOLD alternatifs (nom contient "GOLD")
            if "gold" in path.stem.lower() and path not in gold_docx_list:
                gold_docx_list.append(path)
    
    if not all_docx:
        return None, debug_info
    
    # ✅ V4.1: MODE GOLD STRICT
    candidates_to_score = all_docx
    is_gold_mode = False
    
    if gold_docx_list:
        # Si GOLD présent → scorer UNIQUEMENT les GOLD
        candidates_to_score = gold_docx_list
        is_gold_mode = True
        debug_info["gold_mode"] = True
        debug_info["selected_reasons"].append("🔒 GOLD_STRICT_MODE")
    
    # Scorer les candidats
    scored_candidates = []
    for docx_path in candidates_to_score:
        score, reasons = score_docx_for_training(docx_path, gold_path, is_gold_mode)
        scored_candidates.append({
            "path": str(docx_path),
            "name": docx_path.name,
            "score": score,
            "reasons": reasons,
            "is_gold": docx_path in gold_docx_list
        })
    
    # Trier par score décroissant
    scored_candidates.sort(key=lambda x: x["score"], reverse=True)
    debug_info["candidates"] = scored_candidates
    
    # Sélectionner le meilleur
    if scored_candidates:
        best = scored_candidates[0]
        
        # ✅ V4.1: En mode GOLD, accepter même score bas
        threshold = 0 if is_gold_mode else 30
        
        if best["score"] >= threshold:
            debug_info["selected_path"] = best["path"]
            debug_info["selected_score"] = best["score"]
            debug_info["selected_reasons"] = best["reasons"]
            return Path(best["path"]), debug_info
        else:
            # Score trop bas: fallback sur analyse de contenu des 2 meilleurs
            top2 = scored_candidates[:2]
            best_by_sections = None
            max_sections = 0
            
            for cand in top2:
                try:
                    sections = extract_sections_from_docx(Path(cand["path"]))
                    canonical_count = len(set(s["canonical"] for s in sections if s["canonical"]))
                    if canonical_count > max_sections:
                        max_sections = canonical_count
                        best_by_sections = cand
                except:
                    pass
            
            if best_by_sections:
                debug_info["selected_path"] = best_by_sections["path"]
                debug_info["selected_score"] = best_by_sections["score"]
                debug_info["selected_reasons"] = best_by_sections["reasons"] + [f"(fallback: {max_sections} sections)"]
                return Path(best_by_sections["path"]), debug_info
    
    return None, debug_info


def extract_sections_from_docx(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Extrait les sections d'un fichier DOCX (titres + contenu).
    
    V4.1: 
    - HEADINGS détectés UNIQUEMENT depuis paragraphes
    - CONTENU extrait depuis paragraphes + tables (cellules filtrées)
    - Tables ne créent JAMAIS de nouvelles sections
    
    Returns:
        Liste de {title: str, canonical: str|None, lines: int, content_preview: str}
    """
    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception:
        return []
    
    sections = []
    current_section = None
    current_lines = []  # ✅ Liste de lignes utiles
    
    # ✅ Auto-détection IDENTITY (utilise tables pour keywords)
    has_identity = detect_identity_presence(doc)
    if has_identity:
        sections.append({
            "title": "IDENTITE (AUTO)",
            "canonical": "identity",
            "lines": 1,
            "content_preview": ""
        })
    
    # ✅ ÉTAPE 1: Analyser paragraphes pour HEADINGS et CONTENU
    for para_obj in doc.paragraphs:
        text = para_obj.text.strip()
        
        if not text:
            continue
        
        # Détection heading (paragraphes uniquement)
        is_heading = is_probable_heading(text, para_obj)
        
        # ✅ CORRECTIF B: Ignorer si heading contient PII/formulaire
        if is_heading and is_noise_heading(text):
            is_heading = False
        
        if is_heading:
            # Sauvegarder section précédente SI elle a des lignes utiles
            if current_section and len(current_lines) > 0:
                sections.append({
                    "title": current_section["title"],
                    "canonical": current_section["canonical"],
                    "lines": len(current_lines),
                    "content_preview": '\n'.join(current_lines[:5])[:200]
                })
            
            # Nouvelle section
            canonical = match_title_to_canonical(text)
            current_section = {
                "title": text,
                "canonical": canonical
            }
            current_lines = []
        elif current_section:
            # Ajouter contenu si ligne utile
            if is_useful_line(text):
                current_lines.append(text)
    
    # ✅ ÉTAPE 2: Ajouter contenu des tables à la section active
    # Tables ne créent JAMAIS de headings, seulement du contenu
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    cell_text = para.text.strip()
                    
                    if not cell_text:
                        continue
                    
                    # Filtrer bruit (labels formulaire, PII)
                    if is_noise_cell_text(cell_text):
                        continue
                    
                    # Ajouter à section active si ligne utile
                    if current_section and is_useful_line(cell_text):
                        current_lines.append(cell_text)
    
    # Dernière section SI elle a des lignes utiles
    if current_section and len(current_lines) > 0:
        sections.append({
            "title": current_section["title"],
            "canonical": current_section["canonical"],
            "lines": len(current_lines),
            "content_preview": '\n'.join(current_lines[:5])[:200]
        })
    
    return sections


class DatasetTrainingResult:
    """Résultat d'analyse d'un dataset de training."""
    
    def __init__(self, dataset_id: str):
        self.dataset_id = dataset_id
        self.clients: List[Dict[str, Any]] = []
        self.stats: Dict[str, Any] = {}
        self.patterns: Dict[str, Any] = {}
        self.recommendations: List[str] = []
        self.timestamp = datetime.now().isoformat()
        # ✅ PRIORITÉ 5: Diagnostics GOLD missing
        self.gold_missing_diagnostics_path: Optional[str] = None
        self.gold_missing_count: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        base = {
            "dataset_id": self.dataset_id,
            "clients": self.clients,
            "stats": self.stats,
            "patterns": self.patterns,
            "recommendations": self.recommendations,
            "timestamp": self.timestamp,
        }
        # ✅ PRIORITÉ 5: Ajouter diagnostics si présents
        if self.gold_missing_count > 0:
            base["gold_missing_diagnostics"] = {
                "count": self.gold_missing_count,
                "diagnostics_file": self.gold_missing_diagnostics_path,
            }
        return base


def discover_client_folders(
    root_dir: str,
    scan_depth: int = 3
) -> List[Path]:
    """
    Découvre les dossiers clients dans un dataset.
    
    Supporte 2 structures:
    A) "BATCH 20" : root contient des dossiers "NOM Prénom" avec sous-dossiers 01..06
    B) "580 clients non rangés" : root hétérogène, détecte via présence de sources
    
    Args:
        root_dir: Répertoire racine du dataset
        scan_depth: Profondeur de scan (défaut: 3)
        
    Returns:
        Liste de Path vers les dossiers clients détectés (triée par nom)
    """
    root = Path(root_dir).resolve()
    
    if not root.exists():
        raise FileNotFoundError(f"Dataset introuvable : {root}")
    
    if not root.is_dir():
        raise NotADirectoryError(f"Pas un dossier : {root}")
    
    client_folders = []
    exploitable_extensions = {".docx", ".pdf", ".txt", ".doc", ".msg"}
    typical_subfolders = {
        "01", "02", "03", "04", "05", "06",
        "rapport final", "06 rapport final", "tests et bilans",
        "03 tests et bilans", "entretiens", "02 entretiens"
    }
    
    # Structure A: dossiers directs avec pattern "NOM Prénom"
    for item in root.iterdir():
        if not item.is_dir() or item.name.startswith("."):
            continue
        
        # Vérifier si contient des sous-dossiers typiques
        subfolders_lower = {d.name.lower() for d in item.iterdir() if d.is_dir()}
        has_typical_structure = bool(typical_subfolders & subfolders_lower)
        
        # Ou si contient des sources exploitables (au moins 1)
        has_sources = False
        try:
            for f in item.rglob("*"):
                if (f.is_file() and 
                    f.suffix.lower() in exploitable_extensions and
                    not any(p.startswith(".") for p in f.relative_to(item).parts)):
                    has_sources = True
                    break
        except:
            pass
        
        if has_typical_structure or has_sources:
            client_folders.append(item)
    
    # Structure B: scan récursif jusqu'à scan_depth si rien trouvé
    if not client_folders:
        def scan_recursive(path: Path, current_depth: int = 0):
            if current_depth > scan_depth:
                return
            
            for item in path.iterdir():
                if not item.is_dir() or item.name.startswith("."):
                    continue
                
                # Compter les sources exploitables dans ce dossier
                source_count = 0
                try:
                    for f in item.rglob("*"):
                        if (f.is_file() and 
                            f.suffix.lower() in exploitable_extensions and
                            not any(p.startswith(".") for p in f.relative_to(item).parts)):
                            source_count += 1
                            if source_count >= 2:  # Seuil minimum
                                break
                except:
                    pass
                
                if source_count >= 2:
                    client_folders.append(item)
                else:
                    # Continuer la recherche en profondeur
                    scan_recursive(item, current_depth + 1)
        
        scan_recursive(root)
    
    # Trier alphabétiquement par nom pour faciliter la recherche
    client_folders.sort(key=lambda p: p.name.lower())
    
    return client_folders
    
    return sorted(client_folders, key=lambda p: p.name)


def analyze_dataset(
    root_dir: str,
    out_dir: str = "output/training",
    scan_depth: int = 3,
    limit: Optional[int] = None,
    validation_profile: Optional[ValidationProfile] = None,
    index_msg: bool = True,
    quarantine_empty_sources: bool = False,
) -> DatasetTrainingResult:
    """
    Analyse un dataset de clients et extrait patterns/métriques.
    
    Args:
        root_dir: Répertoire racine du dataset
        out_dir: Dossier de sortie pour les artefacts
        scan_depth: Profondeur de scan
        limit: Limiter le nombre de clients à analyser
        validation_profile: Profil de validation optionnel
        index_msg: Si True, inclure les .msg dans le RAG (défaut: True)
        quarantine_empty_sources: Si True, déplacer les dossiers avec sources=0 vers quarantaine (défaut: False)
        
    Returns:
        DatasetTrainingResult avec toutes les analyses
    """
    root = Path(root_dir).resolve()
    dataset_id = _compute_dataset_id(root)
    
    result = DatasetTrainingResult(dataset_id)
    
    # Découvrir les clients
    client_folders = discover_client_folders(str(root), scan_depth)
    
    if limit:
        client_folders = client_folders[:limit]
    
    print(f"📊 Analyse de {len(client_folders)} clients...")
    
    # Collecteurs de patterns
    all_titles = Counter()
    unknown_titles = Counter()
    section_clients = defaultdict(set)  # section -> {client_uid,...}
    section_lines_per_client = defaultdict(list)  # section -> [max_lines_per_client,...]
    gold_strategies = Counter()
    rag_extensions = Counter()
    coverage_scores = []
    quality_scores = []
    no_go_reasons = Counter()
    
    # ✅ PRIORITÉ 5: Collecteur de diagnostics GOLD missing
    gold_missing_diagnostics = []
    
    # Analyser chaque client
    for i, client_folder in enumerate(client_folders, 1):
        try:
            print(f"  [{i}/{len(client_folders)}] {client_folder.name}")
            
            # Scanner le client (✅ Utiliser le paramètre index_msg)
            scan_result = scan_client_folder(str(client_folder), index_msg=index_msg)
            
            # ✅ ROBUSTESSE: Extraire avec .get() pour éviter KeyError
            rag_sources = scan_result.get("rag_sources") or []
            
            # ✅ PRIORITÉ 5: Diagnostic GOLD missing si non détecté
            gold_from_scan = scan_result.get("gold")
            if not gold_from_scan:
                print(f"    🔍 [GOLD MISSING] Diagnostic en cours...")
                diag = diagnose_gold_missing(client_folder, gold_from_scan)
                gold_missing_diagnostics.append(diag)
            
            # Extraire inventaire sources
            sources_by_type = {}
            for source in rag_sources:
                # Vérifier que path existe
                path = source.get("path")
                if not path:
                    continue
                # ✅ Normaliser extension (lowercase, avec point)
                ext = (source.get("extension") or Path(path).suffix or "").lower().strip()
                if ext and not ext.startswith("."):
                    ext = f".{ext}"
                sources_by_type[ext] = sources_by_type.get(ext, 0) + 1
                rag_extensions[ext] += 1
            
            # ✅ ROBUSTESSE: Détection GOLD avec .get()
            gold = scan_result.get("gold") or None
            gold_path = (gold or {}).get("path") or (gold or {}).get("selected_path")
            gold_score = (gold or {}).get("score")
            gold_strategy = (gold or {}).get("strategy")
            
            gold_info = None
            if gold_path:
                gold_info = {
                    "detected": True,
                    "file": Path(gold_path).name,
                    "score": gold_score,
                    "strategy": gold_strategy,
                }
                if gold_strategy:
                    gold_strategies[gold_strategy] += 1
            
            # ✅ V4: Extraire sections depuis le MEILLEUR DOCX avec scoring
            client_sections = []
            docx_selection_debug = {}
            best_docx, docx_debug = select_best_docx_for_sections(client_folder, scan_result)
            docx_selection_debug = {
                "client": client_folder.name,
                "selected_docx": docx_debug.get("selected_path"),
                "score": docx_debug.get("selected_score", 0),
                "reasons": docx_debug.get("selected_reasons", []),
                "candidates": docx_debug.get("candidates", [])
            }
            
            if best_docx and best_docx.exists():
                try:
                    client_sections = extract_sections_from_docx(best_docx)
                    docx_selection_debug["sections_found"] = len(client_sections)
                    docx_selection_debug["canonical_sections"] = len(set(s["canonical"] for s in client_sections if s["canonical"]))
                except Exception as e:
                    docx_selection_debug["error"] = str(e)
            
            # Sauvegarder debug dans les résultats
            if "docx_selections" not in locals():
                docx_selections = []
            docx_selections.append(docx_selection_debug)
            
            # Collecter titres et sections PAR CLIENT
            client_sections_found = set()
            client_section_max_lines = {}  # section -> max_lines dans ce client
            
            for section in client_sections:
                title = section["title"]
                title_norm = normalize_title(title)
                canonical = section["canonical"]
                lines_count = section["lines"]
                
                if canonical:
                    # Titre mappé: compter variant
                    all_titles[title_norm] += 1
                    
                    # ✅ V4.1: Section présente UNIQUEMENT si lines > 0
                    if lines_count > 0:
                        client_sections_found.add(canonical)
                        # Garder le max de lignes pour cette section dans ce client
                        client_section_max_lines[canonical] = max(
                            client_section_max_lines.get(canonical, 0),
                            lines_count
                        )
                else:
                    # ✅ V4.1 + CORRECTIF B + COPILOT.MD + MICRO-FIX V3: Titre non mappé
                    # Normalisation stricte pour filtrage NOISE/PII (copilot.md section 2)
                    title_for_filter = normalize_heading_for_titles(title)
                    
                    # Filtrer PII en premier (zéro tolérance)
                    if is_pii_title(title_for_filter):
                        continue  # NE PAS compter, NE PAS stocker
                    
                    # Filtrer NOISE ensuite
                    if is_noise_title(title_for_filter):
                        continue  # NE PAS compter, NE PAS stocker
                    
                    # Micro-fix v3: Filtrer conteneurs (ne PAS ouvrir section, ne PAS compter unknown)
                    if is_container_heading(title_for_filter):
                        continue  # NE PAS compter, NE PAS stocker
                    
                    # Micro-fix v3.1: Filtrer sous-titres (questions, listes, phrases longues, étiquettes)
                    if is_subheading(title_for_filter):
                        continue  # NE PAS compter, NE PAS stocker
                    
                    # Garder is_noise_heading() pour rétrocompatibilité (détecte autres patterns)
                    if is_noise_heading(title):
                        continue
                    
                    # PATCH v1.1 (AC5): Vérifier que ce n'est pas un titre admin ignoré
                    if _normalize_title_for_meta(title_for_filter) in META_HEADERS_NORM:
                        continue  # NE PAS compter les titres administratifs
                    
                    # Seulement maintenant => unknown
                    unknown_titles[title_for_filter] += 1
            
            # ✅ V4.1: Enregistrer UNIQUEMENT sections avec lines > 0
            client_uid = f"{client_folder.name}_{i}"  # UID interne non exporté
            for sec in client_sections_found:
                section_clients[sec].add(client_uid)
                section_lines_per_client[sec].append(client_section_max_lines[sec])
            
            # Calculer métriques si debug/metrics disponibles
            client_metrics = None
            
            # ✅ ROBUSTESSE: Gérer warnings qui peut être dict ou list
            warnings = scan_result.get("warnings") or []
            if isinstance(warnings, dict):
                warnings = [warnings]
            
            client_info = {
                "folder_name": client_folder.name,
                "folder_path": str(client_folder),
                "sources_count": len(rag_sources),
                "sources_by_type": sources_by_type,
                "gold": gold_info,
                "sections_extracted": len(client_sections),
                "warnings_count": len(warnings),
                "pipeline_ready": scan_result.get("pipeline_ready", False),
                "metrics": client_metrics,
            }
            
            result.clients.append(client_info)
            
        except Exception as e:
            import traceback
            error_msg = str(e)
            error_type = type(e).__name__
            print(f"    ❌ Erreur ({error_type}): {error_msg}")
            # Optionnel: afficher traceback en debug
            # traceback.print_exc()
            
            result.clients.append({
                "folder_name": client_folder.name,
                "folder_path": str(client_folder),
                "error": error_msg,
                "error_type": error_type,
            })
    
    # Calculer statistiques globales
    total_clients = len(result.clients)
    successful_clients = [c for c in result.clients if "error" not in c]
    gold_detected = sum(1 for c in successful_clients if (c.get("gold") or {}).get("detected"))
    pipeline_ready = sum(1 for c in successful_clients if c.get("pipeline_ready"))
    
    # PATCH v1.1 (AC1): Calculer ready par profil (STRICT/STANDARD/DRAFT)
    # Critères simplifiés basés sur les métriques observées:
    # - STRICT: GOLD détecté + sources>=3 + sections>=8
    # - STANDARD: sources>=2 + sections>=5
    # - DRAFT: sources>=1
    ready_strict = sum(1 for c in successful_clients 
                       if c.get("sources_count", 0) >= 3 
                       and (c.get("gold") or {}).get("detected", False)
                       and c.get("sections_extracted", 0) >= 8)
    
    ready_standard = sum(1 for c in successful_clients 
                         if c.get("sources_count", 0) >= 2 
                         and c.get("sections_extracted", 0) >= 5)
    
    ready_draft = sum(1 for c in successful_clients 
                      if c.get("sources_count", 0) >= 1)
    
    # Distributions
    sources_counts = [c["sources_count"] for c in successful_clients if "sources_count" in c]
    
    # ✅ ROBUSTESSE: Calculer errors_top pour affichage
    error_clients = [c for c in result.clients if "error" in c]
    errors_top = Counter([c.get("error_type", "UnknownError") for c in error_clients]).most_common(5)
    
    # ✅ PRIORITÉ 4: Calculer clients_used et clients_no_sources AVANT de construire result.stats
    clients_used_list = [c for c in successful_clients if c.get('sources_count', 0) > 0]
    clients_used = len(clients_used_list)
    
    # Identifier clients avec sources_count=0
    empty_sources_clients = [c for c in successful_clients if c.get('sources_count', 0) == 0]
    clients_no_sources = len(empty_sources_clients)
    
    # Quarantaine des dossiers vides (optionnel)
    quarantine_manifest = None
    quarantine_base = None
    if quarantine_empty_sources and empty_sources_clients:
        run_id = str(uuid.uuid4())[:8]
        quarantine_base = Path("data/_trash/empty_sources") / run_id
        quarantine_base.mkdir(parents=True, exist_ok=True)
        
        manifest_entries = []
        print(f"\n🗑️  Quarantaine de {len(empty_sources_clients)} clients avec sources=0...")
        
        for client_data in empty_sources_clients:
            client_folder_path = Path(client_data["folder_path"])
            client_id = client_data["folder_name"]
            
            try:
                # Déplacer vers quarantaine
                dest_path = quarantine_base / client_folder_path.name
                shutil.move(str(client_folder_path), str(dest_path))
                
                manifest_entries.append({
                    "client_id": client_id,
                    "path_before": str(client_folder_path),
                    "path_after": str(dest_path),
                    "timestamp": datetime.now().isoformat(),
                    "reason": "sources_count=0",
                })
                
                print(f"  ✅ {client_id} → {dest_path.relative_to(Path('data'))}")
            except Exception as e:
                print(f"  ❌ Erreur quarantaine {client_id}: {e}")
                # Ne pas casser le run, continuer
                continue
        
        # Écrire manifest
        manifest_path = quarantine_base / "manifest.json"
        quarantine_manifest = {
            "run_id": run_id,
            "timestamp": datetime.now().isoformat(),
            "total_quarantined": len(manifest_entries),
            "entries": manifest_entries,
        }
        
        with open(manifest_path, "w", encoding="utf-8") as f:
            json.dump(quarantine_manifest, f, indent=2, ensure_ascii=False)
        
        print(f"  📄 Manifest: {manifest_path}")
    
    # ✅ PRIORITÉ 5: Écrire les diagnostics GOLD missing si présents
    if gold_missing_diagnostics:
        print(f"\n📝 Écriture des diagnostics GOLD missing ({len(gold_missing_diagnostics)} clients)...")
        
        # Créer le dossier de sortie
        out_path = Path(out_dir)
        out_path.mkdir(parents=True, exist_ok=True)
        
        # Écrire JSONL (machine-readable)
        jsonl_path = out_path / "gold_missing_debug.jsonl"
        write_diagnostics_jsonl(gold_missing_diagnostics, jsonl_path)
        print(f"  ✅ JSONL: {jsonl_path}")
        
        # Écrire Markdown (human-readable)
        md_path = out_path / "gold_missing_debug.md"
        write_diagnostics_summary(gold_missing_diagnostics, md_path)
        print(f"  ✅ Markdown: {md_path}")
        
        # Ajouter au résultat pour traçabilité
        result.gold_missing_diagnostics_path = str(jsonl_path)
        result.gold_missing_count = len(gold_missing_diagnostics)
    
    result.stats = {
        "total_clients": total_clients,
        "successful_scans": len(successful_clients),
        "errors": total_clients - len(successful_clients),
        "errors_top": errors_top,  # ✅ Top 5 types d'erreurs
        "clients_used": clients_used,  # ✅ PRIORITÉ 4: Clients avec sources > 0
        "clients_no_sources": clients_no_sources,  # ✅ PRIORITÉ 4: Clients sans sources
        "empty_sources_clients_count": clients_no_sources,  # ✅ Feature: Clients vides
        "empty_sources_clients": [c["folder_name"] for c in empty_sources_clients[:50]],  # Top 50
        "quarantine_manifest_path": str(quarantine_base / "manifest.json") if quarantine_manifest else None,
        "gold_detected": gold_detected,
        "gold_detection_rate": gold_detected / len(successful_clients) if successful_clients else 0,
        "pipeline_ready": pipeline_ready,
        "pipeline_ready_rate": pipeline_ready / len(successful_clients) if successful_clients else 0,
        # PATCH v1.1 (AC1): Ready par profil
        "ready_strict": ready_strict,
        "ready_standard": ready_standard,
        "ready_draft": ready_draft,
        "ready_strict_rate": ready_strict / len(successful_clients) if successful_clients else 0,
        "ready_standard_rate": ready_standard / len(successful_clients) if successful_clients else 0,
        "ready_draft_rate": ready_draft / len(successful_clients) if successful_clients else 0,
        "sources_stats": {
            "mean": statistics.mean(sources_counts) if sources_counts else 0,
            "median": statistics.median(sources_counts) if sources_counts else 0,
            "min": min(sources_counts) if sources_counts else 0,
            "max": max(sources_counts) if sources_counts else 0,
            "p10": _percentile(sources_counts, 10) if sources_counts else 0,
            "p90": _percentile(sources_counts, 90) if sources_counts else 0,
        },
        "extensions_distribution": dict(rag_extensions),
        "gold_strategies": dict(gold_strategies),
    }
    
    # Patterns détectés
    # Construction du section_title_map appris depuis unknown_titles
    learned_title_map = {}
    for title_norm, count in unknown_titles.most_common(100):
        # Tenter de mapper avec heuristiques
        canonical = match_title_to_canonical(title_norm)
        if canonical and title_norm not in SEED_SECTION_TITLE_MAP:
            learned_title_map[title_norm] = canonical
    
    # Stats par section canonique (basées sur les CLIENTS, pas les documents)
    sections_stats = {}
    # Note: clients_used et clients_no_sources déjà calculés plus haut
    
    for canonical in CANONICAL_SECTIONS.keys():
        n_clients = len(section_clients.get(canonical, set()))
        lines = section_lines_per_client.get(canonical, [])
        
        if lines:
            avg_lines = statistics.mean(lines)
            median_lines = statistics.median(lines)
            p90_lines = _percentile(lines, 90)
        else:
            avg_lines = median_lines = p90_lines = 0
        
        # Coverage en pourcentage (0..100)
        coverage_pct = 0 if clients_used == 0 else round(100 * n_clients / clients_used, 1)
        
        sections_stats[canonical] = {
            "title_variants_top": [
                title for title, _ in all_titles.most_common(20)
                if match_title_to_canonical(title) == canonical
            ][:5],
            "avg_lines": round(avg_lines, 1),
            "p50_lines": round(median_lines, 1),
            "p90_lines": round(p90_lines, 1),
            "clients": n_clients,  # ✅ Nombre de clients ayant la section
            "coverage": coverage_pct / 100,  # Pour compatibilité (ratio 0..1)
        }
    
    # ✅ GARDE ANTI-PII (ceinture + bretelles) avant sérialisation JSON (copilot.md section 3)
    filtered_unknown = {}
    pii_removed = 0
    noise_removed = 0
    
    for k, v in unknown_titles.items():
        kk = normalize_heading_for_titles(k)
        
        # Filtrer PII
        if is_pii_title(kk):
            pii_removed += 1
            continue  # NE PAS stocker
        
        # Filtrer NOISE (cohérence)
        if is_noise_title(kk):
            noise_removed += 1
            continue  # NE PAS stocker
        
        # OK, garder le titre normalisé comme clé
        filtered_unknown[kk] = v
    
    # Ajouter warning si PII détecté (sans texte PII)
    if pii_removed > 0:
        logger.warning(f"⚠️ {pii_removed} titres PII filtrés de unknown_titles (non stockés)")
    
    if noise_removed > 0:
        logger.info(f"ℹ️ {noise_removed} titres NOISE filtrés de unknown_titles")
    
    # Reconstruction d'un Counter filtré pour most_common()
    filtered_counter = Counter(filtered_unknown)
    
    result.patterns = {
        "unknown_titles_top10": dict(filtered_counter.most_common(10)),
        "unknown_titles_top": dict(filtered_counter.most_common(50)),  # ✅ Top 50 pour review
        "unknown_titles_count": len(filtered_counter),  # ✅ Nombre de titres distincts inconnus
        "unknown_titles_total_occurrences": sum(filtered_counter.values()),  # ✅ Total occurrences
        "learned_title_map": learned_title_map,
        "sections_stats": sections_stats,
        "common_structures": _detect_common_structures(successful_clients),
        # Métadonnées filtrage (copilot.md section 3)
        "pii_titles_filtered": pii_removed,
        "noise_titles_filtered": noise_removed,
    }
    
    # Recommandations
    result.recommendations = []
    
    if gold_detected / len(successful_clients) < 0.5 if successful_clients else True:
        result.recommendations.append(
            "⚠️ Moins de 50% de GOLD détectés : améliorer la détection (stratégies + patterns)"
        )
    
    # FIX ESSAI 100 (AC3): Avertir si beaucoup de clients sans sources
    if clients_no_sources > 0:
        sources_zero_pct = clients_no_sources / len(successful_clients) * 100 if successful_clients else 0
        if sources_zero_pct > 10:
            result.recommendations.append(
                f"⚠️ {clients_no_sources} clients ({sources_zero_pct:.0f}%) ont sources_count=0 "
                f"→ Non utilisables pour training strict/standard (min sources=1)"
            )
    
    if unknown_titles:
        top_unknown = unknown_titles.most_common(5)
        result.recommendations.append(
            f"📝 Titres inconnus fréquents : {', '.join(t for t, _ in top_unknown)} "
            f"→ Ajouter mappings dans field_specs.py"
        )
    
    avg_sources = statistics.mean(sources_counts) if sources_counts else 0
    if avg_sources < 5:
        result.recommendations.append(
            f"📉 Moyenne de sources faible ({avg_sources:.1f}) : vérifier qualité dataset"
        )
    
    return result


def export_training_artifacts(
    result: DatasetTrainingResult,
    out_dir: str = "output/training",
    merge_existing: bool = False
) -> Dict[str, str]:
    """
    Exporte les artefacts de training.
    
    Args:
        result: Résultat d'analyse
        out_dir: Dossier de sortie
        merge_existing: Si True, fusionne avec training_state existant
        
    Returns:
        Dict avec chemins des fichiers générés
    """
    out_path = Path(out_dir) / result.dataset_id
    out_path.mkdir(parents=True, exist_ok=True)
    
    paths = {}
    
    # 1. dataset_manifest.json
    manifest = {
        "dataset_id": result.dataset_id,
        "timestamp": result.timestamp,
        "total_clients": result.stats["total_clients"],
        "successful_scans": result.stats["successful_scans"],
        "client_folders": [c["folder_name"] for c in result.clients if "error" not in c],
    }
    
    manifest_path = out_path / "dataset_manifest.json"
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)
    paths["manifest"] = str(manifest_path)
    
    # 2. dataset_stats.json
    stats_path = out_path / "dataset_stats.json"
    with open(stats_path, "w", encoding="utf-8") as f:
        json.dump(result.stats, f, indent=2, ensure_ascii=False)
    paths["stats"] = str(stats_path)
    
    # 3. training_report.md (lisible)
    report_md = _generate_training_report_md(result)
    report_path = out_path / "training_report.md"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report_md)
    paths["report"] = str(report_path)
    
    # 4. training_state.json (LE FICHIER IMPORTANT)
    training_state = _build_training_state(result)
    
    # Merge si demandé
    state_path = out_path / "training_state.json"
    if merge_existing and state_path.exists():
        # V4.1 : Merge compatible training_state_v1.0
        try:
            existing_state = load_training_state(str(state_path))
            training_state = _merge_training_states(existing_state, training_state)
            print(f"   ♻️ Fusion avec training_state existant")
        except Exception as e:
            print(f"   ⚠️ Échec du merge, écrasement : {e}")
            # Continue sans merge
    
    with open(state_path, "w", encoding="utf-8") as f:
        json.dump(training_state, f, indent=2, ensure_ascii=False)
    paths["training_state"] = str(state_path)
    
    # 5. artifacts/unknown_titles.csv (nouveau - scalabilité ruleset)
    csv_path = _export_unknown_titles_csv(result, out_path.parent.parent / "artifacts")
    if csv_path:
        paths["unknown_titles_csv"] = csv_path
    
    print(f"\n✅ Artefacts exportés vers : {out_path}")
    print(f"   📄 Manifest : {manifest_path.name}")
    print(f"   📊 Stats : {stats_path.name}")
    print(f"   📝 Rapport : {report_path.name}")
    print(f"   🎯 Training state : {state_path.name}")
    if csv_path:
        print(f"   📊 Unknown titles CSV : {Path(csv_path).relative_to(out_path.parent.parent)}")
    
    return paths


def load_training_state(state_path: str) -> Dict[str, Any]:
    """
    Charge un training_state depuis un fichier JSON.
    
    Args:
        state_path: Chemin vers training_state.json
        
    Returns:
        Dictionnaire avec l'état de training
    """
    path = Path(state_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Training state introuvable : {state_path}")
    
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ============================================================================
# Fonctions privées
# ============================================================================

def _compute_dataset_id(root: Path) -> str:
    """Calcule un ID unique pour le dataset."""
    # Hash basé sur le nom + timestamp (pour versioning)
    content = f"{root.name}_{root.stat().st_mtime}"
    return hashlib.md5(content.encode()).hexdigest()[:12]


def _percentile(data: List[float], p: int) -> float:
    """Calcule le percentile p des données."""
    if not data:
        return 0.0
    sorted_data = sorted(data)
    index = int(len(sorted_data) * p / 100)
    return sorted_data[min(index, len(sorted_data) - 1)]


def _detect_common_structures(clients: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Détecte les structures de dossiers communes."""
    folder_patterns = Counter()
    
    for client in clients:
        if "folder_path" in client:
            # Pattern basé sur le nom de dossier
            folder_name = Path(client["folder_path"]).name
            folder_patterns[folder_name] += 1
    
    return {
        "top_patterns": dict(folder_patterns.most_common(5)),
        "total_patterns": len(folder_patterns),
    }


def _export_unknown_titles_csv(
    result: DatasetTrainingResult,
    artifacts_dir: Path
) -> Optional[str]:
    """
    Exporte unknown_titles vers CSV pour scalabilité ruleset.
    
    Format CSV :
    - title_raw : Titre original (non normalisé, pour lisibilité)
    - title_norm : Titre normalisé (clé unique)
    - count : Nombre d'occurrences
    - suggested_action : MAP_TO_SECTION | MAP_TO_TESTS | SUBHEADING_POLICY | IGNORE
    - suggested_target : Section cible suggérée (ou vide)
    - notes : Commentaires/contexte
    
    Args:
        result: Résultat training
        artifacts_dir: Dossier artifacts/
        
    Returns:
        Chemin du CSV ou None si pas d'unknown_titles
    """
    import csv
    
    # Récupérer unknown_titles depuis patterns
    unknown_top = result.patterns.get("unknown_titles_top", {})
    if not unknown_top:
        return None
    
    # Créer dossier artifacts si nécessaire
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    
    # Timestamp pour versioning
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    csv_path = artifacts_dir / f"unknown_titles_{timestamp}.csv"
    
    # Préparer rows triés par count décroissant
    rows = []
    for title_norm, count in sorted(unknown_top.items(), key=lambda x: x[1], reverse=True):
        # Heuristiques de suggestion
        suggested_action, suggested_target, notes = _suggest_title_action(title_norm, count)
        
        # Title raw = title_norm (on n'a pas l'original, mais on pourrait le tracker)
        # Pour l'instant, utiliser title_norm
        title_raw = title_norm  # TODO: stocker original dans unknown_titles si besoin
        
        rows.append({
            "title_raw": title_raw,
            "title_norm": title_norm,
            "count": count,
            "suggested_action": suggested_action,
            "suggested_target": suggested_target,
            "notes": notes,
        })
    
    # Écrire CSV
    with open(csv_path, "w", encoding="utf-8", newline="") as f:
        fieldnames = ["title_raw", "title_norm", "count", "suggested_action", "suggested_target", "notes"]
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)
    
    return str(csv_path)


def _suggest_title_action(title_norm: str, count: int) -> tuple[str, str, str]:
    """
    Suggère une action pour un unknown_title.
    
    Returns:
        (action, target, notes)
        action ∈ {MAP_TO_SECTION, MAP_TO_TESTS, SUBHEADING_POLICY, IGNORE}
    """
    # Règle 1 : count = 1 → IGNORE (sauf si looks like section canonique)
    if count == 1:
        # Vérifier si ressemble à une section canonique (heuristique simple)
        canonical_keywords = {
            "FORMATION", "COMPETENCES", "PARCOURS", "PROJET", "PISTES", 
            "BILAN", "SYNTHESE", "RECOMMANDATIONS", "FREINS", "ATOUTS"
        }
        if any(kw in title_norm for kw in canonical_keywords):
            return ("MAP_TO_SECTION", "À déterminer", "One-shot mais keywords canoniques détectés")
        return ("IGNORE", "", "One-shot, pas prioritaire")
    
    # Règle 2 : Déjà subheading pattern ? → SUBHEADING_POLICY
    # (Vérifier si correspondrait à is_subheading si on l'appelait)
    if _looks_like_subheading(title_norm):
        return ("SUBHEADING_POLICY", "", f"Pattern subheading détecté, améliorer règles (count={count})")
    
    # Règle 3 : Keywords tests/évaluations → MAP_TO_TESTS
    test_keywords = {
        "TEST", "EVALUATION", "FRANCAIS", "WORD", "EXCEL", "POWERPOINT", 
        "OUTLOOK", "POSITIONNEMENT", "NIVEAU"
    }
    if any(kw in title_norm for kw in test_keywords):
        return ("MAP_TO_TESTS", "tests", f"Keywords tests détectés (count={count})")
    
    # Règle 4 : Keywords sections canoniques → MAP_TO_SECTION
    section_hints = {
        "FORMATION": "formation",
        "COMPETENCES": "competences",
        "PARCOURS": "parcours_professionnel",
        "PROJET": "projet_professionnel",
        "PISTES": "pistes_metiers",
        "BILAN": "bilan",
        "SYNTHESE": "synthese",
        "RECOMMANDATIONS": "recommandations",
        "FREINS": "freins",
        "ATOUTS": "atouts",
    }
    for kw, target in section_hints.items():
        if kw in title_norm:
            return ("MAP_TO_SECTION", target, f"Keyword '{kw}' détecté (count={count})")
    
    # Règle 5 : Fréquence élevée (≥3) → MAP_TO_SECTION
    if count >= 3:
        return ("MAP_TO_SECTION", "À déterminer", f"Fréquence élevée (count={count}), analyser manuellement")
    
    # Fallback : count = 2, pas de pattern évident
    return ("MAP_TO_SECTION", "À déterminer", f"Fréquence moyenne (count={count}), évaluer au cas par cas")


def _looks_like_subheading(title_norm: str) -> bool:
    """Heuristique rapide pour détecter subheading patterns."""
    import re
    
    # Liste numérotée
    if re.match(r'^\d+\.', title_norm):
        return True
    
    # Phrase longue (> 8 mots)
    if len(title_norm.split()) > 8:
        return True
    
    # Étiquette simple (approximation : ≤ 2 mots avant premier espace, reste > 2 mots)
    # Note: title_norm n'a plus de ':', mais on peut détecter pattern "MOT VALEUR VALEUR..."
    tokens = title_norm.split()
    if len(tokens) > 3 and len(tokens[0]) < 15:
        # Heuristique faible, mais mieux que rien
        return False  # Pas assez fiable sans ':'
    
    return False


def _detect_common_structures(clients: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Détecte les structures communes dans les clients."""
    # Pour l'instant, simple comptage des patterns de dossiers
    folder_patterns = Counter()
    
    for client in clients:
        if "sources_by_type" in client:
            pattern = ",".join(sorted(client["sources_by_type"].keys()))
            folder_patterns[pattern] += 1
    
    return {
        "top_patterns": dict(folder_patterns.most_common(5)),
        "total_patterns": len(folder_patterns),
    }


def _build_training_state(result: DatasetTrainingResult) -> Dict[str, Any]:
    """Construit le training_state.json selon schéma v1.0 spec."""
    # run_id format: DATASETNAME_2025-12-27T19:32:37Z_randomhex
    now_iso = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    random_hex = result.dataset_id[:6]
    dataset_name = Path(result.clients[0]["folder_path"]).parent.name if result.clients else "dataset"
    run_id = f"{dataset_name}_{now_iso}_{random_hex}"
    
    # Compter extensions (fichiers totaux)
    ext_counts = result.stats["extensions_distribution"]
    total_clients = result.stats["successful_scans"]
    
    # Simples stats par type de doc
    doc_types_stats = {ext: count for ext, count in ext_counts.items()}
    
    # Gold stats
    gold_detected = result.stats["gold_detected"]
    gold_missing = total_clients - gold_detected
    
    # Construire section_title_map complet (seed + learned)
    full_title_map = {**SEED_SECTION_TITLE_MAP}
    if "learned_title_map" in result.patterns:
        full_title_map.update(result.patterns["learned_title_map"])
    
    # Section stats enrichies (coverage_pct + lines)
    sections_stats_v1 = {}
    field_max_lines = {
        "NAME": 1,
        "SURNAME": 1,
        "NUMERO_AVS": 1,
        "PROFESSION": 4,
        "FORMATION": 10,
        "Ressources_comportementales_Points_d'appui": 4,
        "Ressources_comportementales_Points_de_vigilance": 4,
        # FIX ESSAI 100: Sections canoniques RESSOURCES_* ne doivent pas être 0
        "RESSOURCES_POINTS_APPUI": 6,
        "RESSOURCES_POINTS_VIGILANCE": 6
    }
    
    if "sections_stats" in result.patterns:
        for canonical, stats in result.patterns["sections_stats"].items():
            # ✅ Garde-fou : s'assurer que coverage_pct est entre 0 et 100
            coverage_pct = int(round(float(stats.get("coverage", 0)) * 100))
            coverage_pct = max(0, min(100, coverage_pct))  # Clamp 0-100
            
            sections_stats_v1[canonical.upper()] = {
                "coverage_pct": coverage_pct,
                "clients": stats.get("clients", 0),  # ✅ Nombre de clients
                "lines": {
                    "avg": stats["avg_lines"],
                    "median": int(stats["p50_lines"]),
                    "p90": int(stats["p90_lines"])
                }
            }
            # Ajouter aux field_max_lines (utiliser p90)
            field_max_lines[canonical.upper()] = int(stats["p90_lines"])
    
    # ✅ Warnings : vérifier support .msg
    warnings = []
    
    # Vérifier si extract-msg est disponible
    try:
        from core.extractors.msg_extractor import MSG_SUPPORT_AVAILABLE
    except ImportError:
        MSG_SUPPORT_AVAILABLE = False
    
    # Si .msg présents et extract-msg non installé -> warning
    if ".msg" in ext_counts and ext_counts[".msg"] > 0:
        if not MSG_SUPPORT_AVAILABLE:
            warnings.append({
                "code": "MSG_EXTRACTOR_MISSING",
                "message": "Des fichiers .msg sont présents mais extract-msg n'est pas installé (pip install extract-msg>=0.48.0)",
                "count": ext_counts[".msg"]
            })
    
    root_path = str(Path(result.clients[0]["folder_path"]).parent) if result.clients else ""
    
    return {
        "schema_version": "training_state_v1.0",
        "run_id": run_id,
        "created_at": now_iso,
        
        "dataset": {
            "root_path": root_path,
            "dataset_id": result.dataset_id,
            "clients_scanned": result.stats["total_clients"],
            "clients_used": total_clients,
            
            "doc_types_stats": doc_types_stats,
            
            "gold_stats": {
                "gold_detected_clients": gold_detected,
                "gold_missing_clients": gold_missing
            }
        },
        
        "conventions": {
            "fallback_value": "Non renseigné",
            "status_enum": ["GO", "NO_GO", "DRAFT"],
            "scores": {
                "coverage_range": [0, 100],
                "quality_range": [0, 1],
                "confidence_range": [0, 1]
            }
        },
        
        "profiles": {
            "STRICT": {
                "coverage_min": 85,
                "quality_min": 0.75,
                "confidence_min": 0.70,
                "sources_count_min": 1,
                "profession_or_formation_required": True
            },
            "STANDARD": {
                "coverage_min": 75,
                "quality_min": 0.65,
                "confidence_min": 0.60,
                "sources_count_min": 1,
                "profession_or_formation_required": True
            },
            "DRAFT": {
                "coverage_min": 0,
                "quality_min": 0.0,
                "confidence_min": 0.0,
                "sources_count_min": 0,
                "profession_or_formation_required": False
            }
        },
        
        "patterns": {
            "section_stats": sections_stats_v1,
            "field_max_lines": field_max_lines,
            "section_title_map": full_title_map,  # ✅ Mapping complet (seed + learned)
            "unknown_titles_top": result.patterns.get("unknown_titles_top", {}),  # ✅ Top 50
            "unknown_titles_count": result.patterns.get("unknown_titles_count", 0),  # ✅ Distinct count
            "unknown_titles_total_occurrences": result.patterns.get("unknown_titles_total_occurrences", 0),  # ✅ Total
        },
        
        "warnings": warnings
    }


def _merge_training_states(
    existing: Dict[str, Any],
    new: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Fusionne deux training_states (compatible training_state_v1.0) - VERSION SAFE.
    
    V4.1 Fix: Ne plante JAMAIS même si les schémas diffèrent.
    Base = new (on garde metadata/schema à jour).
    Fusionne uniquement les patterns non-nominatifs de manière défensive.
    """
    import copy
    
    # ✅ Base = copie profonde de new pour éviter mutations
    merged = copy.deepcopy(new)
    
    # ✅ Défensif : vérifier que patterns existe dans merged
    if "patterns" not in merged:
        merged["patterns"] = {}
    
    try:
        # ✅ 1. Fusionner field_max_lines (prendre le max)
        if "patterns" in existing and "field_max_lines" in existing["patterns"]:
            old_max = existing["patterns"]["field_max_lines"]
            new_max = merged["patterns"].get("field_max_lines", {})
            
            if "field_max_lines" not in merged["patterns"]:
                merged["patterns"]["field_max_lines"] = {}
            
            for field, max_val in old_max.items():
                current = merged["patterns"]["field_max_lines"].get(field, 0)
                merged["patterns"]["field_max_lines"][field] = max(max_val, current)
    except Exception as e:
        # Silencieux : si échec, on garde new intact
        pass
    
    try:
        # ✅ 2. Fusionner section_stats (max de p90, max de coverage)
        if "patterns" in existing and "section_stats" in existing["patterns"]:
            old_sections = existing["patterns"]["section_stats"]
            
            if "section_stats" not in merged["patterns"]:
                merged["patterns"]["section_stats"] = {}
            
            new_sections = merged["patterns"]["section_stats"]
            
            # Union des sections
            all_sections = set(old_sections.keys()) | set(new_sections.keys())
            
            for sec in all_sections:
                old_st = old_sections.get(sec, {})
                new_st = new_sections.get(sec, {})
                
                # Initialiser section si absente
                if sec not in merged["patterns"]["section_stats"]:
                    merged["patterns"]["section_stats"][sec] = {}
                
                merged_sec = merged["patterns"]["section_stats"][sec]
                
                # Fusionner lines (prendre max de p90)
                old_lines = old_st.get("lines", {})
                new_lines = new_st.get("lines", {})
                
                if "lines" not in merged_sec:
                    merged_sec["lines"] = {}
                
                # p90 = max
                old_p90 = old_lines.get("p90", 0)
                new_p90 = new_lines.get("p90", 0)
                merged_sec["lines"]["p90"] = max(old_p90, new_p90)
                
                # Autres stats : prendre new par défaut
                for key in ["median", "mean", "min", "max", "p10"]:
                    if key in new_lines:
                        merged_sec["lines"][key] = new_lines[key]
                    elif key in old_lines:
                        merged_sec["lines"][key] = old_lines[key]
                
                # Coverage : prendre le max (meilleur coverage observé)
                old_cov = old_st.get("coverage_pct", 0)
                new_cov = new_st.get("coverage_pct", 0)
                merged_sec["coverage_pct"] = max(old_cov, new_cov)
                
                # Autres champs : garder new
                for key in ["clients_with_section", "total_blocks"]:
                    if key in new_st:
                        merged_sec[key] = new_st[key]
                    elif key in old_st:
                        merged_sec[key] = old_st[key]
    except Exception as e:
        # Silencieux : si échec, on garde new intact
        pass
    
    try:
        # ✅ 3. Fusionner warnings (union par code)
        if "warnings" in existing:
            old_warnings = existing["warnings"]
            
            if "warnings" not in merged:
                merged["warnings"] = []
            
            # Union par code (éviter doublons)
            existing_codes = {w["code"] for w in merged["warnings"] if "code" in w}
            
            for warn in old_warnings:
                if "code" in warn and warn["code"] not in existing_codes:
                    merged["warnings"].append(warn)
                    existing_codes.add(warn["code"])
    except Exception as e:
        # Silencieux : si échec, on garde new intact
        pass
    
    try:
        # ✅ 4. Historique (optionnel)
        history_entry = {
            "run_id": new.get("training_state_id"),
            "timestamp": new.get("generated_at"),
            "clients": new.get("dataset", {}).get("clients_used", 0)
        }
        
        if "history" not in merged:
            merged["history"] = []
        
        # Ajouter historique de existing si présent
        if "history" in existing:
            for entry in existing["history"]:
                if entry not in merged["history"]:
                    merged["history"].append(entry)
        
        # Ajouter nouveau run
        merged["history"].append(history_entry)
    except Exception as e:
        # Silencieux : pas critique
        pass
    
    return merged


def _merge_counters(c1: Dict[str, int], c2: Dict[str, int]) -> Dict[str, int]:
    """Fusionne deux compteurs."""
    merged = c1.copy()
    for key, count in c2.items():
        merged[key] = merged.get(key, 0) + count
    return merged


def _generate_training_report_md(result: DatasetTrainingResult) -> str:
    """Génère un rapport markdown lisible."""
    lines = [
        f"# Training Report - {result.dataset_id}",
        f"",
        f"**Date** : {result.timestamp}",
        f"**Clients analysés** : {result.stats['total_clients']}",
        f"**Scans réussis** : {result.stats['successful_scans']}",
        # FIX ESSAI 100 (AC3): Distinguer clients utilisables (sources>=1) vs sources=0
        f"**Clients utilisables (sources≥1)** : {result.stats['clients_used']} ({result.stats['clients_used']/result.stats['successful_scans']*100:.0f}%)" if result.stats.get('successful_scans', 0) > 0 else f"**Clients utilisables (sources≥1)** : {result.stats['clients_used']}",
        f"**Clients sans sources (sources=0)** : {result.stats['clients_no_sources']}",
        f"",
        f"## 📊 Statistiques Globales",
        f"",
        f"- **GOLD détectés** : {result.stats['gold_detected']} ({result.stats['gold_detection_rate']:.1%})",
        f"- **Pipeline ready** : {result.stats['pipeline_ready']} ({result.stats['pipeline_ready_rate']:.1%})",
        f"",
        f"### Sources",
        f"",
        f"- **Moyenne** : {result.stats['sources_stats']['mean']:.1f}",
        f"- **Médiane** : {result.stats['sources_stats']['median']:.1f}",
        f"- **Min / Max** : {result.stats['sources_stats']['min']} / {result.stats['sources_stats']['max']}",
        f"- **P10 / P90** : {result.stats['sources_stats']['p10']:.1f} / {result.stats['sources_stats']['p90']:.1f}",
        f"",
        f"### Extensions",
        f"",
    ]
    
    for ext, count in sorted(result.stats['extensions_distribution'].items(), key=lambda x: -x[1]):
        lines.append(f"- `{ext}` : {count}")
    
    # ✅ Ajouter section erreurs si présentes
    if result.stats.get('errors', 0) > 0:
        lines.extend([
            f"",
            f"## ❌ Erreurs ({result.stats['errors']} clients)",
            f"",
        ])
        
        errors_top = result.stats.get('errors_top', [])
        if errors_top:
            lines.append("### Top Erreurs")
            lines.append("")
            for error_type, count in errors_top:
                lines.append(f"- **{error_type}** : {count} client(s)")
            lines.append("")
        
        # Lister les clients en erreur avec détails
        error_clients = [c for c in result.clients if "error" in c]
        if error_clients:
            lines.append("### Détail par client")
            lines.append("")
            for client in error_clients[:10]:  # Limiter à 10
                lines.append(f"- **{client['folder_name']}** : {client.get('error_type', 'Error')} - {client['error']}")
            if len(error_clients) > 10:
                lines.append(f"- ... et {len(error_clients) - 10} autres")
    
    # ✅ Ajouter section stats avec clients_used
    if "sections_stats" in result.patterns:
        lines.extend([
            f"",
            f"## 📑 Sections Canoniques",
            f"",
            f"Coverage basée sur **{result.stats['successful_scans']} clients utilisés**.",
            f"",
            f"| Section | Coverage % | Clients | Avg Lines | P90 Lines |",
            f"|---------|------------|---------|-----------|-----------|",
        ])
        
        for canonical, stats in sorted(result.patterns["sections_stats"].items()):
            coverage_pct = int(stats["coverage"] * 100)
            clients = stats.get("clients", 0)
            avg = stats["avg_lines"]
            p90 = stats["p90_lines"]
            lines.append(f"| {canonical.upper()} | {coverage_pct}% | {clients} | {avg:.1f} | {p90:.1f} |")
    
    lines.extend([
        f"",
        f"## 🎯 Patterns Détectés",
        f"",
        f"### Titres Inconnus (Top 10)",
        f"",
    ])
    
    for title, count in result.patterns.get("unknown_titles_top10", {}).items():
        lines.append(f"- `{title}` : {count} occurrences")
    
    lines.extend([
        f"",
        f"## 💡 Recommandations",
        f"",
    ])
    
    for rec in result.recommendations:
        lines.append(f"- {rec}")
    
    lines.append(f"")
    
    return "\n".join(lines)
