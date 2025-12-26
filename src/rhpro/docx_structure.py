"""
DOCX structure extractor — extrait paragraphes avec métadonnées via python-docx
"""
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.text.paragraph import Paragraph as DocxParagraph


@dataclass
class Paragraph:
    """Représente un paragraphe avec ses métadonnées"""
    text: str
    style_name: str
    is_bold: bool
    font_size: Optional[float]  # en points
    is_all_caps: bool
    numbering_prefix: str  # ex: "2.1.1"
    
    def __repr__(self):
        return f"<Paragraph style={self.style_name} bold={self.is_bold} text={self.text[:50]}...>"


class DocxStructureExtractor:
    """Extrait la structure d'un document DOCX"""
    
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.document: Optional[Document] = None
        self.paragraphs: List[Paragraph] = []
    
    def load(self):
        """Charge le document et extrait les paragraphes"""
        if not self.docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {self.docx_path}")
        
        self.document = Document(str(self.docx_path))
        self.paragraphs = self._extract_paragraphs()
    
    def _extract_paragraphs(self) -> List[Paragraph]:
        """Extrait tous les paragraphes avec métadonnées"""
        result = []
        
        for para in self.document.paragraphs:
            # Ignorer les paragraphes vides
            text = para.text.strip()
            if not text:
                continue
            
            # Extraire métadonnées
            style_name = para.style.name if para.style else ""
            is_bold = self._is_bold(para)
            font_size = self._get_font_size(para)
            is_all_caps = text.isupper() and len(text) > 3
            numbering_prefix = self._extract_numbering(para)
            
            result.append(Paragraph(
                text=text,
                style_name=style_name,
                is_bold=is_bold,
                font_size=font_size,
                is_all_caps=is_all_caps,
                numbering_prefix=numbering_prefix
            ))
        
        return result
    
    def _is_bold(self, para: DocxParagraph) -> bool:
        """Détermine si le paragraphe est en gras"""
        # Vérifier les runs
        for run in para.runs:
            if run.bold:
                return True
        return False
    
    def _get_font_size(self, para: DocxParagraph) -> Optional[float]:
        """Récupère la taille de police (moyenne si mixte)"""
        sizes = []
        for run in para.runs:
            if run.font.size:
                # Convertir de twips (1/20 point) en points
                sizes.append(run.font.size.pt)
        
        if sizes:
            return sum(sizes) / len(sizes)
        return None
    
    def _extract_numbering(self, para: DocxParagraph) -> str:
        """Extrait le préfixe de numérotation (ex: "2.1.1")"""
        import re
        text = para.text.strip()
        
        # Pattern pour numérotation hiérarchique
        match = re.match(r'^(\d+(?:\.\d+){1,6})\s*\.?\s+', text)
        if match:
            return match.group(1)
        
        return ""
    
    def get_paragraphs(self) -> List[Paragraph]:
        """Retourne la liste des paragraphes extraits"""
        if not self.paragraphs:
            self.load()
        return self.paragraphs


def extract_paragraphs_from_docx(docx_path: str) -> List[Paragraph]:
    """Helper function pour extraire les paragraphes d'un DOCX"""
    extractor = DocxStructureExtractor(docx_path)
    return extractor.get_paragraphs()
