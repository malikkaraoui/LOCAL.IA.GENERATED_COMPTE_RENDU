"""
Module de génération de comptes-rendus RH-Pro avec remplissage template DOCX.

Intègre RAG + remplissage template + exports structurés.
"""
import json
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import shutil

from docx import Document

from .rag_generator import RAGGenerator
from .validation_profiles import validate_report, ValidationProfile


# Template RH-Pro par défaut (champs attendus)
DEFAULT_TEMPLATE_FIELDS = [
    "nom",
    "prenom",
    "date_naissance",
    "adresse",
    "telephone",
    "email",
    "situation_professionnelle",
    "niveau_formation",
    "experience_professionnelle",
    "competences_principales",
    "competences_transversales",
    "objectifs_professionnels",
    "projet_formation",
    "freins_identifies",
    "atouts_mobilisables",
    "parcours_ai",
    "tests_realises",
    "resultats_tests",
    "bilan_orientation",
    "preconisations",
    "suivi_propose",
]


class RHProReportGenerator:
    """
    Générateur de comptes-rendus RH-Pro.
    """
    
    def __init__(
        self,
        template_path: Optional[str] = None,
        template_fields: Optional[List[str]] = None,
    ):
        """
        Initialise le générateur.
        
        Args:
            template_path: Chemin vers le template DOCX (optionnel)
            template_fields: Liste des champs à remplir (par défaut : DEFAULT_TEMPLATE_FIELDS)
        """
        self.template_path = template_path
        self.template_fields = template_fields or DEFAULT_TEMPLATE_FIELDS
        self.rag_generator = None
    
    def generate_from_client(
        self,
        sources_folder: str,
        gold_path: Optional[str] = None,
        output_dir: str = "output",
        client_name: str = "client",
        strict_mode: bool = True,
        validation_profile: Optional[ValidationProfile] = None,
        training_state: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Génère un compte-rendu complet pour un client.
        
        Pipeline :
        1. Construire index RAG depuis sources/
        2. Extraire champs via RAG avec garde-fous
        3. Remplir template DOCX
        4. Générer outputs : generated.docx, debug.json, metrics.json
        5. Valider selon profil (GO/NO-GO/DRAFT)
        
        Args:
            sources_folder: Dossier contenant les sources RAG
            gold_path: Chemin vers le GOLD (optionnel, pour référence)
            output_dir: Dossier de sortie
            client_name: Nom du client (pour nommage fichiers)
            strict_mode: Interdire l'invention (retourne "Non renseigné" si non trouvé)
            validation_profile: Profil de validation (STRICT/STANDARD/DRAFT)
            training_state: État de training optionnel (patterns appris)
            
        Returns:
            Dict avec chemins des outputs, métriques et statut de validation
        """
        # Créer dossier output
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        # 1. Construire index RAG (avec training_state)
        self.rag_generator = RAGGenerator(
            chunk_size=512,
            chunk_overlap=50,
            llm_model="gpt-4o-mini",
            temperature=0.1,
            training_state=training_state,
        )
        
        index_result = self.rag_generator.build_index_from_sources(sources_folder)
        
        # 2. Générer rapport structuré via RAG
        report_result = self.rag_generator.generate_report(
            template_fields=self.template_fields,
            strict_mode=strict_mode,
        )
        
        # 3. Remplir template DOCX
        docx_path = self._fill_docx_template(
            fields=report_result["fields"],
            output_dir=output_path,
            client_name=client_name,
        )
        
        # 4. Générer debug.json
        debug_path = output_path / f"{client_name}_debug.json"
        self._export_debug_json(
            report_result=report_result,
            index_result=index_result,
            gold_path=gold_path,
            output_path=debug_path,
            client_root=sources_folder,  # Utiliser sources_folder comme client_root
            template_docx=self.template_path,
        )
        
        # 5. Générer metrics.json
        metrics_path = output_path / f"{client_name}_metrics.json"
        self._export_metrics_json(
            metrics=report_result["metrics"],
            output_path=metrics_path,
        )
        
        # 6. Copier le GOLD si fourni (pour comparaison)
        gold_copy_path = None
        if gold_path and Path(gold_path).exists():
            gold_copy_path = output_path / f"{client_name}_gold_reference.docx"
            shutil.copy(gold_path, gold_copy_path)
        
        # 7. Validation selon profil
        validation_result = None
        if validation_profile:
            validation_result = validate_report(
                metrics_path=metrics_path,
                debug_path=debug_path,
                meta_path=None,  # meta.json peut être ajouté si disponible
                profile=validation_profile,
            )
            
            # Exporter validation.json
            validation_path = output_path / f"{client_name}_validation.json"
            with open(validation_path, 'w', encoding='utf-8') as f:
                f.write(validation_result.to_json())
        
        return {
            "success": True,
            "client_name": client_name,
            "outputs": {
                "generated_docx": str(docx_path),
                "debug_json": str(debug_path),
                "metrics_json": str(metrics_path),
                "gold_reference": str(gold_copy_path) if gold_copy_path else None,
                "validation_json": str(output_path / f"{client_name}_validation.json") if validation_result else None,
            },
            "metrics": report_result["metrics"],
            "validation": validation_result.to_dict() if validation_result else None,
            "index_stats": {
                "sources_count": index_result["sources_count"],
                "chunks_created": index_result["chunks_created"],
            },
            "timestamp": datetime.now().isoformat(),
        }
    
    def _fill_docx_template(
        self,
        fields: Dict[str, str],
        output_dir: Path,
        client_name: str,
    ) -> Path:
        """
        Remplit le template DOCX avec les champs générés.
        
        Args:
            fields: Champs remplis par RAG
            output_dir: Dossier de sortie
            client_name: Nom du client
            
        Returns:
            Chemin du DOCX généré
        """
        output_path = output_dir / f"{client_name}_generated.docx"
        
        if self.template_path and Path(self.template_path).exists():
            # Utiliser le template existant
            doc = Document(self.template_path)
            
            # Remplacer les placeholders
            for field, value in fields.items():
                placeholder = f"{{{{{field}}}}}"  # Ex: {{nom}}
                self._replace_placeholder_in_doc(doc, placeholder, value)
        
        else:
            # Créer un document simple si pas de template
            doc = Document()
            doc.add_heading(f"Compte-Rendu RH-Pro - {client_name}", level=1)
            doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph("")
            
            for field, value in fields.items():
                doc.add_heading(field.replace("_", " ").title(), level=2)
                doc.add_paragraph(value)
        
        # Sauvegarder
        doc.save(str(output_path))
        
        return output_path
    
    def _replace_placeholder_in_doc(
        self,
        doc: Document,
        placeholder: str,
        value: str,
    ) -> None:
        """
        Remplace un placeholder dans un document DOCX.
        
        Args:
            doc: Document DOCX
            placeholder: Placeholder à remplacer (ex: {{nom}})
            value: Valeur de remplacement
        """
        # Remplacer dans les paragraphes
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Remplacer dans les tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
        
        # Remplacer dans headers/footers
        for section in doc.sections:
            for header in [section.header, section.footer]:
                for paragraph in header.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)
    
    def _export_debug_json(
        self,
        report_result: Dict[str, Any],
        index_result: Dict[str, Any],
        gold_path: Optional[str],
        output_path: Path,
        client_root: Optional[str] = None,
        template_docx: Optional[str] = None,
    ) -> None:
        """
        Exporte les informations de debug en JSON (schéma v1.0).
        
        Args:
            report_result: Résultat de génération RAG
            index_result: Résultat d'indexation
            gold_path: Chemin vers le GOLD
            output_path: Chemin du fichier JSON de sortie
            client_root: Racine du dossier client
            template_docx: Chemin vers le template DOCX utilisé
        """
        # Générer run_id unique
        now_iso = datetime.now().strftime("%Y%m%dT%H%M%SZ")
        run_id = f"run_{now_iso}_{output_path.stem[-8:]}"
        
        # Structurer les sources avec détails complets
        sources_detailed = []
        for source in index_result.get("sources", []):
            sources_detailed.append({
                "source_id": Path(source["path"]).name,
                "path": source["path"],
                "type": source.get("extension", Path(source["path"]).suffix),
                "loaded": source.get("loaded", True),
                "error": source.get("error", None)
            })
        
        # Structurer les champs avec value, confidence, citations, warnings
        fields_structured = {}
        for field_name, field_data in report_result["debug"].items():
            # Extraire citations depuis les sources
            citations = []
            if isinstance(field_data, dict) and "sources" in field_data:
                for source in field_data["sources"][:3]:  # Max 3 citations
                    citations.append({
                        "source_id": Path(source.get("file", "unknown")).name,
                        "snippet": source.get("snippet", "")[:200],  # Limiter à 200 chars
                        "score": source.get("score", 0.0)
                    })
            
            # Déterminer warnings
            warnings = []
            value = field_data.get("value", "Non renseigné") if isinstance(field_data, dict) else str(field_data)
            confidence = field_data.get("confidence", 0.0) if isinstance(field_data, dict) else 0.0
            
            if value == "Non renseigné" or not value:
                warnings.append("no_evidence")
            elif confidence < 0.5:
                warnings.append("low_confidence")
            
            fields_structured[field_name] = {
                "value": value,
                "confidence": round(confidence, 2),
                "citations": citations,
                "warnings": warnings
            }
        
        # Construire debug.json v1.0
        debug_data = {
            "schema_version": "1.0",
            "artifact_type": "debug",
            "run_id": run_id,
            "created_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"),
            
            "conventions": {
                "language": "fr",
                "fallback_value": "Non renseigné",
                "strict_mode_default": True
            },
            
            "inputs": {
                "dataset_root": None,  # Peut être ajouté si disponible
                "client_root": client_root,
                "template_docx": template_docx or self.template_path
            },
            
            "index": {
                "sources_count": index_result["sources_count"],
                "documents_loaded": len([s for s in sources_detailed if s["loaded"]]),
                "chunks_created": index_result["chunks_created"],
                "sources": sources_detailed
            },
            
            "fields": fields_structured
        }
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(debug_data, f, indent=2, ensure_ascii=False)
    
    def _export_metrics_json(
        self,
        metrics: Dict[str, Any],
        output_path: Path,
    ) -> None:
        """
        Exporte les métriques en JSON.
        
        Args:
            metrics: Métriques de couverture
            output_path: Chemin du fichier JSON de sortie
        """
        metrics_data = {
            "timestamp": datetime.now().isoformat(),
            "required_coverage": metrics["required_coverage_pct"],
            "weighted_coverage": metrics["coverage_pct"],
            "quality_score": metrics["quality_score"],
            "avg_confidence": metrics["avg_confidence"],
            "total_fields": metrics["total_fields"],
            "filled_fields": metrics["filled_fields"],
            "required_fields": metrics["required_fields"],
            "required_filled": metrics["required_filled"],
        }
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(metrics_data, f, indent=2, ensure_ascii=False)
    
    def _generate_warnings(self, report_result: Dict[str, Any]) -> List[str]:
        """
        Génère des warnings basés sur les résultats.
        
        Args:
            report_result: Résultat de génération
            
        Returns:
            Liste de warnings
        """
        warnings = []
        
        metrics = report_result["metrics"]
        
        # Couverture faible
        if metrics["coverage_pct"] < 50:
            warnings.append(
                f"⚠️  Couverture faible : {metrics['coverage_pct']}%"
            )
        
        # Confiance faible
        if metrics["avg_confidence"] < 0.5:
            warnings.append(
                f"⚠️  Confiance moyenne faible : {metrics['avg_confidence']:.2f}"
            )
        
        # Champs requis manquants
        missing_required = metrics["required_fields"] - metrics["required_filled"]
        if missing_required > 0:
            warnings.append(
                f"❌ {missing_required} champ(s) requis manquant(s)"
            )
        
        # Champs sans citations
        debug = report_result["debug"]
        no_citation_count = sum(
            1 for field_info in debug.values()
            if not field_info.get("citations")
        )
        if no_citation_count > len(debug) * 0.3:  # > 30%
            warnings.append(
                f"⚠️  {no_citation_count} champs sans citations"
            )
        
        if not warnings:
            warnings.append("✅ Génération réussie sans avertissements majeurs")
        
        return warnings
    
    def _structure_evidence(self, debug_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Structure les preuves par catégorie selon le pattern evidence.category.field[].
        
        Règle : no-evidence = no-claim
        Chaque champ doit avoir des preuves traçables.
        
        Args:
            debug_info: Informations de debug avec evidence par champ
            
        Returns:
            Dict structuré avec evidence.identity.nom[], evidence.professional.situation[], etc.
        """
        evidence = {
            "identity": {},
            "professional": {},
            "contact": {},
            "other": {},
        }
        
        # Mapping des champs vers les catégories
        field_categories = {
            "nom": "identity",
            "prenom": "identity",
            "civilite": "identity",
            "date_naissance": "identity",
            "numero_avs": "identity",
            "situation_professionnelle": "professional",
            "niveau_formation": "professional",
            "experience_professionnelle": "professional",
            "competences_principales": "professional",
            "telephone": "contact",
            "email": "contact",
            "adresse": "contact",
        }
        
        for field, field_data in debug_info.items():
            if isinstance(field_data, dict) and "evidence" in field_data:
                # Déterminer la catégorie
                category = field_categories.get(field, "other")
                
                # Extraire les preuves
                field_evidence = field_data.get("evidence", [])
                
                # Stocker les preuves uniquement si non vides (no-evidence = no-claim)
                if field_evidence and len(field_evidence) > 0:
                    evidence[category][field] = field_evidence
                else:
                    # Pas de preuves = valeur non fiable
                    evidence[category][field] = []
        
        return evidence


def generate_report_from_normalized(
    normalized_folder: str,
    output_dir: str = "output",
    template_path: Optional[str] = None,
    strict_mode: bool = True,
    training_state: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Génère un rapport depuis un dossier client normalisé.
    
    Structure attendue :
    normalized_folder/
      ├── sources/         ← Sources RAG
      ├── gold/           ← GOLD de référence
      └── meta.json       ← Métadonnées
    
    Args:
        normalized_folder: Dossier client normalisé (sandbox)
        output_dir: Dossier de sortie
        template_path: Template DOCX (optionnel)
        strict_mode: Mode strict (pas d'invention)
        
    Returns:
        Résultat de génération
    """
    normalized_path = Path(normalized_folder)
    
    if not normalized_path.exists():
        raise FileNotFoundError(f"Dossier normalisé introuvable : {normalized_folder}")
    
    # Chemins
    sources_folder = normalized_path / "sources"
    gold_folder = normalized_path / "gold"
    meta_path = normalized_path / "meta.json"
    
    if not sources_folder.exists():
        raise FileNotFoundError(f"Dossier sources introuvable : {sources_folder}")
    
    # Lire métadonnées si disponibles
    client_name = normalized_path.name
    if meta_path.exists():
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
            client_name = meta.get("client_slug", client_name)
    
    # Trouver le GOLD
    gold_path = None
    if gold_folder.exists():
        gold_files = list(gold_folder.glob("*.docx"))
        if gold_files:
            gold_path = str(gold_files[0])
    
    # Générer
    generator = RHProReportGenerator(
        template_path=template_path,
        template_fields=DEFAULT_TEMPLATE_FIELDS,
    )
    
    return generator.generate_from_client(
        sources_folder=str(sources_folder),
        gold_path=gold_path,
        output_dir=output_dir,
        client_name=client_name,
        strict_mode=strict_mode,
        training_state=training_state,
    )
