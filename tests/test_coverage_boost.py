"""Tests supplémentaires pour améliorer la couverture."""

import pytest
from pathlib import Path
from core.extract import extract_pdf, extract_docx, extract_txt, sha256_text, normalize_text
from core.avs import detect_avs_in_text, _normalize_avs
from core.context import make_chunks, build_index
from core.models import Chunk


class TestSha256Text:
    """Tests pour le hachage de texte."""

    def test_generates_hash(self):
        """Génère un hash SHA256."""
        result = sha256_text("test")
        assert isinstance(result, str)
        assert len(result) == 64  # SHA256 hex = 64 caractères

    def test_same_text_same_hash(self):
        """Même texte produit même hash."""
        text = "hello world"
        hash1 = sha256_text(text)
        hash2 = sha256_text(text)
        assert hash1 == hash2

    def test_different_text_different_hash(self):
        """Textes différents produisent hashes différents."""
        hash1 = sha256_text("text1")
        hash2 = sha256_text("text2")
        assert hash1 != hash2


class TestNormalizeTextExtract:
    """Tests pour normalize_text dans extract.py."""

    def test_normalizes_newlines(self):
        """Normalise les retours à la ligne."""
        result = normalize_text("line1\r\nline2\rline3")
        assert "\r" not in result

    def test_reduces_multiple_newlines(self):
        """Réduit les multiples sauts de ligne."""
        result = normalize_text("a\n\n\n\nb")
        assert "\n\n\n" not in result


class TestExtractPdf:
    """Tests pour l'extraction PDF."""

    def test_extracts_pdf_content(self, tmp_path):
        """Extrait le contenu d'un PDF."""
        import fitz
        
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Test PDF Content")
        
        pdf_file = tmp_path / "test.pdf"
        doc.save(pdf_file)
        doc.close()
        
        result = extract_pdf(pdf_file)
        
        assert result.success is True
        assert "text" in result.value
        assert "Test PDF Content" in result.value["text"]


class TestExtractDocx:
    """Tests pour l'extraction DOCX."""

    def test_extracts_docx_content(self, tmp_path):
        """Extrait le contenu d'un DOCX."""
        from docx import Document
        
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test DOCX Content")
        doc.save(docx_file)
        
        result = extract_docx(docx_file)
        
        assert result.success is True
        assert "text" in result.value
        assert "Test DOCX Content" in result.value["text"]

    def test_handles_tables(self, tmp_path):
        """Gère les tableaux."""
        from docx import Document
        
        docx_file = tmp_path / "table.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Cell Content"
        doc.save(docx_file)
        
        result = extract_docx(docx_file)
        
        assert result.success is True
        assert "text" in result.value
        # Le texte du tableau devrait être extrait
        assert "Cell Content" in result.value["text"]


class TestExtractTxt:
    """Tests pour l'extraction de fichiers texte."""

    def test_extracts_utf8(self, tmp_path):
        """Extrait du texte UTF-8."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Contenu français avec accents éàü", encoding="utf-8")
        
        result = extract_txt(txt_file)
        
        assert result.success is True
        assert "text" in result.value
        assert "Contenu français" in result.value["text"]

    def test_handles_large_file(self, tmp_path):
        """Gère un gros fichier."""
        txt_file = tmp_path / "large.txt"
        content = "ligne\n" * 10000
        txt_file.write_text(content, encoding="utf-8")
        
        result = extract_txt(txt_file)
        
        assert result.success is True
        assert "text" in result.value
        # Le fichier ne devrait pas être vide
        assert len(result.value["text"]) > 0


class TestDetectAvsInText:
    """Tests pour la détection d'AVS dans du texte."""

    def test_finds_avs_with_dots(self):
        """Trouve un AVS avec points."""
        text = "Le numéro AVS est 756.1234.5678.90 dans ce document."
        result = detect_avs_in_text(text)
        assert result is not None
        assert "756" in result

    def test_finds_avs_without_dots(self):
        """Trouve un AVS sans points."""
        text = "AVS: 7561234567890"
        result = detect_avs_in_text(text)
        assert result is not None

    def test_returns_none_when_not_found(self):
        """Retourne None si pas trouvé."""
        text = "Aucun numéro AVS ici."
        result = detect_avs_in_text(text)
        assert result is None


class TestNormalizeAvs:
    """Tests pour la normalisation d'AVS."""

    def test_formats_without_dots(self):
        """Formate un AVS sans points."""
        result = _normalize_avs("7561234567890")
        assert "." in result
        assert result.startswith("756")
        assert len(result) == 16  # 756.XXXX.XXXX.XX

    def test_formats_with_dots(self):
        """Formate un AVS avec points."""
        result = _normalize_avs("756.1234.5678.90")
        assert result == "756.1234.5678.90"

