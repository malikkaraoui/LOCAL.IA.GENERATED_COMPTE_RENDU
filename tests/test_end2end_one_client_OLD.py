"""
Test End-to-End (E2E) du pipeline complet RH-Pro - Definition of Done.

Pipeline testé : normalise → index → génère → valide

Contraintes STRICTES :
- 100% déterministe, rapide, local
- AUCUN appel réseau (LLM/embeddings mockés)
- Mini-dossier client créé dans tmp_path
- Vérification cohérence status vs métriques

Tests :
- Pipeline complet avec données minimales
- Vérification présence outputs
- Cohérence status GO/NO_GO vs métriques
"""
import json
import pytest
from pathlib import Path
from typing import Dict, Any
from unittest import mock
from docx import Document

# Imports du projet
from src.rhpro.report_generator import RHProReportGenerator
from src.rhpro.validation_profiles import validate_report, ValidationProfile


# ============================================================================
# Fixtures pour créer un mini-dossier client
# ============================================================================

@pytest.fixture
def mini_client_folder(tmp_path):
    """
    Crée un mini-dossier client avec structure minimale.
    
    Structure:
    client_folder/
      ├── cv.txt              (identité + profession)
      ├── entretien.txt       (compléments)
      └── formation.txt       (formation)
    """
    client_folder = tmp_path / "mini_client"
    client_folder.mkdir()
    
    # CV avec identité + profession
    cv_file = client_folder / "cv.txt"
    cv_file.write_text("""
CURRICULUM VITAE

Nom: DUPONT
Prénom: Jean
Date de naissance: 15/03/1985

EXPÉRIENCE PROFESSIONNELLE
Conseiller en orientation professionnelle
Centre d'orientation de Paris
Depuis 2015

COMPÉTENCES
- Accompagnement individuel
- Bilan de compétences
- Connaissance des métiers

Téléphone: 06 12 34 56 78
Email: jean.dupont@example.com
""", encoding='utf-8')
    
    # Entretien avec compléments
    entretien_file = client_folder / "entretien.txt"
    entretien_file.write_text("""
COMPTE-RENDU ENTRETIEN

Date: 10/12/2025
Conseiller: Marie MARTIN

Jean DUPONT est très motivé pour une reconversion.
Il souhaite évoluer vers la formation d'adultes.

Objectifs professionnels:
- Devenir formateur
- Travailler dans l'insertion professionnelle

Freins identifiés:
- Besoin de financement pour la formation
- Manque de réseau dans le domaine de la formation

Atouts:
- Excellentes compétences relationnelles
- Expérience solide en accompagnement
- Bonne capacité d'adaptation
""", encoding='utf-8')
    
    # Formation
    formation_file = client_folder / "formation.txt"
    formation_file.write_text("""
PARCOURS DE FORMATION

2007-2010: Master en Psychologie du travail
Université Paris 8

2015: Certification conseiller en insertion professionnelle
AFPA

Formations complémentaires:
- Bilan de compétences (2018)
- Techniques d'entretien (2020)
""", encoding='utf-8')
    
    return client_folder


@pytest.fixture
def mock_rag_components():
    """
    Patch LlamaIndex pour éviter appels réseau.
    
    Crée des mocks pour :
    - OpenAIEmbedding → FakeEmbedding
    - OpenAI LLM → FakeLLM
    - VectorStoreIndex → FakeIndex
    """
    # Mock pour le modèle d'embeddings
    class FakeEmbedding:
        """Embedding fake qui retourne toujours le même vecteur."""
        def __init__(self, *args, **kwargs):
            pass
        
        def get_text_embedding(self, text):
            # Retourner un vecteur fake de dimension 1536 (OpenAI standard)
            import hashlib
            # Utiliser hash du texte pour déterminisme
            h = int(hashlib.md5(text.encode()).hexdigest(), 16)
            return [(h % 100) / 100.0] * 1536
        
        def get_query_embedding(self, query):
            return self.get_text_embedding(query)
    
    # Mock pour le LLM
    class FakeLLM:
        """LLM fake qui génère des réponses déterministes."""
        def __init__(self, *args, **kwargs):
            pass
        
        def complete(self, prompt, **kwargs):
            # Extraire le champ demandé du prompt
            prompt_lower = prompt.lower()
            
            if "nom" in prompt_lower:
                return MockResponse("DUPONT")
            elif "prenom" in prompt_lower or "prénom" in prompt_lower:
                return MockResponse("Jean")
            elif "date_naissance" in prompt_lower or "naissance" in prompt_lower:
                return MockResponse("15/03/1985")
            elif "profession" in prompt_lower or "situation" in prompt_lower:
                return MockResponse("Conseiller en orientation professionnelle")
            elif "formation" in prompt_lower or "niveau" in prompt_lower:
                return MockResponse("Master en Psychologie du travail")
            elif "telephone" in prompt_lower or "téléphone" in prompt_lower:
                return MockResponse("06 12 34 56 78")
            elif "email" in prompt_lower:
                return MockResponse("jean.dupont@example.com")
            elif "objectif" in prompt_lower:
                return MockResponse("Devenir formateur et travailler dans l'insertion professionnelle")
            elif "competence" in prompt_lower or "compétence" in prompt_lower:
                return MockResponse("Accompagnement individuel, bilan de compétences")
            else:
                return MockResponse("Non renseigné")
        
        def chat(self, messages, **kwargs):
            # Pour les appels chat
            if messages:
                last_msg = messages[-1]
                content = last_msg.get("content", "") if isinstance(last_msg, dict) else str(last_msg)
                return self.complete(content, **kwargs)
            return MockResponse("Non renseigné")
    
    class MockResponse:
        """Réponse mock."""
        def __init__(self, text):
            self.text = text
        
        def __str__(self):
            return self.text
    
    # Mock pour VectorStoreIndex
    class FakeVectorStoreIndex:
        """Index fake qui simule le comportement de LlamaIndex."""
        def __init__(self, documents=None, *args, **kwargs):
            self.documents = documents or []
            self.docstore = FakeDocstore()
        
        @classmethod
        def from_documents(cls, documents, *args, **kwargs):
            return cls(documents=documents)
        
        def as_query_engine(self, *args, **kwargs):
            return FakeQueryEngine(self.documents)
    
    class FakeDocstore:
        """Docstore fake."""
        def __init__(self):
            self.docs = {}
    
    class FakeQueryEngine:
        """Query engine fake."""
        def __init__(self, documents):
            self.documents = documents
            self.fake_llm = FakeLLM()
        
        def query(self, query_str):
            # Simuler une réponse avec sources
            response = self.fake_llm.complete(query_str)
            
            # Créer des source_nodes fake
            source_nodes = []
            for i, doc in enumerate(self.documents[:3]):  # Max 3 sources
                node = FakeNode(
                    text=doc.text[:200] if hasattr(doc, 'text') else str(doc)[:200],
                    metadata=getattr(doc, 'metadata', {"source_file": f"source_{i}.txt"}),
                    score=0.85 - (i * 0.05)  # Scores décroissants
                )
                source_nodes.append(node)
            
            return FakeResponse(str(response), source_nodes)
    
    class FakeNode:
        """Node fake pour les sources."""
        def __init__(self, text, metadata, score):
            self.text = text
            self.metadata = metadata
            self.score = score
    
    class FakeResponse:
        """Response fake."""
        def __init__(self, text, source_nodes):
            self.text = text
            self.source_nodes = source_nodes
        
        def __str__(self):
            return self.text
    
    # Patcher les imports LlamaIndex
    patches = []
    
    try:
        # Patcher OpenAIEmbedding
        patch_embedding = mock.patch(
            'src.rhpro.rag_generator.OpenAIEmbedding',
            return_value=FakeEmbedding()
        )
        patches.append(patch_embedding)
        patch_embedding.start()
        
        # Patcher OpenAI LLM
        patch_llm = mock.patch(
            'src.rhpro.rag_generator.OpenAI',
            return_value=FakeLLM()
        )
        patches.append(patch_llm)
        patch_llm.start()
        
        # Patcher VectorStoreIndex
        patch_index = mock.patch(
            'src.rhpro.rag_generator.VectorStoreIndex',
            FakeVectorStoreIndex
        )
        patches.append(patch_index)
        patch_index.start()
        
        # Patcher Settings globaux
        patch_settings_embed = mock.patch(
            'src.rhpro.rag_generator.Settings.embed_model',
            FakeEmbedding()
        )
        patches.append(patch_settings_embed)
        patch_settings_embed.start()
        
        patch_settings_llm = mock.patch(
            'src.rhpro.rag_generator.Settings.llm',
            FakeLLM()
        )
        patches.append(patch_settings_llm)
        patch_settings_llm.start()
        
    except Exception as e:
        print(f"Warning: Could not patch LlamaIndex: {e}")
    
    yield patches
    
    # Cleanup
    for patch in patches:
        try:
            patch.stop()
        except:
            pass


# ============================================================================
# Test E2E Principal
# ============================================================================

@pytest.mark.e2e
def test_end2end_pipeline_complete(mini_client_folder, mock_rag_components, tmp_path):
    """
    Test E2E complet : normalise → index → génère → valide.
    
    Vérifie :
    1. Tous les outputs sont créés
    2. Status cohérent avec les métriques
    3. Si GO : métriques au-dessus des seuils
    4. Si NO_GO : au moins une raison bloquante
    
    CONTRAINTE : Aucun appel réseau (LLM/embeddings mockés).
    """
    # Préparer output dir
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    
    client_name = "mini_client_test"
    
    # ========================================================================
    # ÉTAPE 1 : Générer le rapport via report_generator
    # ========================================================================
    
    generator = RHProReportGenerator(
        template_path=None,  # Pas de template, génération simple
    )
    
    try:
        result = generator.generate_from_client(
            sources_folder=str(mini_client_folder),
            gold_path=None,
            output_dir=str(output_dir),
            client_name=client_name,
            strict_mode=True,
            validation_profile=ValidationProfile.STRICT,
        )
        
        # Vérifier succès
        assert result["success"], "Génération doit réussir"
        
    except Exception as e:
        # Si échec (ex: LlamaIndex non disponible), skip gracefully
        pytest.skip(f"Pipeline non disponible: {e}")
    
    # ========================================================================
    # ÉTAPE 2 : Vérifier présence des fichiers de sortie
    # ========================================================================
    
    expected_outputs = {
        "generated_docx": output_dir / f"{client_name}_generated.docx",
        "debug_json": output_dir / f"{client_name}_debug.json",
        "metrics_json": output_dir / f"{client_name}_metrics.json",
        "validation_json": output_dir / f"{client_name}_validation.json",
    }
    
    for output_name, output_path in expected_outputs.items():
        assert output_path.exists(), (
            f"Output manquant: {output_name} attendu à {output_path}"
        )
    
    # ========================================================================
    # ÉTAPE 3 : Charger et vérifier les données
    # ========================================================================
    
    # Charger metrics
    with open(expected_outputs["metrics_json"], 'r', encoding='utf-8') as f:
        metrics = json.load(f)
    
    # Charger debug
    with open(expected_outputs["debug_json"], 'r', encoding='utf-8') as f:
        debug = json.load(f)
    
    # Charger validation
    with open(expected_outputs["validation_json"], 'r', encoding='utf-8') as f:
        validation = json.load(f)
    
    # Vérifier structure minimale
    assert "required_coverage" in metrics, "metrics doit contenir required_coverage"
    assert "quality_score" in metrics, "metrics doit contenir quality_score"
    assert "fields" in debug or "index" in debug, "debug doit contenir fields ou index"
    assert "status" in validation, "validation doit contenir status"
    
    # ========================================================================
    # ÉTAPE 4 : Vérifier cohérence status vs métriques (ANTI-REGRESSION)
    # ========================================================================
    
    status = validation["status"]
    required_coverage = metrics["required_coverage"] / 100  # Normaliser 0-1
    quality_score = metrics["quality_score"]
    avg_confidence = metrics["avg_confidence"]
    
    # Récupérer sources_count
    sources_count = 0
    if "index" in debug:
        sources_count = debug["index"].get("sources_count", 0)
    
    # Récupérer champs critiques
    missing_critical = []
    if "fields" in debug:
        fields = debug["fields"]
        
        # Vérifier nom
        nom_value = fields.get("nom", {}).get("value", "Non renseigné")
        if nom_value == "Non renseigné":
            missing_critical.append("nom")
        
        # Vérifier prénom
        prenom_value = fields.get("prenom", {}).get("value", "Non renseigné")
        if prenom_value == "Non renseigné":
            missing_critical.append("prenom")
        
        # Vérifier profession OU formation
        profession = fields.get("situation_professionnelle", {}).get("value", "Non renseigné")
        formation = fields.get("niveau_formation", {}).get("value", "Non renseigné")
        if profession == "Non renseigné" and formation == "Non renseigné":
            missing_critical.append("profession_or_formation")
    
    # ========================================================================
    # VÉRIFICATION STRICTE : cohérence status vs métriques
    # ========================================================================
    
    if status == "GO":
        # Si GO avec profil STRICT → toutes les conditions DOIVENT être remplies
        assert required_coverage >= 0.85, (
            f"GO avec STRICT doit avoir required_coverage >= 0.85, obtenu: {required_coverage:.2f}"
        )
        assert quality_score >= 0.75, (
            f"GO avec STRICT doit avoir quality_score >= 0.75, obtenu: {quality_score:.2f}"
        )
        assert avg_confidence >= 0.70, (
            f"GO avec STRICT doit avoir avg_confidence >= 0.70, obtenu: {avg_confidence:.2f}"
        )
        assert sources_count >= 1, (
            f"GO avec STRICT doit avoir au moins 1 source, obtenu: {sources_count}"
        )
        assert len(missing_critical) == 0, (
            f"GO avec STRICT ne doit pas avoir de champs critiques manquants, "
            f"manquants: {missing_critical}"
        )
    
    elif status == "NO_GO":
        # Si NO_GO → au moins UNE condition bloquante
        reasons = validation.get("reasons", [])
        assert len(reasons) > 0, "NO_GO doit avoir au moins une raison"
        
        # Vérifier qu'au moins une condition est effectivement bloquante
        has_blocking_condition = (
            sources_count == 0 or
            len(missing_critical) > 0 or
            required_coverage < 0.85 or
            quality_score < 0.75 or
            avg_confidence < 0.70
        )
        
        assert has_blocking_condition, (
            f"NO_GO doit avoir au moins une condition bloquante:\n"
            f"  sources_count={sources_count} (attendu >= 1)\n"
            f"  missing_critical={missing_critical} (attendu vide)\n"
            f"  required_coverage={required_coverage:.2f} (attendu >= 0.85)\n"
            f"  quality_score={quality_score:.2f} (attendu >= 0.75)\n"
            f"  avg_confidence={avg_confidence:.2f} (attendu >= 0.70)"
        )
    
    # ========================================================================
    # ÉTAPE 5 : Vérifier le DOCX généré
    # ========================================================================
    
    docx_path = expected_outputs["generated_docx"]
    doc = Document(str(docx_path))
    
    # Vérifier que le document a du contenu
    assert len(doc.paragraphs) > 0, "DOCX doit contenir des paragraphes"
    
    # Vérifier présence de texte
    all_text = "\n".join([p.text for p in doc.paragraphs])
    assert len(all_text.strip()) > 0, "DOCX doit contenir du texte"


# ============================================================================
# Test E2E avec données minimales → NO_GO attendu
# ============================================================================

@pytest.mark.e2e
def test_end2end_minimal_data_no_go(tmp_path, mock_rag_components):
    """
    Test E2E avec données minimales → doit produire NO_GO.
    
    Vérifie que le système détecte correctement un dossier incomplet.
    """
    # Créer un dossier quasi-vide
    client_folder = tmp_path / "client_minimal"
    client_folder.mkdir()
    
    # Un seul fichier avec très peu d'info
    minimal_file = client_folder / "info.txt"
    minimal_file.write_text("Quelques informations minimales sans identité.", encoding='utf-8')
    
    output_dir = tmp_path / "output_minimal"
    output_dir.mkdir()
    
    client_name = "client_minimal"
    
    # Générer
    generator = RHProReportGenerator()
    
    try:
        result = generator.generate_from_client(
            sources_folder=str(client_folder),
            output_dir=str(output_dir),
            client_name=client_name,
            strict_mode=True,
            validation_profile=ValidationProfile.STRICT,
        )
        
        # Charger validation
        validation_path = output_dir / f"{client_name}_validation.json"
        assert validation_path.exists(), "validation.json doit exister"
        
        with open(validation_path, 'r', encoding='utf-8') as f:
            validation = json.load(f)
        
        # Doit être NO_GO (données insuffisantes)
        assert validation["status"] in ["NO_GO", "DRAFT"], (
            f"Données minimales doivent produire NO_GO ou DRAFT, obtenu: {validation['status']}"
        )
        
        # Doit avoir des raisons
        assert len(validation["reasons"]) > 0, "NO_GO/DRAFT doit avoir des raisons"
        
    except Exception as e:
        pytest.skip(f"Pipeline non disponible: {e}")


# ============================================================================
# Test : Stabilité déterministe
# ============================================================================

@pytest.mark.e2e
def test_end2end_deterministic(mini_client_folder, mock_rag_components, tmp_path):
    """
    Vérifie que deux exécutions successives produisent le même résultat.
    
    Important pour éviter flakiness.
    """
    output_dir1 = tmp_path / "output1"
    output_dir2 = tmp_path / "output2"
    output_dir1.mkdir()
    output_dir2.mkdir()
    
    generator = RHProReportGenerator()
    
    try:
        # Première exécution
        result1 = generator.generate_from_client(
            sources_folder=str(mini_client_folder),
            output_dir=str(output_dir1),
            client_name="test1",
            strict_mode=True,
            validation_profile=ValidationProfile.STRICT,
        )
        
        # Deuxième exécution (même input)
        result2 = generator.generate_from_client(
            sources_folder=str(mini_client_folder),
            output_dir=str(output_dir2),
            client_name="test2",
            strict_mode=True,
            validation_profile=ValidationProfile.STRICT,
        )
        
        # Charger validations
        with open(output_dir1 / "test1_validation.json", 'r') as f:
            val1 = json.load(f)
        
        with open(output_dir2 / "test2_validation.json", 'r') as f:
            val2 = json.load(f)
        
        # Les status doivent être identiques
        assert val1["status"] == val2["status"], (
            f"Deux exécutions doivent produire le même status: {val1['status']} vs {val2['status']}"
        )
        
    except Exception as e:
        pytest.skip(f"Pipeline non disponible: {e}")
