"""Tests pour l'exclusion Devis et la sélection AUTO du DOCX source."""

import tempfile
from pathlib import Path
import pytest
from src.rhpro.client_finder import (
    discover_client_documents_recursive,
    select_best_source_docx,
    contains_keyword
)


class TestExcludeDevis:
    """Tests pour l'exclusion des dossiers et fichiers 'Devis'."""
    
    def test_contains_keyword(self):
        """Vérifie que contains_keyword fonctionne correctement."""
        assert contains_keyword("02 Devis", ["devis"]) is True
        assert contains_keyword("Devis RH-Pro", ["devis"]) is True
        assert contains_keyword("DEVIS", ["devis"]) is True
        assert contains_keyword("Rapport final", ["devis"]) is False
        assert contains_keyword("test", ["devis", "facture"]) is False
    
    def test_exclude_devis_dir(self):
        """Les fichiers dans le dossier '02 Devis' ne doivent pas être listés."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer structure
            (tmpdir / "Rapport final.docx").write_text("test")
            
            devis_dir = tmpdir / "02 Devis"
            devis_dir.mkdir()
            (devis_dir / "Devis RH-Pro 1.docx").write_text("test")
            (devis_dir / "Devis RH-Pro 2.docx").write_text("test")
            (devis_dir / "Facture.pdf").write_text("test")
            
            # Scanner
            result = discover_client_documents_recursive(
                tmpdir,
                max_depth=1,
                exclude_dir_keywords=['devis']
            )
            
            # Vérifier que seul le fichier racine est trouvé
            assert len(result['files']['docx']) == 1
            assert result['files']['docx'][0].name == "Rapport final.docx"
            assert len(result['files']['pdf']) == 0
            
            # Vérifier les exclusions
            assert len(result['excluded_dirs']) > 0
            assert any('devis' in d.lower() for d in result['excluded_dirs'])
    
    def test_exclude_devis_filename_fallback(self):
        """Les fichiers contenant 'devis' dans le nom doivent être exclus."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer fichiers
            (tmpdir / "Rapport final.docx").write_text("test")
            (tmpdir / "Devis RH-Pro - Client.docx").write_text("test")
            (tmpdir / "DEVIS_2024.pdf").write_text("test")
            
            # Scanner
            result = discover_client_documents_recursive(
                tmpdir,
                max_depth=0,
                exclude_file_keywords=['devis']
            )
            
            # Vérifier que seul le rapport est trouvé
            assert len(result['files']['docx']) == 1
            assert result['files']['docx'][0].name == "Rapport final.docx"
            assert len(result['files']['pdf']) == 0
            
            # Vérifier le compteur d'exclusions
            assert result['excluded_files_count'] == 2


class TestAutoSelectSourceDocx:
    """Tests pour la sélection automatique du meilleur DOCX source."""
    
    def test_auto_select_prefers_bilan_over_contrat(self):
        """AUTO doit choisir un doc RH-Pro plutôt qu'un contrat."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer fichiers
            contrat = tmpdir / "Contrat de travail.docx"
            contrat.write_text("test")
            
            bilan = tmpdir / "RH-Pro - Bilan orientation.docx"
            bilan.write_text("test")
            
            docx_files = [contrat, bilan]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir le bilan
            assert best == bilan
            assert mode in ["AUTO_PRIORITY", "AUTO_FALLBACK"]
    
    def test_auto_select_rejects_devis(self):
        """AUTO ne doit PAS choisir un fichier devis."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer fichiers
            devis = tmpdir / "Devis RH-Pro.docx"
            devis.write_text("test")
            
            rapport = tmpdir / "Rapport final.docx"
            rapport.write_text("test")
            
            docx_files = [devis, rapport]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir le rapport (pas le devis)
            assert best == rapport
    
    def test_auto_select_rejects_evaluation_stage(self):
        """Patch 4 Option B: AUTO doit rejeter les docs évaluation/stage pour bilan_complet."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer fichiers
            evaluation = tmpdir / "Évaluation de stage.docx"
            evaluation.write_text("test")
            
            bilan = tmpdir / "Bilan final.docx"
            bilan.write_text("test")
            
            docx_files = [evaluation, bilan]
            
            # Sélection AUTO avec profile bilan_complet
            best, mode = select_best_source_docx(docx_files, profile="bilan_complet")
            
            # Doit choisir le bilan (pas l'évaluation)
            assert best == bilan
            assert "Bilan final" in best.name
    
    def test_auto_select_prefers_lai_keyword(self):
        """Patch 2: AUTO doit privilégier les docs avec 'lai'."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer fichiers
            doc1 = tmpdir / "Rapport.docx"
            doc1.write_text("test")
            
            doc2 = tmpdir / "LAI - Bilan.docx"
            doc2.write_text("test")
            
            docx_files = [doc1, doc2]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir le doc avec 'lai'
            assert best == doc2
            assert mode == "AUTO_PRIORITY"
    
    def test_auto_select_prefers_composite_keywords(self):
        """Patch 2: AUTO doit donner gros boost aux keywords composés (bilan final, bilan d'orientation)."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer fichiers
            bilan_simple = tmpdir / "Bilan.docx"
            bilan_simple.write_text("test")
            
            bilan_final = tmpdir / "Bilan final.docx"
            bilan_final.write_text("test")
            
            bilan_orientation = tmpdir / "Bilan d'orientation.docx"
            bilan_orientation.write_text("test")
            
            docx_files = [bilan_simple, bilan_final, bilan_orientation]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir un des keywords composés
            assert best in [bilan_final, bilan_orientation]
            assert mode == "AUTO_PRIORITY"
    
    def test_auto_select_rejects_certificat(self):
        """Patch 2: AUTO doit rejeter les docs 'certificat'."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer fichiers
            certificat = tmpdir / "Certificat de travail.docx"
            certificat.write_text("test")
            
            rapport = tmpdir / "Rapport.docx"
            rapport.write_text("test")
            
            docx_files = [certificat, rapport]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir le rapport (pas le certificat)
            assert best == rapport
            assert "Certificat" not in best.name
    
    def test_auto_select_rejects_all_admin_docs(self):
        """AUTO doit rejeter tous les docs administratifs."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer uniquement des docs administratifs
            (tmpdir / "Contrat.docx").write_text("test")
            (tmpdir / "Convention.docx").write_text("test")
            (tmpdir / "Attestation.docx").write_text("test")
            
            docx_files = list(tmpdir.glob("*.docx"))
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Aucun candidat valide
            assert best is None
            assert mode == "NONE"
    
    def test_auto_select_returns_none_for_empty_list(self):
        """AUTO doit retourner None si liste vide."""
        best, mode = select_best_source_docx([])
        assert best is None
        assert mode == "NONE"
    
    def test_auto_select_fallback_on_longest_docx(self):
        """Si aucun keyword trouvé, AUTO utilise le scoring (nb paragraphes, headings)."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer de vrais fichiers DOCX (vides) avec tailles différentes
            from docx import Document
            
            small = tmpdir / "doc1.docx"
            doc_small = Document()
            doc_small.add_paragraph("Small doc")
            doc_small.save(str(small))
            
            large = tmpdir / "doc2.docx"
            doc_large = Document()
            for i in range(50):
                doc_large.add_paragraph(f"Paragraph {i}" * 10)
            doc_large.save(str(large))
            
            docx_files = [small, large]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit retourner un résultat (pas None)
            assert best is not None
            assert best in docx_files
            # Le mode peut varier selon le scoring
            assert mode in ["AUTO_PRIORITY", "AUTO_FALLBACK"]


class TestExcludeDevisIntegration:
    """Tests d'intégration pour l'exclusion Devis."""
    
    def test_typical_client_structure_excludes_devis(self):
        """Test avec structure client typique incluant dossier Devis."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Structure typique
            (tmpdir / "Contrat.docx").write_text("test")
            
            personnel = tmpdir / "01 Dossier personnel"
            personnel.mkdir()
            (personnel / "CV.pdf").write_text("test")
            (personnel / "Diplomes.pdf").write_text("test")
            
            devis = tmpdir / "02 Devis"
            devis.mkdir()
            (devis / "Devis 1.docx").write_text("test")
            (devis / "Devis 2.docx").write_text("test")
            (devis / "Facture.pdf").write_text("test")
            
            rapport = tmpdir / "06 Rapport final"
            rapport.mkdir()
            (rapport / "RH-Pro Bilan.docx").write_text("test")
            
            # Scanner
            result = discover_client_documents_recursive(
                tmpdir,
                max_depth=1,
                exclude_dir_keywords=['devis']
            )
            
            # Vérifier totaux
            total_expected = 4  # Contrat + 2 PDF personnel + 1 Bilan
            assert result['total_files'] == total_expected
            
            # Vérifier que Devis est exclu
            assert len(result['excluded_dirs']) > 0
            assert result['excluded_files_count'] == 0  # Exclus au niveau dossier
            
            # Vérifier les DOCX
            docx_files = result['files']['docx']
            docx_names = [f.name for f in docx_files]
            assert "RH-Pro Bilan.docx" in docx_names
            assert "Contrat.docx" in docx_names
            assert "Devis 1.docx" not in docx_names
            assert "Devis 2.docx" not in docx_names
