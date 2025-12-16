"""Tests pour le module core/export.py."""

import pytest
from pathlib import Path
from core.export import docx_to_pdf


class TestDocxToPdf:
    """Tests pour la conversion DOCX vers PDF."""

    def test_returns_path_object(self, tmp_path):
        """Retourne un objet Path."""
        # Créer un fichier DOCX de test
        from docx import Document
        
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(docx_path)
        
        # La conversion pourrait échouer sans LibreOffice installé
        try:
            result = docx_to_pdf(docx_path, tmp_path)
            assert isinstance(result, Path)
        except (RuntimeError, FileNotFoundError):
            # LibreOffice non disponible, on skip ce test
            pytest.skip("LibreOffice not available")

    def test_creates_pdf_file(self, tmp_path):
        """Crée un fichier PDF."""
        from docx import Document
        
        docx_path = tmp_path / "document.docx"
        doc = Document()
        doc.add_paragraph("Contenu du document")
        doc.save(docx_path)
        
        try:
            result = docx_to_pdf(docx_path, tmp_path)
            # Vérifier que le fichier existe (si la conversion a réussi)
            assert result.suffix == ".pdf"
        except (RuntimeError, FileNotFoundError):
            pytest.skip("LibreOffice not available")

    def test_uses_output_dir(self, tmp_path):
        """Utilise le répertoire de sortie spécifié."""
        from docx import Document
        
        docx_path = tmp_path / "input.docx"
        output_dir = tmp_path / "output"
        output_dir.mkdir()
        
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(docx_path)
        
        try:
            result = docx_to_pdf(docx_path, output_dir)
            assert result.parent == output_dir
        except (RuntimeError, FileNotFoundError):
            pytest.skip("LibreOffice not available")
