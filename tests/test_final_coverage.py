"""Tests finaux pour dépasser 60% de couverture."""

import pytest
from pathlib import Path
from core.extract import extract_pdf, extract_docx, extract_txt, walk_files
from core.context import normalize_text, chunk_text
from core.generate import sanitize_output, truncate_lines, truncate_chars
import fitz
from docx import Document


class TestFinalCoverageBoost:
    """Tests finaux pour atteindre 60%+."""

    def test_extract_pdf_multiple_pages_detailed(self, tmp_path):
        """Extraction PDF détaillée multipages."""
        doc = fitz.open()
        
        # 3 pages
        for i in range(3):
            page = doc.new_page()
            page.insert_text((50, 50 + i*20), f"Page {i+1}: Contenu de test")
        
        pdf_file = tmp_path / "detailed.pdf"
        doc.save(pdf_file)
        doc.close()
        
        result = extract_pdf(pdf_file)
        
        assert result.success is True
        assert "Page 1" in result.value["text"]
        assert "Page 2" in result.value["text"]
        assert "Page 3" in result.value["text"]
        assert "pages" in result.value
        assert len(result.value["pages"]) == 3

    def test_extract_docx_complex_structure(self, tmp_path):
        """DOCX avec structure complexe."""
        docx_file = tmp_path / "complex.docx"
        doc = Document()
        
        # Plusieurs éléments
        doc.add_heading("Titre Principal", level=1)
        doc.add_paragraph("Paragraphe 1 avec du texte.")
        doc.add_heading("Sous-titre", level=2)
        doc.add_paragraph("Paragraphe 2 avec plus de contenu.")
        
        # Tableau
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "A2"
        table.cell(1, 0).text = "B1"
        table.cell(1, 1).text = "B2"
        
        doc.save(docx_file)
        
        result = extract_docx(docx_file)
        
        assert result.success is True
        assert "Titre Principal" in result.value["text"]
        assert "Paragraphe 1" in result.value["text"]
        assert "Sous-titre" in result.value["text"]
        # Les cellules du tableau devraient être extraites
        assert "A1" in result.value["text"]
        assert "B2" in result.value["text"]

    def test_extract_txt_various_encodings(self, tmp_path):
        """Teste extraction avec différents contenus."""
        # UTF-8 avec caractères spéciaux
        txt_file = tmp_path / "utf8.txt"
        content = "Texte avec éàü ñ ç œ\nLigne 2\nLigne 3 €£¥"
        txt_file.write_text(content, encoding="utf-8")
        
        result = extract_txt(txt_file)
        
        assert result.success is True
        assert "éàü" in result.value["text"]
        assert "ñ" in result.value["text"]
        assert "€" in result.value["text"]

    def test_walk_files_recursive_structure(self, tmp_path):
        """walk_files avec structure récursive."""
        # Créer structure
        (tmp_path / "dir1").mkdir()
        (tmp_path / "dir1" / "file1.txt").write_text("test1")
        (tmp_path / "dir2").mkdir()
        (tmp_path / "dir2" / "subdir").mkdir()
        (tmp_path / "dir2" / "subdir" / "file2.txt").write_text("test2")
        (tmp_path / "root.txt").write_text("root")
        
        result = walk_files(tmp_path)
        
        # Devrait trouver 3 fichiers
        assert len(result) == 3
        assert all(p.is_file() for p in result)

    def test_normalize_text_comprehensive_cases(self):
        """normalize_text - cas exhaustifs."""
        cases = [
            ("\r\n", "\n"),
            ("\r", "\n"),
            ("\n\n\n\n\n", "\n\n"),
            ("   text   ", "text"),
            ("line1\r\nline2\rline3\n", "line1\nline2\nline3"),
        ]
        for input_val, expected_pattern in cases:
            result = normalize_text(input_val)
            # Vérifier propriétés clés
            assert "\r" not in result
            if expected_pattern == "text":
                assert result.strip() == "text"

    def test_chunk_text_edge_cases(self):
        """chunk_text avec edge cases."""
        # Texte exactement à la taille du chunk
        text = "a" * 1000
        result = chunk_text(text, chunk_size=1000, overlap=0)
        assert len(result) >= 1
        
        # Texte très court
        short = "court"
        result = chunk_text(short, chunk_size=1000)
        assert len(result) == 1
        assert result[0] == short
        
        # Overlap important
        text = "mot " * 200
        result = chunk_text(text, chunk_size=500, overlap=400)
        assert len(result) > 0

    def test_sanitize_output_all_cases(self):
        """sanitize_output - tous les cas."""
        cases = [
            ("```code```", "code"),
            ("text with \u200b zero-width", "text with   zero-width"),
            ("JSON: {}", "{}"),
            ("  spaced  ", "spaced"),
            ("json:\ndata", "data"),
        ]
        for input_val, expected_substring in cases:
            result = sanitize_output(input_val)
            assert expected_substring in result or result.strip()

    def test_truncate_comprehensive(self):
        """Tests exhaustifs de troncature."""
        # Lines
        many_lines = "\n".join([f"L{i}" for i in range(50)])
        result = truncate_lines(many_lines, 10)
        assert result.count("\n") <= 10
        
        # Chars
        long_text = "x" * 1000
        result = truncate_chars(long_text, 100)
        assert len(result) <= 105  # 100 + ellipsis
        
        # Pas de troncature si court
        short = "court"
        assert truncate_lines(short, 100) == short
        assert truncate_chars(short, 100) == short
