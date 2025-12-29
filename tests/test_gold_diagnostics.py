"""
Tests pour PRIORITÉ 5 — Diagnostic GOLD missing

Vérifie que:
- Le diagnostic est produit pour les clients sans GOLD
- Aucun diagnostic n'est produit pour les clients avec GOLD
- Les fichiers JSONL et Markdown sont créés
- Le format de diagnostic est correct
"""
import pytest
import json
from pathlib import Path
from docx import Document as DocxDocument

from src.rhpro.gold_diagnostics import (
    diagnose_gold_missing,
    write_diagnostics_jsonl,
    write_diagnostics_summary,
    extract_text_snippets,
    analyze_candidate_rejection,
)


class TestGoldDiagnostics:
    """Tests du système de diagnostic GOLD missing"""
    
    def test_diagnose_gold_missing_emits_diagnostic(self, tmp_path):
        """
        Test qu'un diagnostic est produit pour un client sans GOLD
        """
        # Créer un dossier client sans GOLD évident
        client_dir = tmp_path / "Client_NoGold"
        client_dir.mkdir()
        
        # Créer un fichier DOCX qui ne match pas les patterns GOLD
        docx_path = client_dir / "notes.docx"
        doc = DocxDocument()
        doc.add_paragraph("Ceci est un fichier sans ancres GOLD")
        doc.save(str(docx_path))
        
        # Exécuter le diagnostic
        diagnostic = diagnose_gold_missing(client_dir, gold_result=None)
        
        # Assertions
        assert diagnostic is not None
        assert diagnostic["client_id"] == "Client_NoGold"
        assert diagnostic["gold_detected"] == False
        assert "candidates" in diagnostic
        assert len(diagnostic["candidates"]) > 0
        
        # Vérifier le candidat
        candidate = diagnostic["candidates"][0]
        assert "path" in candidate
        assert "gold_score" in candidate
        assert "gold_pass" in candidate
        assert "reject_reasons" in candidate
        
        print(f"✅ Diagnostic produit:")
        print(f"   Client: {diagnostic['client_id']}")
        print(f"   Candidats: {len(diagnostic['candidates'])}")
        print(f"   Notes: {diagnostic['notes']}")
    
    def test_diagnose_gold_detected_no_diagnostic(self, tmp_path):
        """
        Test qu'aucun diagnostic détaillé n'est produit si GOLD détecté
        """
        client_dir = tmp_path / "Client_WithGold"
        client_dir.mkdir()
        
        # Simuler un résultat GOLD détecté
        gold_result = {
            "path": str(client_dir / "bilan_final.docx"),
            "score": 0.95,
            "strategy": "06_rapport_final",
        }
        
        # Exécuter le diagnostic
        diagnostic = diagnose_gold_missing(client_dir, gold_result=gold_result)
        
        # Assertions
        assert diagnostic["gold_detected"] == True
        assert "gold_detected_ok" in diagnostic["notes"]
        assert len(diagnostic["candidates"]) == 0  # Pas de scan inutile
        
        print(f"✅ Pas de diagnostic inutile pour GOLD détecté")
    
    def test_extract_text_snippets_from_docx(self, tmp_path):
        """
        Test l'extraction de snippets d'un DOCX
        """
        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("IDENTITÉ")
        doc.add_paragraph("Jean Dupont")
        doc.add_paragraph("FORMATION")
        doc.add_paragraph("Diplôme d'ingénieur")
        doc.save(str(docx_path))
        
        # Extraire snippets
        snippets = extract_text_snippets(docx_path, max_snippets=3, snippet_length=50)
        
        assert len(snippets) > 0
        assert "IDENTITÉ" in snippets[0] or "Jean Dupont" in snippets[0]
        
        print(f"✅ Snippets extraits: {len(snippets)}")
        for i, snippet in enumerate(snippets[:3], 1):
            print(f"   {i}. {snippet[:50]}...")
    
    def test_analyze_candidate_rejection(self, tmp_path):
        """
        Test l'analyse des raisons de rejet d'un candidat
        """
        # Créer un fichier avec score faible
        file_path = tmp_path / "notes_diverses.docx"
        file_path.touch()
        
        # Analyser
        reasons = analyze_candidate_rejection(file_path, score=0.15, threshold=0.5)
        
        assert len(reasons) > 0
        assert any("below_threshold" in r for r in reasons)
        
        print(f"✅ Raisons de rejet identifiées:")
        for reason in reasons:
            print(f"   - {reason}")
    
    def test_write_diagnostics_jsonl(self, tmp_path):
        """
        Test l'écriture du fichier JSONL
        """
        diagnostics = [
            {
                "client_id": "Client1",
                "client_path": "/path/to/client1",
                "gold_detected": False,
                "candidates": [
                    {
                        "path": "notes.docx",
                        "gold_score": 0.12,
                        "gold_pass": False,
                        "reject_reasons": ["no_gold_keywords_found"],
                    }
                ],
                "notes": ["no_high_priority_match"],
            }
        ]
        
        output_path = tmp_path / "diagnostics.jsonl"
        write_diagnostics_jsonl(diagnostics, output_path)
        
        # Vérifier le fichier
        assert output_path.exists()
        
        # Lire et valider
        with open(output_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        
        assert len(lines) == 1
        parsed = json.loads(lines[0])
        assert parsed["client_id"] == "Client1"
        assert parsed["gold_detected"] == False
        
        print(f"✅ JSONL écrit: {output_path}")
    
    def test_write_diagnostics_summary_markdown(self, tmp_path):
        """
        Test l'écriture du résumé Markdown
        """
        diagnostics = [
            {
                "client_id": "Client_Test",
                "client_path": "/path/to/client",
                "gold_detected": False,
                "candidates": [
                    {
                        "path": "rapport/notes.docx",
                        "gold_score": 0.25,
                        "gold_pass": False,
                        "reject_reasons": ["no_high_priority_keywords"],
                        "snippets": ["Extrait 1", "Extrait 2"],
                        "is_ignored": False,
                        "size_bytes": 12345,
                    }
                ],
                "notes": ["1_docx_files_scanned"],
            }
        ]
        
        output_path = tmp_path / "diagnostics.md"
        write_diagnostics_summary(diagnostics, output_path)
        
        # Vérifier le fichier
        assert output_path.exists()
        
        # Lire contenu
        content = output_path.read_text(encoding="utf-8")
        
        assert "# Diagnostic GOLD Missing" in content
        assert "Client_Test" in content
        assert "notes.docx" in content
        
        print(f"✅ Markdown écrit: {output_path}")
        print(f"   Taille: {len(content)} chars")


class TestGoldDiagnosticsIntegration:
    """Tests d'intégration avec le pipeline complet"""
    
    def test_diagnostic_structure_complete(self, tmp_path):
        """
        Test que le diagnostic contient tous les champs requis
        """
        # Créer un client minimal
        client_dir = tmp_path / "Client_Integration"
        client_dir.mkdir()
        
        # Créer plusieurs fichiers DOCX avec différents scores
        for filename in ["bilan.docx", "notes.docx", "cv.docx"]:
            docx_path = client_dir / filename
            doc = DocxDocument()
            doc.add_paragraph(f"Contenu de {filename}")
            doc.save(str(docx_path))
        
        # Exécuter diagnostic
        diagnostic = diagnose_gold_missing(client_dir, gold_result=None)
        
        # Vérifier structure complète
        required_fields = ["client_id", "client_path", "gold_detected", "timestamp", "candidates", "notes"]
        for field in required_fields:
            assert field in diagnostic, f"Champ manquant: {field}"
        
        # Vérifier structure des candidats
        assert len(diagnostic["candidates"]) == 3
        
        for candidate in diagnostic["candidates"]:
            required_candidate_fields = [
                "path", "absolute_path", "type", "size_bytes",
                "is_ignored", "gold_score", "gold_pass", "reject_reasons"
            ]
            for field in required_candidate_fields:
                assert field in candidate, f"Champ candidat manquant: {field}"
        
        # Vérifier tri par score
        scores = [c["gold_score"] for c in diagnostic["candidates"]]
        assert scores == sorted(scores, reverse=True), "Candidats non triés par score"
        
        print(f"✅ Structure complète validée:")
        print(f"   Candidats: {len(diagnostic['candidates'])}")
        print(f"   Scores: {scores}")
    
    def test_diagnostic_with_high_score_file(self, tmp_path):
        """
        Test diagnostic avec un fichier qui a un bon score mais non sélectionné
        """
        client_dir = tmp_path / "Client_HighScore"
        client_dir.mkdir()
        
        # Créer un fichier avec keywords GOLD mais pas assez
        docx_path = client_dir / "rapport_intermediaire.docx"
        doc = DocxDocument()
        doc.add_paragraph("BILAN INTERMÉDIAIRE")
        doc.add_paragraph("Synthèse partielle")
        doc.save(str(docx_path))
        
        # Diagnostic
        diagnostic = diagnose_gold_missing(client_dir, gold_result=None)
        
        # Le fichier devrait avoir un score > 0
        best_candidate = diagnostic["candidates"][0]
        assert best_candidate["gold_score"] > 0.0
        
        # Mais pas assez pour être "gold_pass"
        if not best_candidate["gold_pass"]:
            assert "below_threshold" in " ".join(best_candidate["reject_reasons"])
        
        print(f"✅ Diagnostic haute granularité:")
        print(f"   Meilleur score: {best_candidate['gold_score']}")
        print(f"   Pass: {best_candidate['gold_pass']}")
        print(f"   Raisons: {best_candidate['reject_reasons']}")
