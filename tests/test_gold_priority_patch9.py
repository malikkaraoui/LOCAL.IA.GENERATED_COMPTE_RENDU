"""
Tests PATCH 9: Priorité GOLD pour bilan_complet

Objectif: "Journal" et "Evaluation de stage" ne doivent PAS être choisis comme GOLD.
Prioriser: "Bilan final", "Rapport final", "Bilan général", "Synthèse", "Bilan d'orientation"
puis "Rapport RH-Pro"
et seulement ensuite journal/évaluations/tests
"""
import pytest
from pathlib import Path
import tempfile
from docx import Document

from src.rhpro.client_scanner import score_gold_candidate
from src.rhpro.client_finder import select_best_source_docx


class TestPatch9GoldPriority:
    """Tests pour la priorité GOLD (PATCH 9)"""
    
    def test_exclude_journal_from_gold(self):
        """Journal ne doit JAMAIS être GOLD"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            journal = tmpdir / "Journal de bilan.docx"
            journal.touch()
            
            score = score_gold_candidate(journal)
            
            # Score doit être 0 (exclu)
            assert score == 0.0
    
    def test_exclude_evaluation_stage_from_gold(self):
        """Evaluation de stage ne doit JAMAIS être GOLD"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            eval_stage = tmpdir / "Évaluation de stage.docx"
            eval_stage.touch()
            
            score = score_gold_candidate(eval_stage)
            
            # Score doit être 0 (exclu)
            assert score == 0.0
    
    def test_bilan_final_high_priority(self):
        """Bilan final doit avoir score élevé"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            bilan_final = tmpdir / "Bilan final.docx"
            bilan_final.touch()
            
            score = score_gold_candidate(bilan_final)
            
            # Score doit être > 0.5 (high priority keyword)
            assert score > 0.5
    
    def test_rapport_final_high_priority(self):
        """Rapport final doit avoir score élevé"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            rapport_final = tmpdir / "Rapport final.docx"
            rapport_final.touch()
            
            score = score_gold_candidate(rapport_final)
            
            # Score doit être > 0.5
            assert score > 0.5
    
    def test_bilan_orientation_high_priority(self):
        """Bilan d'orientation doit avoir score élevé"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            bilan_ori = tmpdir / "Bilan d'orientation.docx"
            bilan_ori.touch()
            
            score = score_gold_candidate(bilan_ori)
            
            # Score doit être > 0.5
            assert score > 0.5
    
    def test_rapport_rhpro_medium_priority(self):
        """Rapport RH-Pro doit avoir score moyen"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            rapport_rhpro = tmpdir / "Rapport RH-Pro.docx"
            rapport_rhpro.touch()
            
            score = score_gold_candidate(rapport_rhpro)
            
            # Score doit être > 0.3 mais < bilan final
            assert score > 0.3
    
    def test_priority_order_bilan_final_vs_journal(self):
        """Bilan final doit être préféré à Journal"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            journal = tmpdir / "Journal de bilan.docx"
            doc_journal = Document()
            doc_journal.add_paragraph("Journal")
            doc_journal.save(str(journal))
            
            bilan_final = tmpdir / "Bilan final.docx"
            doc_bilan = Document()
            doc_bilan.add_paragraph("Bilan")
            doc_bilan.save(str(bilan_final))
            
            docx_files = [journal, bilan_final]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir le bilan final (pas journal)
            assert best == bilan_final
            assert "Journal" not in best.name
    
    def test_priority_order_bilan_final_vs_evaluation(self):
        """Bilan final doit être préféré à Evaluation de stage"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            evaluation = tmpdir / "Évaluation de stage.docx"
            doc_eval = Document()
            doc_eval.add_paragraph("Evaluation")
            doc_eval.save(str(evaluation))
            
            bilan_final = tmpdir / "Bilan final.docx"
            doc_bilan = Document()
            doc_bilan.add_paragraph("Bilan")
            doc_bilan.save(str(bilan_final))
            
            docx_files = [evaluation, bilan_final]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir le bilan final
            assert best == bilan_final
    
    def test_priority_order_rapport_final_vs_journal(self):
        """Rapport final > Journal"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            journal = tmpdir / "Journal.docx"
            doc_journal = Document()
            doc_journal.add_paragraph("Journal")
            doc_journal.save(str(journal))
            
            rapport = tmpdir / "Rapport final.docx"
            doc_rapport = Document()
            doc_rapport.add_paragraph("Rapport")
            doc_rapport.save(str(rapport))
            
            docx_files = [journal, rapport]
            
            best, mode = select_best_source_docx(docx_files)
            
            assert best == rapport
    
    def test_priority_order_bilan_general_vs_journal(self):
        """Bilan général > Journal"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            journal = tmpdir / "Journal de bord.docx"
            doc_journal = Document()
            doc_journal.add_paragraph("Journal")
            doc_journal.save(str(journal))
            
            bilan_gen = tmpdir / "Bilan général.docx"
            doc_bilan = Document()
            doc_bilan.add_paragraph("Bilan")
            doc_bilan.save(str(bilan_gen))
            
            docx_files = [journal, bilan_gen]
            
            best, mode = select_best_source_docx(docx_files)
            
            assert best == bilan_gen
    
    def test_priority_order_rapport_rhpro_vs_journal(self):
        """Rapport RH-Pro > Journal"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            journal = tmpdir / "Journal.docx"
            doc_journal = Document()
            doc_journal.add_paragraph("Journal")
            doc_journal.save(str(journal))
            
            rapport_rhpro = tmpdir / "Rapport RH-Pro.docx"
            doc_rapport = Document()
            doc_rapport.add_paragraph("Rapport")
            doc_rapport.save(str(rapport_rhpro))
            
            docx_files = [journal, rapport_rhpro]
            
            best, mode = select_best_source_docx(docx_files)
            
            assert best == rapport_rhpro
    
    def test_only_journal_returns_none(self):
        """Si seulement Journal disponible, retourner None"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            journal = tmpdir / "Journal.docx"
            doc_journal = Document()
            doc_journal.add_paragraph("Journal")
            doc_journal.save(str(journal))
            
            docx_files = [journal]
            
            best, mode = select_best_source_docx(docx_files)
            
            # Aucun candidat valide
            assert best is None
            assert mode == "NONE"
    
    def test_only_evaluation_returns_none(self):
        """Si seulement Evaluation de stage, retourner None"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            evaluation = tmpdir / "Évaluation de stage.docx"
            doc_eval = Document()
            doc_eval.add_paragraph("Evaluation")
            doc_eval.save(str(evaluation))
            
            docx_files = [evaluation]
            
            best, mode = select_best_source_docx(docx_files)
            
            assert best is None
            assert mode == "NONE"
    
    def test_complete_priority_cascade(self):
        """Test cascade complète de priorité"""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            
            # Créer tous les types
            journal = tmpdir / "Journal.docx"
            evaluation = tmpdir / "Évaluation de stage.docx"
            rapport_simple = tmpdir / "Rapport.docx"
            rapport_rhpro = tmpdir / "Rapport RH-Pro.docx"
            bilan_final = tmpdir / "Bilan final.docx"
            
            for doc_path in [journal, evaluation, rapport_simple, rapport_rhpro, bilan_final]:
                doc = Document()
                doc.add_paragraph(doc_path.stem)
                doc.save(str(doc_path))
            
            docx_files = [journal, evaluation, rapport_simple, rapport_rhpro, bilan_final]
            
            # Sélection AUTO
            best, mode = select_best_source_docx(docx_files)
            
            # Doit choisir Bilan final (priorité max)
            assert best == bilan_final
            assert mode == "AUTO_PRIORITY"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
