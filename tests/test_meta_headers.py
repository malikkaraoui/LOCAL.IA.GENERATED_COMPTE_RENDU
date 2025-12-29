"""
Tests pour PRIORITÉ 3 — Meta Headers (titres administratifs ignorés)

Vérifie que "PARTICIPATION AU PROGRAMME" et variantes:
- Ne comptent PAS dans unknown_titles
- Ne créent PAS de section structurée
- Sont ignorés avec normalisation identique à section_title_map
"""
import pytest
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt

from src.rhpro.parse_bilan import parse_bilan_docx_to_normalized
from src.rhpro.dataset_training import META_HEADERS_NORM, _normalize_title_for_meta


# Helper pour parser avec ruleset par défaut
def parse_bilan(docx_path):
    """Helper pour parser un DOCX avec le ruleset par défaut"""
    ruleset_path = Path(__file__).parent.parent / "config" / "rulesets" / "rhpro_v1.yaml"
    return parse_bilan_docx_to_normalized(str(docx_path), str(ruleset_path))


class TestMetaHeadersNormalization:
    """Tests de normalisation des meta headers"""
    
    def test_meta_headers_norm_contains_participation(self):
        """Vérifie que PARTICIPATION AU PROGRAMME est dans META_HEADERS_NORM"""
        normalized = _normalize_title_for_meta("PARTICIPATION AU PROGRAMME")
        assert normalized in META_HEADERS_NORM, \
            f"PARTICIPATION AU PROGRAMME normalisé devrait être dans META_HEADERS_NORM: {normalized}"
    
    def test_normalize_removes_accents(self):
        """Vérifie que la normalisation retire les accents"""
        # Test avec accents
        with_accent = _normalize_title_for_meta("PARTICIPÁTION AU PRÓGRÁMMÉ")
        without_accent = _normalize_title_for_meta("PARTICIPATION AU PROGRAMME")
        assert with_accent == without_accent, "La normalisation devrait retirer les accents"
    
    def test_normalize_strips_punctuation(self):
        """Vérifie que la normalisation retire la ponctuation finale"""
        variants = [
            "PARTICIPATION AU PROGRAMME",
            "PARTICIPATION AU PROGRAMME:",
            "PARTICIPATION AU PROGRAMME…",
            "PARTICIPATION AU PROGRAMME :  ",
        ]
        normalized_set = {_normalize_title_for_meta(v) for v in variants}
        assert len(normalized_set) == 1, \
            f"Toutes les variantes devraient être normalisées identiquement: {normalized_set}"


class TestMetaHeaderParticipationIgnored:
    """Tests d'intégration: PARTICIPATION AU PROGRAMME ne doit pas être dans unknown_titles"""
    
    def test_meta_header_participation_ignored_not_unknown(self, tmp_path):
        """
        Test principal: PARTICIPATION AU PROGRAMME et variantes sont ignorées
        et ne remontent PAS en unknown_titles
        """
        # Créer un DOCX avec plusieurs variantes du meta header
        docx_path = tmp_path / "test_participation.docx"
        doc = DocxDocument()
        
        # Titre identity (section connue)
        para_identity = doc.add_paragraph("IDENTITÉ")
        para_identity.runs[0].bold = True
        para_identity.runs[0].font.size = Pt(14)
        doc.add_paragraph("Jean Dupont")
        
        # Meta header - variante 1 (exacte)
        para_meta1 = doc.add_paragraph("PARTICIPATION AU PROGRAMME")
        para_meta1.runs[0].bold = True
        para_meta1.runs[0].font.size = Pt(14)
        doc.add_paragraph("Contenu administratif à ignorer")
        
        # Meta header - variante 2 (avec ponctuation)
        para_meta2 = doc.add_paragraph("PARTICIPATION AU PROGRAMME:")
        para_meta2.runs[0].bold = True
        para_meta2.runs[0].font.size = Pt(14)
        doc.add_paragraph("Autre contenu administratif")
        
        # Meta header - variante 3 (avec ellipse)
        para_meta3 = doc.add_paragraph("PARTICIPATION AU PROGRAMME…")
        para_meta3.runs[0].bold = True
        para_meta3.runs[0].font.size = Pt(14)
        doc.add_paragraph("Encore du contenu")
        
        # Titre inconnu légitime (devrait rester unknown)
        para_unknown = doc.add_paragraph("TITRE VRAIMENT INCONNU")
        para_unknown.runs[0].bold = True
        para_unknown.runs[0].font.size = Pt(14)
        doc.add_paragraph("Contenu légitime")
        
        doc.save(str(docx_path))
        
        # Parser le document
        result = parse_bilan(docx_path)
        report = result['report']
        unknown_titles = report.get('unknown_titles', [])
        
        # ✅ ASSERTIONS PRINCIPALES
        # 1. Les variantes de PARTICIPATION AU PROGRAMME ne doivent PAS être dans unknown_titles
        unknown_str = ' '.join(unknown_titles).upper()
        assert 'PARTICIPATION' not in unknown_str, \
            f"PARTICIPATION AU PROGRAMME trouvé dans unknown_titles: {unknown_titles}"
        
        # 2. Le titre vraiment inconnu DOIT être dans unknown_titles
        assert any('INCONNU' in title.upper() for title in unknown_titles), \
            f"Le titre inconnu légitime devrait être dans unknown_titles: {unknown_titles}"
        
        print(f"✅ Test réussi:")
        print(f"   Unknown titles: {unknown_titles}")
        print(f"   PARTICIPATION AU PROGRAMME correctement ignoré (non présent)")
    
    def test_meta_header_with_extra_spaces(self, tmp_path):
        """Test avec espaces supplémentaires dans le titre"""
        docx_path = tmp_path / "test_participation_spaces.docx"
        doc = DocxDocument()
        
        # Meta header avec espaces
        para_meta = doc.add_paragraph("  PARTICIPATION   AU   PROGRAMME  ")
        para_meta.runs[0].bold = True
        para_meta.runs[0].font.size = Pt(14)
        doc.add_paragraph("Contenu")
        
        doc.save(str(docx_path))
        
        result = parse_bilan(docx_path)
        unknown_titles = result['report'].get('unknown_titles', [])
        
        # Ne devrait pas être dans unknown_titles
        unknown_str = ' '.join(unknown_titles).upper()
        assert 'PARTICIPATION' not in unknown_str, \
            f"Meta header avec espaces trouvé dans unknown_titles: {unknown_titles}"
    
    def test_meta_header_does_not_create_section(self, tmp_path):
        """Vérifie que le meta header ne crée pas de contenu dans participation_programme"""
        docx_path = tmp_path / "test_participation_no_section.docx"
        doc = DocxDocument()
        
        # Titre identity
        para_identity = doc.add_paragraph("IDENTITÉ")
        para_identity.runs[0].bold = True
        para_identity.runs[0].font.size = Pt(14)
        doc.add_paragraph("Jean Dupont")
        
        # Meta header
        para_meta = doc.add_paragraph("PARTICIPATION AU PROGRAMME")
        para_meta.runs[0].bold = True
        para_meta.runs[0].font.size = Pt(14)
        doc.add_paragraph("Contenu administratif")
        
        doc.save(str(docx_path))
        
        result = parse_bilan(docx_path)
        normalized = result['normalized']
        
        # La clé participation_programme existe dans le template, mais doit rester VIDE
        # (car le segment a été filtré dans segmenter)
        if 'participation_programme' in normalized:
            # Si la clé existe, elle doit être vide (string vide ou dict vide)
            content = normalized['participation_programme']
            assert content == '' or content == {}, \
                f"participation_programme devrait être vide, got: {content}"
        
        # Vérifier que identity existe bien et n'est pas vide
        assert 'identity' in normalized, \
            "La section identity devrait être présente"


class TestMetaHeadersIntegrationWithMapping:
    """Tests d'intégration avec le mapper"""
    
    def test_meta_header_confidence_zero(self, tmp_path):
        """
        Si le meta header était mappé (ce qui ne devrait pas arriver),
        il ne devrait avoir aucune confidence
        """
        # Ce test est plus défensif - vérifie qu'aucun segment avec
        # PARTICIPATION AU PROGRAMME n'a de mapped_section_id
        docx_path = tmp_path / "test_participation_mapping.docx"
        doc = DocxDocument()
        
        para_meta = doc.add_paragraph("PARTICIPATION AU PROGRAMME")
        para_meta.runs[0].bold = True
        para_meta.runs[0].font.size = Pt(14)
        doc.add_paragraph("Contenu")
        
        doc.save(str(docx_path))
        
        result = parse_bilan(docx_path)
        
        # Vérifier dans le report qu'aucune section "participation" n'est trouvée
        found_sections = result['report'].get('found_sections', [])
        participation_sections = [
            s for s in found_sections
            if 'PARTICIPATION' in s.get('title', '').upper()
        ]
        
        assert len(participation_sections) == 0, \
            f"PARTICIPATION AU PROGRAMME ne devrait pas être dans found_sections: {participation_sections}"
