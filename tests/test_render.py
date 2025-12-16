"""Tests pour le module core/render.py."""

import pytest
from pathlib import Path
from core.render import (
    _norm,
    build_moustache_mapping,
    _stringify_answer,
    replace_text_everywhere
)
from docx import Document


class TestNormFunction:
    """Tests pour la fonction de normalisation de texte."""

    def test_strips_whitespace(self):
        """Supprime les espaces."""
        assert _norm("  texte  ") == "texte"

    def test_removes_colons(self):
        """Supprime les deux-points."""
        assert _norm("titre:") == "titre"

    def test_lowercases_text(self):
        """Convertit en minuscules."""
        assert _norm("TEXTE") == "texte"

    def test_normalizes_spaces(self):
        """Normalise les espaces multiples."""
        assert _norm("a  b   c") == "a b c"

    def test_replaces_nbsp(self):
        """Remplace les espaces insécables."""
        assert _norm("a\u00a0b") == "a b"


class TestStringifyAnswer:
    """Tests pour la conversion de réponses en chaînes."""

    def test_converts_string(self):
        """Convertit une chaîne."""
        assert _stringify_answer("test") == "test"

    def test_converts_integer(self):
        """Convertit un entier."""
        assert _stringify_answer(42) == "42"

    def test_converts_float(self):
        """Convertit un flottant."""
        result = _stringify_answer(3.14)
        assert "3.14" in result

    def test_converts_list_to_bullets(self):
        """Convertit une liste en points."""
        result = _stringify_answer(["item1", "item2"])
        assert "item1" in result
        assert "item2" in result

    def test_handles_none(self):
        """Gère None."""
        assert _stringify_answer(None) == ""


class TestBuildMoustacheMapping:
    """Tests pour la construction du mapping de placeholders."""

    def test_creates_mapping_from_dict(self):
        """Crée un mapping depuis un dictionnaire."""
        answers = {"NOM": "DUPONT", "AGE": 25}
        result = build_moustache_mapping(answers)
        
        assert "{{NOM}}" in result
        assert result["{{NOM}}"] == "DUPONT"
        assert "{{AGE}}" in result

    def test_handles_empty_dict(self):
        """Gère un dictionnaire vide."""
        result = build_moustache_mapping({})
        assert result == {}

    def test_converts_values_to_strings(self):
        """Convertit les valeurs en chaînes."""
        answers = {"COUNT": 42, "ACTIVE": True}
        result = build_moustache_mapping(answers)
        
        assert isinstance(result["{{COUNT}}"], str)


class TestReplaceTextEverywhere:
    """Tests pour le remplacement de texte dans un document."""

    def test_replaces_in_paragraphs(self, tmp_path):
        """Remplace dans les paragraphes."""
        doc_path = tmp_path / "test.docx"
        
        doc = Document()
        doc.add_paragraph("Bonjour {{NOM}}")
        doc.save(doc_path)
        
        doc = Document(doc_path)
        replace_text_everywhere(doc, {"{{NOM}}": "MARTIN"})
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "MARTIN" in text
        assert "{{NOM}}" not in text

    def test_replaces_multiple_placeholders(self, tmp_path):
        """Remplace plusieurs placeholders."""
        doc_path = tmp_path / "test.docx"
        
        doc = Document()
        doc.add_paragraph("{{NOM}} {{PRÉNOM}}")
        doc.save(doc_path)
        
        doc = Document(doc_path)
        replace_text_everywhere(doc, {
            "{{NOM}}": "DUPONT",
            "{{PRÉNOM}}": "Marie"
        })
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "DUPONT" in text
        assert "Marie" in text
