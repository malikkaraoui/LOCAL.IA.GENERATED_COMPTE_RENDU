"""Tests pour augmenter la couverture de core/render.py."""

import pytest
from pathlib import Path
from core.render import (
    _norm,
    _stringify_answer,
    build_moustache_mapping,
    replace_text_everywhere
)
from docx import Document


class TestRenderHelpers:
    """Tests pour les fonctions helper de rendu."""

    def test_norm_comprehensive(self):
        """Test exhaustif de _norm."""
        cases = [
            ("UPPER", "upper"),
            ("  spaces  ", "spaces"),
            ("title:", "title"),
            ("a\u00a0b", "a b"),  # nbsp
            ("multiple   spaces", "multiple spaces"),
        ]
        for input_val, expected in cases:
            result = _norm(input_val)
            assert result == expected

    def test_stringify_various_types(self):
        """Stringify différents types."""
        assert _stringify_answer("text") == "text"
        assert _stringify_answer(42).isdigit()
        assert _stringify_answer(3.14)
        assert _stringify_answer(None) == ""
        assert _stringify_answer(True)

    def test_stringify_list_items(self):
        """Stringify des listes."""
        result = _stringify_answer(["a", "b", "c"])
        assert "a" in result
        assert "b" in result
        assert "c" in result

    def test_stringify_nested_structures(self):
        """Stringify structures imbriquées."""
        data = {"key": ["val1", "val2"]}
        result = _stringify_answer(data)
        # Devrait convertir en texte
        assert isinstance(result, str)

    def test_moustache_mapping_comprehensive(self):
        """Test exhaustif du mapping."""
        data = {
            "NOM": "DUPONT",
            "AGE": 30,
            "ACTIF": True,
            "NOTES": ["note1", "note2"]
        }
        result = build_moustache_mapping(data)
        
        assert "{{NOM}}" in result
        assert "{{AGE}}" in result
        assert "{{ACTIF}}" in result
        assert "{{NOTES}}" in result
        
        # Toutes les valeurs doivent être des chaînes
        for val in result.values():
            assert isinstance(val, str)

    def test_replace_text_in_multiple_paragraphs(self, tmp_path):
        """Remplace dans plusieurs paragraphes."""
        doc_path = tmp_path / "multi.docx"
        
        doc = Document()
        doc.add_paragraph("Premier {{NOM}}")
        doc.add_paragraph("Deuxième {{PRÉNOM}}")
        doc.add_paragraph("Troisième {{NOM}} encore")
        doc.save(doc_path)
        
        doc = Document(doc_path)
        replace_text_everywhere(doc, {
            "{{NOM}}": "MARTIN",
            "{{PRÉNOM}}": "Sophie"
        })
        
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "MARTIN" in all_text
        assert "Sophie" in all_text
        assert "{{NOM}}" not in all_text
        assert "{{PRÉNOM}}" not in all_text

    def test_replace_text_handles_empty_mapping(self, tmp_path):
        """Gère un mapping vide."""
        doc_path = tmp_path / "empty_map.docx"
        
        doc = Document()
        doc.add_paragraph("Texte {{PLACEHOLDER}}")
        doc.save(doc_path)
        
        doc = Document(doc_path)
        replace_text_everywhere(doc, {})
        
        # Le texte ne devrait pas changer
        text = doc.paragraphs[0].text
        assert "{{PLACEHOLDER}}" in text
