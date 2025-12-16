"""Tests pour le module core/template_fields.py."""

import pytest
from pathlib import Path
from core.template_fields import extract_placeholders_from_docx
from docx import Document


class TestExtractPlaceholders:
    """Tests pour l'extraction des placeholders."""

    def test_extracts_simple_placeholder(self, tmp_path):
        """Extrait un placeholder simple."""
        template_path = tmp_path / "template.docx"
        
        doc = Document()
        doc.add_paragraph("Bonjour {{NOM}}")
        doc.save(template_path)
        
        result = extract_placeholders_from_docx(template_path)
        assert "NOM" in result

    def test_extracts_multiple_placeholders(self, tmp_path):
        """Extrait plusieurs placeholders."""
        template_path = tmp_path / "template.docx"
        
        doc = Document()
        doc.add_paragraph("{{NOM}} {{PRÉNOM}} habite à {{VILLE}}")
        doc.save(template_path)
        
        result = extract_placeholders_from_docx(template_path)
        assert "NOM" in result
        assert "PRÉNOM" in result
        assert "VILLE" in result

    def test_removes_duplicates(self, tmp_path):
        """Supprime les doublons."""
        template_path = tmp_path / "template.docx"
        
        doc = Document()
        doc.add_paragraph("{{NOM}} est {{NOM}}")
        doc.save(template_path)
        
        result = extract_placeholders_from_docx(template_path)
        assert result.count("NOM") == 1

    def test_handles_empty_document(self, tmp_path):
        """Gère un document vide."""
        template_path = tmp_path / "template.docx"
        
        doc = Document()
        doc.save(template_path)
        
        result = extract_placeholders_from_docx(template_path)
        assert result == []

    def test_extracts_from_tables(self, tmp_path):
        """Extrait depuis les tableaux."""
        template_path = tmp_path / "template.docx"
        
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "{{TABLEAU}}"
        doc.save(template_path)
        
        result = extract_placeholders_from_docx(template_path)
        assert "TABLEAU" in result
