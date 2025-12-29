"""
Test d'intégration end-to-end pour vérifier l'impact des règles regex
sur la réduction des unknown_titles.
"""
import pytest
from pathlib import Path
from src.rhpro.parse_bilan import parse_bilan_docx_to_normalized
from src.rhpro.ruleset_loader import RulesetLoader


def test_integration_title_rules_reduce_unknown_titles(tmp_path):
    """
    Test d'intégration : les règles regex doivent réduire unknown_titles
    pour les titres de tests qui n'étaient pas dans le mapping exact.
    """
    # Créer un document de test avec des titres qui doivent matcher les règles
    from docx import Document
    
    doc = Document()
    doc.add_heading("IDENTITÉ", level=1)
    doc.add_paragraph("Sophie Martin — 756.1234.5678.90")
    
    # Titres qui doivent maintenant matcher via title_rules
    doc.add_heading("FRANÇAIS - TEST DE POSITIONNEMENT NIVEAU", level=2)
    doc.add_paragraph("Résultat: Niveau B1")
    
    doc.add_heading("CALCUL - NIVEAU 3", level=2)
    doc.add_paragraph("Score: 85%")
    
    doc.add_heading("ANGLAIS POSITIONNEMENT", level=2)
    doc.add_paragraph("Résultat: A2")
    
    doc.add_heading("DIMENSIONS - MESURES", level=2)
    doc.add_paragraph("Test passé avec succès")
    
    # Titre qui restera unknown (non mappé)
    doc.add_heading("TITRE INCONNU LEGITIME", level=2)
    doc.add_paragraph("Contenu quelconque")
    
    # Sauvegarder le document
    docx_path = tmp_path / "test_title_rules.docx"
    doc.save(str(docx_path))
    
    # Charger le ruleset
    ruleset_path = Path(__file__).parent.parent / "config" / "rulesets" / "rhpro_v1.yaml"
    
    # Parser le document
    result = parse_bilan_docx_to_normalized(
        str(docx_path),
        str(ruleset_path)
    )
    
    report = result['report']
    unknown_titles = report.get('unknown_titles', [])
    found_sections = report.get('found_sections', [])
    
    # Vérifications
    # 1. Les titres de tests ne doivent PAS être dans unknown_titles
    unknown_titles_str = ' '.join(unknown_titles)
    assert 'FRANCAIS' not in unknown_titles_str.upper(), \
        f"FRANÇAIS test trouvé dans unknown_titles: {unknown_titles}"
    assert 'CALCUL' not in unknown_titles_str.upper(), \
        f"CALCUL test trouvé dans unknown_titles: {unknown_titles}"
    assert 'ANGLAIS' not in unknown_titles_str.upper(), \
        f"ANGLAIS test trouvé dans unknown_titles: {unknown_titles}"
    assert 'DIMENSIONS' not in unknown_titles_str.upper(), \
        f"DIMENSIONS test trouvé dans unknown_titles: {unknown_titles}"
    
    # 2. Les titres de tests doivent être dans found_sections sous 'tests'
    found_section_ids = [s.get('section_id') for s in found_sections]
    assert 'tests' in found_section_ids, \
        f"Section 'tests' non trouvée. Sections trouvées: {found_section_ids}"
    
    # 3. unknown_titles doit être réduit (seulement le titre légitime inconnu)
    assert len(unknown_titles) <= 1, \
        f"Trop d'unknown_titles: {unknown_titles}. Attendu: ≤1"
    
    # 4. Le titre légitime inconnu doit être présent
    if len(unknown_titles) > 0:
        assert any('TITRE INCONNU' in t.upper() for t in unknown_titles), \
            f"Titre légitime inconnu manquant: {unknown_titles}"
    
    print(f"✅ Test réussi:")
    print(f"   Found sections: {found_section_ids}")
    print(f"   Unknown titles: {unknown_titles}")
    print(f"   Unknown count: {len(unknown_titles)}")


def test_title_rules_method_shown_in_report():
    """
    Vérifie que la méthode 'title_rule' est enregistrée quand une règle match.
    """
    from src.rhpro.mapper import TitleMapper
    from src.rhpro.segmenter import Segment
    from src.rhpro.ruleset_loader import RulesetLoader
    
    # Charger le ruleset
    ruleset_path = Path(__file__).parent.parent / "config" / "rulesets" / "rhpro_v1.yaml"
    ruleset = RulesetLoader(str(ruleset_path))
    
    # Créer un mapper
    mapper = TitleMapper(ruleset)
    
    # Créer un segment avec un titre qui doit matcher via title_rule
    segment = Segment(
        raw_title="FRANÇAIS - POSITIONNEMENT NIVEAU",
        normalized_title="FRANÇAIS - POSITIONNEMENT NIVEAU",
        level=2,
        paragraphs=[]
    )
    
    # Mapper le segment
    segments = mapper.map_segments([segment])
    
    # Vérifications
    assert len(segments) == 1
    mapped_segment = segments[0]
    
    # Le segment doit être mappé vers 'tests'
    assert mapped_segment.mapped_section_id == 'tests', \
        f"Section attendue: 'tests', obtenue: {mapped_segment.mapped_section_id}"
    
    # La confidence doit être 0.80 (fallback regex)
    assert mapped_segment.confidence == 0.80, \
        f"Confidence attendue: 0.80, obtenue: {mapped_segment.confidence}"
    
    print(f"✅ Titre mappé via title_rule:")
    print(f"   Section: {mapped_segment.mapped_section_id}")
    print(f"   Confidence: {mapped_segment.confidence}")


if __name__ == '__main__':
    pytest.main([__file__, '-v', '-s'])
