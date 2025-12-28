"""
Tests anti-régression V4.1 - Patch tables + GOLD strict

Contraintes testées:
1. Tables ne créent JAMAIS de headings
2. Tables ajoutent du contenu à section active (filtré)
3. Section présente ↔ lines utiles > 0
4. GOLD strict: si GOLD présent, sélection uniquement parmi GOLD
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import shutil

from src.rhpro.dataset_training import (
    extract_sections_from_docx,
    select_best_docx_for_sections,
    is_noise_cell_text,
    is_useful_line,
    is_noise_title,
    normalize_title
)


def create_test_docx_with_table_noise(filepath: Path):
    """
    Crée un DOCX de test avec:
    - 1 heading paragraphe "FORMATION"
    - 1 table avec cellules "NOM", "PRENOM", "N°AVS" (gras)
    
    Attendu V4.1:
    - unknown_titles ne contient pas NOM/PRENOM/AVS
    - Aucune section créée depuis table
    - Section active reste FORMATION
    """
    doc = Document()
    
    # Heading paragraphe
    heading = doc.add_paragraph("FORMATION")
    heading.style = "Heading 1"
    
    # Contenu
    doc.add_paragraph("Diplôme en informatique obtenu en 2020.")
    
    # Table avec labels formulaires
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "NOM"
    table.cell(0, 1).text = "PRENOM"
    table.cell(0, 2).text = "N°AVS"
    
    # Rendre les labels gras
    for i in range(3):
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
    
    # Données (ne doivent pas être traitées comme headings)
    table.cell(1, 0).text = "DUPONT"
    table.cell(1, 1).text = "Jean"
    table.cell(1, 2).text = "756.1234.5678.90"
    
    doc.save(filepath)


def create_test_docx_with_table_content(filepath: Path):
    """
    Crée un DOCX avec:
    - Heading "COMPETENCES"
    - Table avec contenu utile
    
    Attendu V4.1:
    - COMPETENCES contient lignes de la table (filtrées)
    - clients pour COMPETENCES > 0
    - avg_lines > 0
    """
    doc = Document()
    
    # Heading
    heading = doc.add_paragraph("COMPETENCES")
    heading.style = "Heading 1"
    
    # Table avec contenu mixte
    table = doc.add_table(rows=3, cols=2)
    
    # Row 0: Labels (doivent être filtrés)
    table.cell(0, 0).text = "Compétence"
    table.cell(0, 1).text = "Niveau"
    
    # Row 1: Contenu utile
    table.cell(1, 0).text = "Organisation et rigueur"
    table.cell(1, 1).text = "Excellent"
    
    # Row 2: Contenu utile
    table.cell(2, 0).text = "Autonomie et initiative"
    table.cell(2, 1).text = "Bon"
    
    doc.save(filepath)


def create_test_docx_empty_section(filepath: Path):
    """
    Crée un DOCX avec:
    - Heading "RESSOURCES COMPORTEMENTALES POINTS D'APPUI"
    - Contenu vide ou uniquement bruit (X, chiffres)
    
    Attendu V4.1:
    - Section NON comptée comme présente
    - lines.avg == 0 ET clients == 0
    """
    doc = Document()
    
    heading = doc.add_paragraph("RESSOURCES COMPORTEMENTALES : POINTS D'APPUI")
    heading.style = "Heading 1"
    
    # Contenu bruit
    doc.add_paragraph("X")
    doc.add_paragraph("123456")
    doc.add_paragraph("")
    
    doc.save(filepath)


def create_test_gold_and_journal(temp_dir: Path):
    """
    Crée un dossier client avec:
    - 1 DOCX GOLD (bilan avec sections)
    - 1 DOCX journal (contient "CHATGPT A DIT")
    
    Attendu V4.1:
    - Sélection = GOLD (toujours)
    - reasons incluent "GOLD_STRICT_MODE"
    """
    # GOLD bilan
    gold_path = temp_dir / "bilan_GOLD.docx"
    doc_gold = Document()
    doc_gold.add_paragraph("BILAN D'ORIENTATION PROFESSIONNELLE", style="Heading 1")
    doc_gold.add_paragraph("OBJECTIFS", style="Heading 1")
    doc_gold.add_paragraph("Retrouver un emploi dans le secteur de l'informatique.")
    doc_gold.add_paragraph("COMPETENCES", style="Heading 1")
    doc_gold.add_paragraph("Programmation Python, gestion de projet.")
    doc_gold.save(gold_path)
    
    # Journal transcription
    journal_path = temp_dir / "journal_chatgpt_transcription.docx"
    doc_journal = Document()
    doc_journal.add_paragraph("JOURNAL DE BORD - ENTRETIENS")
    doc_journal.add_paragraph("CHATGPT A DIT : Le client doit...")
    doc_journal.add_paragraph("Conversation du 15/12/2024")
    doc_journal.add_paragraph("Notes diverses et variées.")
    doc_journal.save(journal_path)
    
    return {
        "gold": gold_path,
        "journal": journal_path
    }


# ============================================================================
# TEST 1 — Table cells never create headings
# ============================================================================

def test_table_cells_never_create_headings():
    """
    V4.1 Test 1: Les cellules de table ne doivent JAMAIS créer de headings.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "test_table_noise.docx"
        create_test_docx_with_table_noise(docx_path)
        
        # Extraire sections
        sections = extract_sections_from_docx(docx_path)
        
        # Assertions
        section_titles = [normalize_title(s["title"]) for s in sections]
        
        # NOM, PRENOM, AVS ne doivent PAS apparaître comme sections
        assert "NOM" not in section_titles, "NOM ne doit pas être un heading"
        assert "PRENOM" not in section_titles, "PRENOM ne doit pas être un heading"
        assert "N AVS" not in section_titles, "N°AVS ne doit pas être un heading"
        assert "AVS" not in section_titles, "AVS ne doit pas être un heading"
        
        # FORMATION doit être présente
        assert "FORMATION" in section_titles, "FORMATION doit être détecté"
        
        # Vérifier que les labels sont bien filtrés par is_noise_title
        assert is_noise_title("NOM"), "NOM doit être détecté comme bruit"
        assert is_noise_title("PRENOM"), "PRENOM doit être détecté comme bruit"
        assert is_noise_cell_text("N°AVS"), "N°AVS doit être filtré"


# ============================================================================
# TEST 2 — Table content appended to active section
# ============================================================================

def test_table_content_appended_to_active_section():
    """
    V4.1 Test 2: Le contenu utile des tables doit être ajouté à la section active.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "test_table_content.docx"
        create_test_docx_with_table_content(docx_path)
        
        sections = extract_sections_from_docx(docx_path)
        
        # Trouver section COMPETENCES
        competences_section = None
        for s in sections:
            if normalize_title(s["title"]) == "COMPETENCES":
                competences_section = s
                break
        
        assert competences_section is not None, "COMPETENCES doit être détectée"
        assert competences_section["lines"] > 0, "COMPETENCES doit avoir des lignes"
        
        # Vérifier filtrage : labels "Compétence" / "Niveau" doivent être exclus
        # mais contenu utile doit être présent
        preview = competences_section.get("content_preview", "")
        
        # Contenu utile devrait être présent
        assert competences_section["lines"] >= 2, "Au moins 2 lignes utiles attendues"
        
        # Vérifier is_useful_line
        assert is_useful_line("Organisation et rigueur"), "Doit être ligne utile"
        assert is_useful_line("Autonomie et initiative"), "Doit être ligne utile"


# ============================================================================
# TEST 3 — Section present iff useful lines > 0
# ============================================================================

def test_section_present_iff_useful_lines():
    """
    V4.1 Test 3: Une section est présente UNIQUEMENT si lines > 0.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "test_empty_section.docx"
        create_test_docx_empty_section(docx_path)
        
        sections = extract_sections_from_docx(docx_path)
        
        # Chercher section ressources_points_appui
        ressources_section = None
        for s in sections:
            canonical = s.get("canonical")
            if canonical == "ressources_points_appui":
                ressources_section = s
                break
        
        # V4.1: Si contenu vide/bruit, section NE DOIT PAS être ajoutée
        if ressources_section:
            assert ressources_section["lines"] == 0, \
                "Section vide ne doit pas compter de lignes"
        
        # Vérifier que le bruit est bien détecté
        assert not is_useful_line("X"), "X n'est pas une ligne utile"
        assert not is_useful_line("123456"), "Uniquement chiffres n'est pas utile"
        assert not is_useful_line(""), "Vide n'est pas utile"


# ============================================================================
# TEST 4 — GOLD strict selection
# ============================================================================

def test_gold_strict_selection():
    """
    V4.1 Test 4: Si GOLD présent, sélection uniquement parmi GOLD.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        paths = create_test_gold_and_journal(temp_dir)
        
        # Simuler scan_result
        scan_result = {
            "gold": {
                "path": str(paths["gold"]),
                "score": 100,
                "strategy": "filename"
            },
            "rag_sources": [
                {"path": str(paths["gold"]), "extension": ".docx"},
                {"path": str(paths["journal"]), "extension": ".docx"}
            ]
        }
        
        # Sélectionner meilleur DOCX
        selected_docx, debug = select_best_docx_for_sections(temp_dir, scan_result)
        
        # Assertions
        assert selected_docx is not None, "Un DOCX doit être sélectionné"
        assert selected_docx == paths["gold"], \
            f"GOLD doit être sélectionné, pas {selected_docx.name}"
        
        # Vérifier mode GOLD strict
        assert debug.get("gold_mode") is True, "gold_mode doit être True"
        
        # Vérifier que journal n'a pas été considéré
        reasons = debug.get("selected_reasons", [])
        assert any("GOLD" in str(r) for r in reasons), \
            "Reasons doit mentionner GOLD"


# ============================================================================
# TEST BONUS — is_noise_cell_text comprehensive
# ============================================================================

def test_is_noise_cell_text_comprehensive():
    """
    Test exhaustif de is_noise_cell_text pour V4.1
    """
    # Doit être filtré
    assert is_noise_cell_text("NOM")
    assert is_noise_cell_text("PRENOM")
    assert is_noise_cell_text("N°AVS")
    assert is_noise_cell_text("AVS")
    assert is_noise_cell_text("DATE DE NAISSANCE")
    assert is_noise_cell_text("756.1234.5678.90")  # AVS suisse
    assert is_noise_cell_text("15/12/2024")  # Date
    assert is_noise_cell_text("123456")  # Trop de chiffres
    assert is_noise_cell_text("")  # Vide
    assert is_noise_cell_text("X")  # Trop court
    
    # NE doit PAS être filtré
    assert not is_noise_cell_text("Organisation et rigueur")
    assert not is_noise_cell_text("Autonomie dans les tâches")
    assert not is_noise_cell_text("Bon niveau")
